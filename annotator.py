#!/usr/bin/env python3
"""
Manual Speech Annotation Tool
================================

Build ground-truth CSV files compatible with app.py.

Workflow:
  1. Pick language -> wordlist auto-loads from wordlists/<iso>.txt
     (or from the matching Dr__Abhishek_..._Text_Material.docx the first time).
  2. Load a .wav file -> waveform appears.
  3. Mark a region. Three ways:
       a. DRAG  on the waveform (click + drag horizontally)
       b. CLICK on the waveform: LEFT-click = set start, RIGHT-click = set end
       c. TYPE  numeric seconds into the Start / End spinboxes
  4. Pick a label: select from the word list, click a quick-button
     (s1..s5, paragraph, or any char button), or type freeform.
  5. New word? Type it in the "Add new word" box and click Add — it's
     appended to wordlists/<iso>.txt and stays for next session.
  6. Click "Add Segment" -> appears in the bottom list.
  7. Click "Save CSV" -> writes a 3-column CSV (label, start, end) with
     no header — exactly what app.py's "Load Ground Truth" expects.
  8. Optionally click "Export transcripts/<iso>.txt" to also produce the
     sentence/paragraph reference file your app.py uses for accuracy
     scoring on s1..s5 / paragraph labels.

Dependencies:
    pip install PyQt5 matplotlib numpy soundfile sounddevice python-docx scipy
    (librosa is an optional fallback if soundfile is missing)
"""

import os
import re
import sys
import csv
import time
import traceback
from pathlib import Path
from functools import lru_cache

import numpy as np
import sounddevice as sd

import matplotlib
matplotlib.use("QtAgg")
import matplotlib.pyplot as plt
from matplotlib import font_manager as fm
from matplotlib.patches import Rectangle
from matplotlib.backends.backend_qtagg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.backends.backend_qtagg import NavigationToolbar2QT as NavigationToolbar

from PyQt5.QtCore import Qt, QThread, pyqtSignal, QTimer
from PyQt5.QtGui import QColor, QFont, QKeySequence
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QFileDialog, QMessageBox,
    QVBoxLayout, QHBoxLayout, QGridLayout, QLabel, QPushButton, QComboBox,
    QListWidget, QListWidgetItem, QSplitter, QGroupBox, QDoubleSpinBox,
    QLineEdit, QTabWidget, QFrame, QScrollArea, QInputDialog, QCheckBox,
    QScrollBar, QShortcut, QDialog, QDialogButtonBox, QRadioButton,
    QButtonGroup, QPlainTextEdit, QSpinBox
)

# ---------- Optional deps ----------
try:
    import soundfile as sf
    _SF = True
except ImportError:
    _SF = False

try:
    from scipy.signal import resample_poly, stft
    from math import gcd
    _SCIPY = True
except ImportError:
    _SCIPY = False

try:
    from docx import Document
    _DOCX = True
except ImportError:
    _DOCX = False

try:
    import openpyxl
    _XLSX = True
except ImportError:
    _XLSX = False

try:
    from pypdf import PdfReader
    _PDF = True
except ImportError:
    try:
        from PyPDF2 import PdfReader   # type: ignore
        _PDF = True
    except ImportError:
        _PDF = False

# ASR backends — optional. If neither is installed, auto-detect is disabled.
try:
    import torch
    import torchaudio  # noqa: F401  (sometimes required transitively)
    from transformers import (
        AutoModel,
        AutoModelForSpeechSeq2Seq,
        AutoProcessor,
    )
    _HF = True
except Exception:
    _HF = False

try:
    from faster_whisper import WhisperModel as _FWModel
    _FW = True
except Exception:
    _FW = False

_ASR_AVAILABLE = _HF or _FW


# ---------- Constants ----------
APP_DIR = Path(__file__).parent.resolve()
WORDLISTS_DIR = APP_DIR / "wordlists"
TRANSCRIPTS_DIR = APP_DIR / "transcripts"
OUTPUT_DIR = APP_DIR / "annotations"
AUTOSAVE_DIR = OUTPUT_DIR / ".autosave"  # hidden dir; one autosave per audio file

# Audio constants used by the ASR engines (both prefer 16 kHz mono).
SR16 = 16000
CPU_THR = max(1, min(os.cpu_count() or 4, 8))

# Which model to use for each language. Same selection as your app.py:
#   Indic languages (gu/hi/mr) → IndicConformer-600M (IIT Madras / AI4Bharat)
#   English → faster-whisper large-v3 if installed, else HF Whisper large-v3
INDIC_MODEL = ("indicconformer", "ai4bharat/indic-conformer-600m-multilingual") if _HF else None
if _FW:
    ENGLISH_MODEL = ("fw", "large-v3")
elif _HF:
    ENGLISH_MODEL = ("hf_whisper", "openai/whisper-large-v3")
else:
    ENGLISH_MODEL = None

LANGUAGES = [
    ("Gujarati", "gu", "gujarati"),
    ("Hindi", "hi", "hindi"),
    ("Marathi", "mr", "marathi"),
    ("English", "en", "english"),
]

QUICK_LABELS = ["s1", "s2", "s3", "s4", "s5", "paragraph"]

# Multilingual section header keywords for docx parsing
SECTION_KEYWORDS = {
    "chars": [
        "alphabet", "alphabets", "vowels", "consonants",
        "स्वर", "व्यंजन", "वर्णमाला", "मूळाक्षरे", "अक्षरे",
        "મૂળાક્ષરો", "સ્વર", "વ્યંજન",
    ],
    "words": [
        "words", "word",
        "शब्द", "शब्दों",
        "શબ્દો", "શબ્દ",
    ],
    "sentences": [
        "sentences", "sentence",
        "वाक्य", "वाक्ये", "वाक्यांश",
        "વાક્યો", "વાક્ય",
    ],
    "paragraph": [
        "paragraph", "passage",
        "अनुच्छेद", "गद्यांश", "परिच्छेद",
        "ફકરો", "પેરેગ્રાફ", "ગદ્યાંશ",
    ],
}


def log(msg):
    print(f"[{time.strftime('%H:%M:%S')}] {msg}", flush=True)


# ---------- Time formatting (HH:MM:SS.fff) ----------
def format_hms(seconds, with_ms=True):
    """Format seconds as 'HH:MM:SS.fff' (or 'HH:MM:SS' if with_ms=False).

    Examples:
        0.5          -> '00:00:00.500'
        65.123       -> '00:01:05.123'
        3725.0       -> '01:02:05.000'
    """
    if seconds is None:
        return ""
    try:
        s_val = float(seconds)
    except (TypeError, ValueError):
        return ""
    sign = "-" if s_val < 0 else ""
    s_val = abs(s_val)
    h = int(s_val // 3600)
    m = int((s_val % 3600) // 60)
    s = s_val - h * 3600 - m * 60
    if with_ms:
        # 06.3f -> 'SS.fff' (width 6: 2 digits + . + 3 decimals)
        return f"{sign}{h:02d}:{m:02d}:{s:06.3f}"
    return f"{sign}{h:02d}:{m:02d}:{int(round(s)):02d}"


def parse_hms(text):
    """Parse a time string into seconds. Accepts:
       'HH:MM:SS.fff' / 'HH:MM:SS'
       'MM:SS.fff'    / 'MM:SS'
       plain numeric seconds (e.g. '1.234' or '0.5')
    Returns None on failure."""
    if text is None:
        return None
    s = str(text).strip().replace(",", ".")  # tolerate ',' decimals
    if not s:
        return None
    sign = 1.0
    if s.startswith("-"):
        sign = -1.0
        s = s[1:]
    # Plain float?
    try:
        return sign * float(s)
    except ValueError:
        pass
    # HH:MM:SS variants
    parts = s.split(":")
    try:
        if len(parts) == 3:
            return sign * (int(parts[0]) * 3600 + int(parts[1]) * 60 + float(parts[2]))
        if len(parts) == 2:
            return sign * (int(parts[0]) * 60 + float(parts[1]))
    except ValueError:
        pass
    return None


def load_audio(path):
    """Load audio as float32 mono with native sample rate."""
    if _SF:
        try:
            data, sr = sf.read(str(path), dtype="float32", always_2d=True)
            y = data.mean(axis=1) if data.shape[1] > 1 else data[:, 0]
            return y.astype(np.float32), sr
        except Exception:
            pass
    import librosa
    y, sr = librosa.load(str(path), sr=None, mono=True)
    return y.astype(np.float32), sr


def resample_audio(y, src_sr, tgt_sr):
    """Resample to target rate using scipy.signal.resample_poly, librosa fallback."""
    if src_sr == tgt_sr:
        return y.astype(np.float32)
    if _SCIPY:
        g = gcd(int(src_sr), int(tgt_sr))
        up = tgt_sr // g
        down = src_sr // g
        return resample_poly(y, up, down).astype(np.float32)
    import librosa
    return librosa.resample(y, orig_sr=src_sr, target_sr=tgt_sr).astype(np.float32)


# ---------- Voice Activity Detection (energy-based, same approach as app.py) ----------
def vad_segment(y, sr, silence_ms=300, thresh_db=-40, min_ms=200, max_ms=10000):
    """Split audio into speech segments by short-time RMS energy.
    Returns list of {'start': sec, 'end': sec}."""
    frame = int(sr * 0.025)
    hop = int(sr * 0.010)
    n_frames = 1 + (len(y) - frame) // hop
    if n_frames <= 0:
        return []
    idx = np.arange(frame)[None, :] + np.arange(n_frames)[:, None] * hop
    np.clip(idx, 0, len(y) - 1, out=idx)
    rms = np.sqrt(np.mean(y[idx] ** 2, axis=1) + 1e-12)
    rms_db = 20 * np.log10(rms / (rms.max() + 1e-12) + 1e-12)
    voiced = rms_db > thresh_db
    padded = np.concatenate([[False], voiced, [False]])
    diffs = np.diff(padded.astype(np.int8))
    starts, ends = np.where(diffs == 1)[0], np.where(diffs == -1)[0]
    sil_f = max(1, int(silence_ms / 1000 * sr / hop))
    min_f = max(1, int(min_ms / 1000 * sr / hop))
    max_f = max(1, int(max_ms / 1000 * sr / hop))
    merged = []
    i = 0
    while i < len(starts):
        s, e = starts[i], ends[i]
        while i + 1 < len(starts) and (starts[i + 1] - e) < sil_f:
            i += 1
            e = ends[i]
        merged.append((s, e))
        i += 1
    out = []
    for s, e in merged:
        if (e - s) < min_f:
            continue
        for cs in range(s, e, max_f):
            ce = min(cs + max_f, e)
            if (ce - cs) >= min_f:
                out.append({
                    "start": round(cs * hop / sr, 4),
                    "end":   round(min(ce * hop / sr, len(y) / sr), 4),
                })
    return out


# ---------- Light text cleanup for ASR output ----------
_INDIC_RE = r"\u0900-\u097F\uA8E0-\uA8FF\u0A80-\u0AFF"
_RE_CLEAN = re.compile(rf"[^A-Za-z0-9\s{_INDIC_RE}.,!?'\-]", re.UNICODE)
_RE_SPACES = re.compile(r"\s+")


def clean_asr_text(text):
    """Normalize whitespace + strip unusual punctuation from ASR output."""
    if not text:
        return ""
    t = _RE_CLEAN.sub("", str(text).strip())
    return _RE_SPACES.sub(" ", t).strip()


# ---------- Segment type classification (adapted from app.py) ----------
# Decide whether a label is a single character, a single word, a full
# sentence, or a paragraph. Sequence labels (s1..s5, p1..p3, paragraph)
# are detected so we can tag them specially on the waveform.
_RE_SEQ_SENT = re.compile(r"^s\d+$", re.I)            # s1, s2, ...
_RE_SEQ_PARA = re.compile(r"^p\d+$", re.I)            # p1, p2, ... (partial para)
_RE_NORM_CMP = re.compile(rf"[^a-z0-9{_INDIC_RE}]", re.UNICODE)


def _norm_for_compare(s):
    """Lowercase + strip everything except letters/digits, for fair comparison."""
    return _RE_NORM_CMP.sub("", str(s).strip().lower())


def classify_segment_type(label):
    """Return one of: 'char', 'word', 'sentence', 'paragraph'.
    Mirrors the logic in app.py's classify_segment(), minus the GT lookup.
    """
    label = str(label).strip()
    ll = label.lower()
    # Explicit sequence labels
    if _RE_SEQ_SENT.match(label):
        return "sentence"
    if _RE_SEQ_PARA.match(label) or "para" in ll or "passage" in ll:
        return "paragraph"
    # Heuristic by content
    chars = re.sub(r"\s+", "", label)
    if len(chars) == 1:
        return "char"
    if len(label.split()) == 1:
        return "word"
    # Multi-word: sentence vs paragraph by length
    n_words = len(label.split())
    return "paragraph" if n_words > 25 else "sentence"


# Short type tag shown on the waveform per segment. For sequence labels we
# preserve the s1/s2/p1/p2 identity; for plain content we show char/word/sent.
def short_type_tag(seg, ref_sentences=None, ref_paragraph=None):
    """Build the on-waveform tag string for a segment.
    Examples: 'char', 'word', 's3', 'p2 (40%)', 'sent'.
    `seg` may carry a '_para_fraction' set by the partial-paragraph matcher.
    """
    label = str(seg.get("label", "")).strip()
    # Explicit sequence labels keep their identity
    m_s = _RE_SEQ_SENT.match(label)
    if m_s:
        return label.lower()
    m_p = _RE_SEQ_PARA.match(label)
    if m_p:
        frac = seg.get("_para_fraction")
        if frac is not None and frac < 0.99:
            return f"{label.lower()} ({frac*100:.0f}%)"
        return label.lower()
    t = classify_segment_type(label)
    if t == "char":
        return "char"
    if t == "word":
        return "word"
    if t == "paragraph":
        frac = seg.get("_para_fraction")
        if frac is not None and frac < 0.99:
            return f"para ({frac*100:.0f}%)"
        return "para"
    return "sent"


def _fuzzy_score(a, b):
    """Normalised similarity in [0, 1] using character edit distance.
    Case- and whitespace-tolerant. Returns 1.0 on exact match (after normalize)."""
    if a is None or b is None:
        return 0.0
    a = re.sub(r"\s+", "", str(a).strip().lower())
    b = re.sub(r"\s+", "", str(b).strip().lower())
    if not a and not b:
        return 1.0
    if not a or not b:
        return 0.0
    if a == b:
        return 1.0
    # Classic edit distance, O(len(a)*len(b)) but our strings are short
    na, nb = len(a), len(b)
    dp = list(range(nb + 1))
    for i in range(1, na + 1):
        prev, dp[0] = dp[0], i
        for j in range(1, nb + 1):
            tmp = dp[j]
            cost = 0 if a[i - 1] == b[j - 1] else 1
            dp[j] = min(prev + cost, dp[j] + 1, dp[j - 1] + 1)
            prev = tmp
    return max(0.0, 1.0 - dp[nb] / max(na, nb))


def measure_paragraph_progress(prediction, reference_paragraph):
    """Given the model's prediction for a (possibly partial) paragraph read,
    estimate how far through the reference paragraph the speaker got.

    Returns dict:
        spoken_words   : int  — how many reference words appear to be spoken
        total_words    : int  — total words in the reference paragraph
        fraction       : float in [0, 1]
        last_word      : str  — the last reference word that was matched
        matched_text   : str  — the reference text up to the spoken point

    Strategy: align the predicted word sequence against the reference word
    sequence with a longest-common-subsequence-style sweep. The spoken extent
    is the index of the furthest reference word that has a fuzzy match to some
    predicted word, scanning left to right.
    """
    ref_words = (reference_paragraph or "").split()
    pred_words = (prediction or "").split()
    total = len(ref_words)
    if total == 0:
        return {"spoken_words": 0, "total_words": 0, "fraction": 0.0,
                "last_word": "", "matched_text": ""}
    if not pred_words:
        return {"spoken_words": 0, "total_words": total, "fraction": 0.0,
                "last_word": "", "matched_text": ""}

    ref_norm = [_norm_for_compare(w) for w in ref_words]
    pred_norm = [_norm_for_compare(w) for w in pred_words]
    pred_set_index = {}
    for idx, w in enumerate(pred_norm):
        pred_set_index.setdefault(w, []).append(idx)

    # Sweep reference words; a ref word is "spoken" if it fuzzy-matches a
    # predicted word at a monotonically non-decreasing predicted position.
    last_matched_ref = -1
    pred_cursor = 0
    for r_idx, rw in enumerate(ref_norm):
        if not rw:
            continue
        # Look for rw among predicted words at/after pred_cursor
        best_j = -1
        for j in range(pred_cursor, len(pred_norm)):
            if not pred_norm[j]:
                continue
            if rw == pred_norm[j] or _fuzzy_score(rw, pred_norm[j]) >= 0.80:
                best_j = j
                break
        if best_j >= 0:
            last_matched_ref = r_idx
            pred_cursor = best_j + 1

    spoken = last_matched_ref + 1
    spoken = max(0, min(spoken, total))
    frac = spoken / total if total else 0.0
    return {
        "spoken_words": spoken,
        "total_words": total,
        "fraction": frac,
        "last_word": ref_words[last_matched_ref] if last_matched_ref >= 0 else "",
        "matched_text": " ".join(ref_words[:spoken]),
    }


# ---------- ASR Engine (Whisper + IndicConformer, mirrors app.py) ----------
class AsrEngine:
    """Thin wrapper around faster-whisper / HF Whisper / IndicConformer.
    Loads one model at a time; switching languages may require reloading.
    """
    def __init__(self):
        self.backend = None
        self.processor = None
        self.device = None
        self.engine_key = None
        self._ck = None  # (engine_key, model_id) currently loaded

    @property
    def is_loaded(self):
        return self.backend is not None

    @property
    def loaded_key(self):
        return self._ck

    def load(self, engine_key, model_id):
        ck = (engine_key, model_id)
        if self._ck == ck and self.backend is not None:
            return True, f"Already loaded: {model_id}", 0.0
        t0 = time.time()
        self.backend = None
        self.processor = None
        self.engine_key = engine_key
        try:
            if engine_key == "indicconformer":
                self._load_indic(model_id)
            elif engine_key == "fw":
                self._load_fw(model_id)
            elif engine_key == "hf_whisper":
                self._load_hf_whisper(model_id)
            else:
                return False, f"Unknown engine: {engine_key}", 0.0
            self._ck = ck
            elapsed = time.time() - t0
            return True, f"Loaded {model_id} in {elapsed:.1f}s", elapsed
        except Exception as e:
            traceback.print_exc()
            self.backend = None
            return False, str(e), time.time() - t0

    def _load_indic(self, model_id):
        if not _HF:
            raise RuntimeError("transformers/torch not installed")
        self.device = "cuda" if torch.cuda.is_available() else "cpu"
        log(f"Loading IndicConformer on {self.device}…")
        self.backend = AutoModel.from_pretrained(model_id, trust_remote_code=True)
        if hasattr(self.backend, "to"):
            try:
                self.backend = self.backend.to(self.device)
            except Exception:
                self.device = "cpu"

    def _load_fw(self, size):
        if not _FW:
            raise RuntimeError("faster-whisper not installed")
        has_cuda = (_HF and torch.cuda.is_available()) if _HF else False
        log(f"Loading faster-whisper '{size}' (cuda={has_cuda})…")
        self.backend = _FWModel(
            size,
            device="cuda" if has_cuda else "cpu",
            compute_type="float16" if has_cuda else "int8",
            cpu_threads=CPU_THR,
        )

    def _load_hf_whisper(self, model_id):
        if not _HF:
            raise RuntimeError("transformers/torch not installed")
        dv = "cuda:0" if torch.cuda.is_available() else "cpu"
        dt = torch.float16 if torch.cuda.is_available() else torch.float32
        log(f"Loading HF Whisper '{model_id}' on {dv}…")
        self.backend = AutoModelForSpeechSeq2Seq.from_pretrained(
            model_id, torch_dtype=dt, low_cpu_mem_usage=True, use_safetensors=True,
        ).to(dv)
        self.processor = AutoProcessor.from_pretrained(model_id)
        self.device = dv

    def transcribe(self, audio_16k, lang_iso, lang_name):
        """Run ASR on a mono float32 numpy array at 16 kHz. Returns predicted text."""
        if self.backend is None or audio_16k is None or len(audio_16k) == 0:
            return ""
        # Pad short clips so the model doesn't choke
        if len(audio_16k) < int(0.1 * SR16):
            audio_16k = np.pad(audio_16k, (0, int(0.1 * SR16) - len(audio_16k)))
        a = audio_16k.astype(np.float32)
        if self.engine_key == "indicconformer":
            return self._tx_indic(a, lang_iso)
        if self.engine_key == "fw":
            return self._tx_fw(a, lang_iso)
        if self.engine_key == "hf_whisper":
            return self._tx_hf(a, lang_name)
        return ""

    def transcribe_with_words(self, audio_16k, lang_iso, lang_name,
                              clip_offset=0.0):
        """Run ASR and also return per-word timestamps as
        [(word, start_sec, end_sec), ...], in audio time.

        `clip_offset` is added to every word's start/end so the times are
        in absolute audio coordinates (caller passes the segment start).

        Returns (full_text, [(word, start, end), ...]).
        For backends that don't expose word timestamps (IndicConformer,
        HF Whisper without timestamps=True), word times are *estimated*
        by evenly distributing the segment's duration across the predicted
        words. Estimated word times are flagged via a parallel
        `timestamps_are_estimated` boolean on the engine for the caller.
        """
        if self.backend is None or audio_16k is None or len(audio_16k) == 0:
            return "", []
        if len(audio_16k) < int(0.1 * SR16):
            audio_16k = np.pad(audio_16k, (0, int(0.1 * SR16) - len(audio_16k)))
        a = audio_16k.astype(np.float32)
        dur_s = len(a) / SR16
        if self.engine_key == "fw":
            return self._tx_fw_words(a, lang_iso, clip_offset)
        if self.engine_key == "indicconformer":
            text = self._tx_indic(a, lang_iso)
            return text, _estimate_word_times(text, clip_offset, dur_s)
        if self.engine_key == "hf_whisper":
            text = self._tx_hf(a, lang_name)
            return text, _estimate_word_times(text, clip_offset, dur_s)
        return "", []

    def _tx_indic(self, a, lang_iso):
        w = torch.tensor(a, dtype=torch.float32).unsqueeze(0)
        if self.device and self.device != "cpu":
            w = w.to(self.device)
        with torch.inference_mode():
            t = self.backend(w, lang_iso, "ctc")
        if isinstance(t, (list, tuple)):
            t = t[0] if t else ""
        return clean_asr_text(str(t))

    def _tx_fw(self, a, lang_iso):
        segments, _ = self.backend.transcribe(
            a, language=lang_iso, task="transcribe",
            beam_size=1, best_of=1,
            vad_filter=False, without_timestamps=True,
            condition_on_previous_text=False,
        )
        text = " ".join(s.text.strip() for s in list(segments) if s.text.strip())
        return clean_asr_text(text)

    def _tx_fw_words(self, a, lang_iso, clip_offset):
        """faster-whisper with native word_timestamps=True.
        Returns (full_text, [(word, abs_start, abs_end), ...])."""
        segments, _ = self.backend.transcribe(
            a, language=lang_iso, task="transcribe",
            beam_size=1, best_of=1,
            vad_filter=False, word_timestamps=True,
            condition_on_previous_text=False,
        )
        words_out = []
        text_parts = []
        for seg in list(segments):
            text_parts.append(seg.text.strip())
            for w in (getattr(seg, "words", None) or []):
                token = (w.word or "").strip()
                if not token:
                    continue
                # Whisper's word offsets are within the clip; add clip_offset
                # so the result is in absolute audio coordinates.
                words_out.append((
                    token,
                    float(w.start) + float(clip_offset),
                    float(w.end) + float(clip_offset),
                ))
        text = clean_asr_text(" ".join(text_parts))
        return text, words_out

    def _tx_hf(self, a, lang_name):
        inp = self.processor(a, sampling_rate=SR16,
                              return_tensors="pt", return_attention_mask=True)
        ft = inp.input_features.to(self.device)
        mk = getattr(inp, "attention_mask", None)
        if mk is not None:
            mk = mk.to(self.device)
        gk = {"input_features": ft, "task": "transcribe",
              "language": lang_name, "return_timestamps": False}
        if mk is not None:
            gk["attention_mask"] = mk
        with torch.inference_mode():
            ids = self.backend.generate(**gk)
        text = self.processor.batch_decode(ids, skip_special_tokens=True)[0]
        return clean_asr_text(text)


def _estimate_word_times(text, clip_offset, clip_dur_s):
    """For backends without word-level timestamps, split the clip's duration
    evenly across the predicted words. Returns [(word, start, end), ...].

    This is an approximation — the actual word boundaries inside the clip
    are unknown to us. Useful for displaying *which words were predicted*
    in time order and for rough error estimation.
    """
    text = (text or "").strip()
    if not text or clip_dur_s <= 0:
        return []
    # Indic + English alike: split on whitespace
    words = [w for w in text.split() if w]
    n = len(words)
    if n == 0:
        return []
    per = clip_dur_s / n
    out = []
    for i, w in enumerate(words):
        s = clip_offset + i * per
        e = clip_offset + (i + 1) * per
        out.append((w, s, e))
    return out


ENGINE = AsrEngine()


# ---------- Spectrogram backends (linear, log dB, mel) ----------
def _hz_to_mel(hz):
    return 2595.0 * np.log10(1.0 + np.asarray(hz, dtype=np.float64) / 700.0)


def _mel_to_hz(mel):
    return 700.0 * (10.0 ** (np.asarray(mel, dtype=np.float64) / 2595.0) - 1.0)


@lru_cache(maxsize=8)
def _mel_filterbank(sr, n_fft, n_mels=80, fmin=0.0, fmax=None):
    """Triangular mel filterbank, shape: (n_mels, n_fft//2 + 1)."""
    if fmax is None:
        fmax = sr / 2.0
    n_bins = n_fft // 2 + 1
    mel_points = np.linspace(_hz_to_mel(fmin), _hz_to_mel(fmax), n_mels + 2)
    hz_points = _mel_to_hz(mel_points)
    bin_pts = np.floor((n_fft + 1) * hz_points / sr).astype(int)
    bin_pts = np.clip(bin_pts, 0, n_bins - 1)
    fb = np.zeros((n_mels, n_bins), dtype=np.float32)
    for m in range(1, n_mels + 1):
        l, c, r = bin_pts[m - 1], bin_pts[m], bin_pts[m + 1]
        if c > l:
            fb[m - 1, l:c + 1] = np.linspace(0, 1, c - l + 1, dtype=np.float32)
        if r > c:
            fb[m - 1, c:r + 1] = np.linspace(1, 0, r - c + 1, dtype=np.float32)
    return fb


def compute_spectrogram(y, sr, mode="log_db", n_fft=1024, hop=256, n_mels=80):
    """Compute one of three spectrogram views.

    mode='linear'  → magnitude spectrogram on a linear amplitude scale
    mode='log_db'  → 20·log10(magnitude) on linear-frequency axis
    mode='mel'     → log-magnitude mel spectrogram (dB), 80 mel bins

    Returns (freqs, times, S) where S has shape (len(freqs), len(times)).
    """
    if not _SCIPY:
        raise RuntimeError("scipy is required for spectrograms — pip install scipy")
    # scipy.signal.stft handles padding/windowing for us
    f, t, Z = stft(y.astype(np.float64), fs=sr,
                   nperseg=n_fft, noverlap=n_fft - hop,
                   padded=False, boundary=None, window="hann")
    mag = np.abs(Z).astype(np.float32)
    if mode == "linear":
        return f, t, mag
    if mode == "log_db":
        return f, t, (20.0 * np.log10(mag + 1e-8)).astype(np.float32)
    if mode == "mel":
        fb = _mel_filterbank(sr, n_fft, n_mels=n_mels)  # (n_mels, n_bins)
        mel = fb @ mag                                    # (n_mels, n_t)
        mel_db = (20.0 * np.log10(mel + 1e-8)).astype(np.float32)
        # Use approximate Hz centers for the y-axis instead of bin indices
        mel_pts = np.linspace(_hz_to_mel(0), _hz_to_mel(sr / 2.0), n_mels + 2)
        mel_centers = _mel_to_hz(mel_pts[1:-1])
        return mel_centers, t, mel_db
    raise ValueError(f"Unknown spectrogram mode: {mode}")


# ---------- DOCX text material parser ----------
def _extract_docx_paragraphs(path):
    """Return list of cleaned non-empty paragraph strings from a .docx.

    Word stores Shift+Enter "soft" line breaks as `\\n` inside a single
    paragraph object — but each visual line is semantically a separate
    entry for our purposes (header on one line, body on the next). We
    split on `\\n` so the section parser sees the right structure.
    Also reads text inside tables so docs that put content in a 1-cell
    table (like the English file's blank table) don't drop anything.
    """
    if not _DOCX:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")
    doc = Document(str(path))
    out = []

    def _push(raw):
        if not raw:
            return
        for sub in raw.split("\n"):
            sub = sub.strip()
            if sub:
                out.append(sub)

    for p in doc.paragraphs:
        _push(p.text)
    # Include any text trapped inside tables, in document order is hard
    # to guarantee with python-docx; for the Dr. Abhishek files the tables
    # are empty, so we just append after paragraphs.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for cp in cell.paragraphs:
                    _push(cp.text)
    return out


def _looks_like_header(line):
    """Strip ornament chars and return a short normalized version, or '' if not header-ish."""
    # Strip markdown stars, parens, colons, brackets, dashes, asterisks
    clean = re.sub(r"[\*_:()\[\]\-\.]", " ", line)
    clean = re.sub(r"\s+", " ", clean).strip().lower()
    # Strip leading "a)", "b)", "part-i", etc.
    clean = re.sub(r"^[a-z]\)\s*", "", clean)
    clean = re.sub(r"^part[\s\-]*[ivx]+\s*", "", clean)
    # Headers are short
    return clean if len(clean) <= 40 else ""


def _classify_header(line):
    """If the line is a section header, return its category (chars/words/sentences/paragraph), else None."""
    norm = _looks_like_header(line)
    if not norm:
        return None
    tokens = set(norm.split())
    for category, kws in SECTION_KEYWORDS.items():
        for kw in kws:
            kw_norm = kw.lower()
            if kw_norm in tokens or kw_norm == norm:
                return category
            # Indic keywords may be one of several tokens
            if any(kw_norm == t for t in tokens):
                return category
    return None


def _extract_pdf_paragraphs(path):
    """Return list of cleaned non-empty paragraph strings from a .pdf.

    PDFs of the Dr. Abhishek text material have one big page of running text
    per section. We split each page's extracted text into lines, drop empties,
    and treat each line as a paragraph — same shape as the docx extractor.

    PDF extraction often runs section headers and body content together on
    one line (e.g. "ગુજરાતી લખાણ સામગ્રી મૂળાક્ષરો અ આ ઇ..."). When we
    spot a known section keyword inside a longer line, we split it so the
    header sits on its own line and the body follows.
    """
    if not _PDF:
        raise RuntimeError(
            "pypdf not installed. Run:  pip install pypdf"
        )
    # Flatten all keywords from SECTION_KEYWORDS for inline splitting
    all_kws = []
    for kws in SECTION_KEYWORDS.values():
        all_kws.extend(kws)
    # Build a regex that captures any of these keywords as a split point.
    # Sort longer first so multi-word keywords win over single tokens.
    all_kws_sorted = sorted(set(all_kws), key=len, reverse=True)
    kw_pat = re.compile(
        "(" + "|".join(re.escape(k) for k in all_kws_sorted) + ")",
        re.IGNORECASE,
    )
    reader = PdfReader(str(path))
    out = []
    for page in reader.pages:
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
        for line in page_text.splitlines():
            line = line.strip()
            if not line:
                continue
            # If a section keyword appears mid-line, split before it so
            # the parser's _classify_header() picks it up as a header.
            parts = kw_pat.split(line)
            # parts: alternating [pre_text, keyword, between_text, keyword, ...]
            buf = ""
            for chunk in parts:
                if not chunk:
                    continue
                if chunk.lower() in [k.lower() for k in all_kws_sorted]:
                    # Flush buffer first, then emit the keyword alone
                    if buf.strip():
                        out.append(buf.strip())
                    out.append(chunk.strip())
                    buf = ""
                else:
                    buf += " " + chunk
            if buf.strip():
                out.append(buf.strip())
    return out


def parse_text_material(path):
    """
    Parse a Dr. Abhishek-style text-material file (.docx or .pdf) and return:
      {
        'chars':     [str],
        'words':     [str],
        'sentences': {'s1': str, 's2': str, ...},
        'paragraph': str,
      }
    Best-effort across English/Hindi/Marathi/Gujarati layouts.
    """
    ext = Path(path).suffix.lower()
    if ext == ".pdf":
        paragraphs = _extract_pdf_paragraphs(path)
    else:
        paragraphs = _extract_docx_paragraphs(path)
    out = {"chars": [], "words": [], "sentences": {}, "paragraph": ""}

    # Find section header lines
    sections = []  # list of (idx, category)
    seen_cats = set()
    for i, line in enumerate(paragraphs):
        cat = _classify_header(line)
        if cat and cat not in seen_cats:
            sections.append((i, cat))
            seen_cats.add(cat)

    if not sections:
        # Fallback: dump everything into words
        for line in paragraphs:
            for tok in re.split(r"[\s,;।]+", line):
                tok = tok.strip(" *_:.-,()")
                if tok and not re.fullmatch(r"[\W_]+", tok):
                    out["words"].append(tok)
        out["words"] = list(dict.fromkeys(out["words"]))
        return out

    n = len(paragraphs)
    for idx, (start_i, cat) in enumerate(sections):
        end_i = sections[idx + 1][0] if idx + 1 < len(sections) else n
        body = paragraphs[start_i + 1 : end_i]
        # Skip lines inside body that are themselves header-like
        # (e.g. "b) व्यंजन:" in Hindi after "a) स्वर:", or "Words (Take a pause)"
        # right after "Read the words twice" — both got classified as section
        # starts but only the first wins; the rest must not pollute body.)
        body = [ln for ln in body if _classify_header(ln) is None]

        if cat == "chars":
            for line in body:
                # Tokens are usually whitespace-separated single chars
                for tok in re.split(r"[\s,;।]+", line):
                    tok = tok.strip(" *_:.-,()")
                    # Allow ligatures (Indic combining chars can make a "char" 2-3 codepoints)
                    if tok and 1 <= len(tok) <= 4 and not re.fullmatch(r"[\W_]+", tok):
                        out["chars"].append(tok)

        elif cat == "words":
            for line in body:
                # Strip leading bullets and stars
                line = re.sub(r"^[\-\*•]+\s*", "", line)
                for tok in re.split(r"[\s,;।]+", line):
                    tok = tok.strip(" *_:.,()")
                    if tok and not re.fullmatch(r"[\W_]+", tok) and len(tok) >= 1:
                        # Skip tokens that are just numbers or just punctuation
                        out["words"].append(tok)

        elif cat == "sentences":
            # Each non-empty line in this section is one sentence. If the
            # line carries an explicit (sN) prefix, use it as the key;
            # otherwise auto-number as s1, s2, s3, ... in document order.
            # This handles all three observed shapes:
            #   "(s1) Each untimely…"       (English, Marathi)
            #   "- पुलिस को देखते ही…"        (legacy bulleted text)
            #   "ધીરુભાઈ અંબાણી એક…"          (Word native list items, Gujarati/Hindi)
            auto_idx = 0
            used_labels = set()
            for line in body:
                line_stripped = line.strip()
                if not line_stripped:
                    continue
                m = re.match(r"^\s*\*?\s*\(?(s\d+)\)?\s*[:\.]?\s*(.+)$",
                             line_stripped, re.IGNORECASE)
                if m and m.group(2).strip():
                    key = m.group(1).lower()
                    text = m.group(2).strip()
                else:
                    # Strip any stray leading bullet character
                    text = re.sub(r"^[\-\*•]\s+", "", line_stripped).strip()
                    if not text:
                        continue
                    auto_idx += 1
                    # Find next free auto key
                    while f"s{auto_idx}" in used_labels:
                        auto_idx += 1
                    key = f"s{auto_idx}"
                used_labels.add(key)
                # If same key appears more than once, append to existing
                if key in out["sentences"]:
                    out["sentences"][key] += " " + text
                else:
                    out["sentences"][key] = text

        elif cat == "paragraph":
            joined = " ".join(l.strip() for l in body if l.strip())
            out["paragraph"] = re.sub(r"\s+", " ", joined).strip()

    # Dedupe preserving order
    out["chars"] = list(dict.fromkeys(out["chars"]))
    out["words"] = list(dict.fromkeys(out["words"]))
    return out


# ---------- Tabular file I/O (CSV / TSV / XLSX) ----------
def _read_table(path):
    """Read CSV / TSV / XLSX / XLS / XLSM and return list-of-lists of strings.
    No header is assumed (matches app.py's `header=None` convention).
    """
    p = Path(path)
    ext = p.suffix.lower()
    rows = []
    if ext in (".xlsx", ".xlsm", ".xls"):
        if not _XLSX:
            raise RuntimeError(
                "openpyxl is required to read Excel files.\n"
                "Install it with:  pip install openpyxl"
            )
        wb = openpyxl.load_workbook(str(p), read_only=True, data_only=True)
        # Use the first sheet, matching app.py's sheet_name=0
        ws = wb.worksheets[0]
        for raw_row in ws.iter_rows(values_only=True):
            # Stop at the first fully-empty row (Excel often pads with blanks)
            if raw_row is None or all(v is None or str(v).strip() == "" for v in raw_row):
                continue
            rows.append(["" if v is None else str(v) for v in raw_row])
        wb.close()
    elif ext == ".tsv":
        with open(p, "r", encoding="utf-8-sig", newline="") as f:
            rows = list(csv.reader(f, delimiter="\t"))
    else:  # .csv and anything else: try CSV
        with open(p, "r", encoding="utf-8-sig", newline="") as f:
            rows = list(csv.reader(f))
    return rows


def _write_table(path, data_rows, header=None):
    """Write rows as CSV / TSV / XLSX based on file extension.
    `header` is optional — if provided, written as the first row."""
    p = Path(path)
    ext = p.suffix.lower()
    if ext in (".xlsx", ".xlsm"):
        if not _XLSX:
            raise RuntimeError(
                "openpyxl is required to write Excel files.\n"
                "Install it with:  pip install openpyxl"
            )
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "annotations"
        if header:
            ws.append(list(header))
            # Bold the header row so it's visually distinct in Excel
            for cell in ws[1]:
                cell.font = openpyxl.styles.Font(bold=True)
        for row in data_rows:
            ws.append(list(row))
        # Reasonable column widths so the file is readable straight away
        widths = {"A": 36, "B": 14, "C": 14, "D": 50}
        for col, w in widths.items():
            if col in ws.column_dimensions:
                ws.column_dimensions[col].width = w
            else:
                ws.column_dimensions[col].width = w
        # Freeze the header row if we wrote one
        if header:
            ws.freeze_panes = "A2"
        wb.save(str(p))
    else:
        delim = "\t" if ext == ".tsv" else ","
        with open(p, "w", newline="", encoding="utf-8-sig") as f:
            w = csv.writer(f, delimiter=delim)
            if header:
                w.writerow(list(header))
            for row in data_rows:
                w.writerow(list(row))


# ---------- Persistence ----------
def ensure_dirs():
    for d in (WORDLISTS_DIR, TRANSCRIPTS_DIR, OUTPUT_DIR, AUTOSAVE_DIR):
        d.mkdir(exist_ok=True, parents=True)


def load_wordlist(iso):
    """Load wordlists/<iso>.txt as a list of strings (preserves order, dedupes)."""
    p = WORDLISTS_DIR / f"{iso}.txt"
    if not p.exists():
        return []
    items = []
    seen = set()
    for line in p.read_text(encoding="utf-8").splitlines():
        w = line.strip()
        if w and w not in seen:
            seen.add(w)
            items.append(w)
    return items


def save_wordlist(iso, words):
    p = WORDLISTS_DIR / f"{iso}.txt"
    p.write_text("\n".join(words) + ("\n" if words else ""), encoding="utf-8")


def save_transcripts_file(iso, sentences, paragraph):
    """Write transcripts/<iso>.txt in the format app.py expects."""
    p = TRANSCRIPTS_DIR / f"{iso}.txt"
    lines = [f"# Reference text for {iso}", ""]
    for key in sorted(sentences.keys(), key=lambda k: (k != "paragraph", k)):
        lines.append(f"[{key}]")
        lines.append(sentences[key])
        lines.append("")
    if paragraph:
        lines.append("[paragraph]")
        lines.append(paragraph)
        lines.append("")
    p.write_text("\n".join(lines), encoding="utf-8")


def load_transcripts_file(iso):
    """Read transcripts/<iso>.txt (same format as app.py).

    Format:
        # Optional comment
        [s1]
        sentence 1 text...
        possibly multi-line
        [s2]
        ...
        [paragraph]
        ...

    Returns (sentences_dict, paragraph_str). Both empty if the file is missing.
    """
    p = TRANSCRIPTS_DIR / f"{iso}.txt"
    if not p.exists():
        return {}, ""
    sentences = {}
    paragraph = ""
    current_key = None
    current_buf = []
    try:
        text = p.read_text(encoding="utf-8")
    except Exception as e:
        log(f"load_transcripts_file({iso}) failed: {e}")
        return {}, ""
    for raw_line in text.splitlines():
        stripped = raw_line.strip()
        if not stripped or stripped.startswith("#"):
            # Section breaks: blank lines end the current section
            continue
        m = re.match(r"^\[(.+?)\]$", stripped)
        if m:
            # Commit previous section before starting the new one
            if current_key is not None:
                _commit_transcript_section(current_key, current_buf,
                                           sentences, lambda v: None)
                if current_key == "paragraph":
                    paragraph = "\n".join(current_buf).strip()
            current_key = m.group(1).strip().lower()
            current_buf = []
        elif current_key is not None:
            current_buf.append(raw_line.rstrip())
    if current_key is not None:
        if current_key == "paragraph":
            paragraph = "\n".join(current_buf).strip()
        else:
            sentences[current_key] = "\n".join(current_buf).strip()
    return sentences, paragraph


def _commit_transcript_section(key, buf, sentences_dict, _para_setter):
    """Internal helper: stash buf into the right place based on key."""
    if not buf:
        return
    text = "\n".join(buf).strip()
    if not text:
        return
    if key == "paragraph":
        # Paragraph is handled by caller (returned separately)
        return
    sentences_dict[key] = text


# ---------- Font helper (for matplotlib waveform labels in Indic) ----------
@lru_cache(maxsize=8)
def get_font(iso):
    fams = {
        "gu": ["Noto Sans Gujarati", "Lohit Gujarati", "Shruti", "Arial Unicode MS"],
        "hi": ["Noto Sans Devanagari", "Lohit Devanagari", "Mangal", "Arial Unicode MS"],
        "mr": ["Noto Sans Devanagari", "Lohit Devanagari", "Mangal", "Arial Unicode MS"],
        "en": ["Arial", "Helvetica", "DejaVu Sans"],
    }.get(iso, ["DejaVu Sans"])
    installed = {f.name for f in fm.fontManager.ttflist}
    for f in fams:
        if f in installed:
            return fm.FontProperties(family=f)
    return None


# ---------- Background threads ----------
class AudioLoadThread(QThread):
    # Emits (y, sr, y16, path). y16 is None if scipy unavailable & sr already 16k.
    done = pyqtSignal(np.ndarray, int, object, str)
    error = pyqtSignal(str)

    def __init__(self, path):
        super().__init__()
        self.path = path

    def run(self):
        try:
            y, sr = load_audio(self.path)
            try:
                y16 = resample_audio(y, sr, SR16) if sr != SR16 else y.astype(np.float32)
            except Exception as e:
                log(f"Resample failed ({e}); ASR will be disabled for this file")
                y16 = None
            self.done.emit(y, sr, y16, self.path)
        except Exception as e:
            traceback.print_exc()
            self.error.emit(str(e))


class ModelLoadThread(QThread):
    done = pyqtSignal(bool, str, float)
    error = pyqtSignal(str)

    def __init__(self, engine_key, model_id):
        super().__init__()
        self.engine_key = engine_key
        self.model_id = model_id

    def run(self):
        try:
            ok, msg, el = ENGINE.load(self.engine_key, self.model_id)
            self.done.emit(ok, msg, el)
        except Exception as e:
            traceback.print_exc()
            self.error.emit(str(e))


class AutoSegmentThread(QThread):
    """Runs VAD then transcribes each chunk so each segment arrives with
    a predicted text label, ready for the user to verify or edit.
    """
    progress = pyqtSignal(int, int)             # done, total
    segment = pyqtSignal(dict)                  # {'label': str, 'start': float, 'end': float}
    done = pyqtSignal(int, float)               # total_segments, elapsed_seconds
    error = pyqtSignal(str)

    def __init__(self, y, sr, y16, lang_iso, lang_name, vad_params, do_asr=True):
        super().__init__()
        self.y = y
        self.sr = sr
        self.y16 = y16
        self.lang_iso = lang_iso
        self.lang_name = lang_name
        self.vad_params = vad_params
        self.do_asr = do_asr
        self._stop = False

    def request_stop(self):
        self._stop = True

    def run(self):
        try:
            t0 = time.time()
            # VAD runs on the *original* audio so timestamps line up exactly
            chunks = vad_segment(self.y, self.sr, **self.vad_params)
            total = len(chunks)
            if total == 0:
                self.done.emit(0, time.time() - t0)
                return
            # Decide whether to transcribe (need an engine + 16k audio)
            run_asr = self.do_asr and ENGINE.is_loaded and self.y16 is not None
            for i, ch in enumerate(chunks):
                if self._stop:
                    break
                pred = ""
                if run_asr:
                    s16 = max(0, int(ch["start"] * SR16))
                    e16 = min(len(self.y16), int(ch["end"] * SR16))
                    if e16 > s16:
                        try:
                            pred = ENGINE.transcribe(
                                self.y16[s16:e16], self.lang_iso, self.lang_name,
                            )
                        except Exception as e:
                            log(f"  ASR failed on chunk {i+1}: {e}")
                # `label` is what the user will see/edit (pre-filled with pred);
                # `prediction` is the *original* model output, kept verbatim so
                # we can always show GT vs PRED side-by-side, even after edits.
                label = pred if pred else f"seg_{i+1:04d}"
                self.segment.emit({
                    "label": label,
                    "start": float(ch["start"]),
                    "end": float(ch["end"]),
                    "prediction": pred,
                })
                self.progress.emit(i + 1, total)
            self.done.emit(total, time.time() - t0)
        except Exception as e:
            traceback.print_exc()
            self.error.emit(str(e))


# ---------- Forced alignment: GT label list -> audio timestamps ----------
# Given a list of expected GT labels (in the order the speaker reads them) and
# the audio, we want to find the timestamps for each. Approach:
#   1. VAD splits the audio into speech chunks (one chunk ≈ one utterance)
#   2. ASR transcribes each chunk → predicted text per chunk
#   3. Dynamic programming aligns the chunk-sequence to the GT-sequence,
#      allowing both sides to skip (silence/noise chunks AND missing GT labels)
# Output: a list of (gt_idx, chunk_idx, similarity) for the best alignment.
def align_labels_to_chunks(chunk_preds, gt_labels,
                           match_threshold=0.30,
                           skip_chunk_cost=-0.05,
                           skip_gt_cost=-0.20):
    """Dynamic-programming alignment of GT labels to ASR-transcribed VAD chunks.
    Both sides may have extras: a chunk can be silence/noise (skip it), a GT
    label can be unspoken or merged into a neighbour (skip it). Similarity
    below `match_threshold` is treated as a no-match.

    Returns list of (gt_idx, chunk_idx, similarity), in order.
    """
    n = len(chunk_preds)
    m = len(gt_labels)
    if n == 0 or m == 0:
        return []
    # Precompute the similarity matrix once
    sim = [[_fuzzy_score(chunk_preds[i], gt_labels[j]) for j in range(m)]
           for i in range(n)]
    NEG_INF = float("-inf")
    # dp[i][j] = best score aligning first i chunks with first j gt labels
    dp = [[0.0] * (m + 1) for _ in range(n + 1)]
    # trace[i][j]: 'M' (match i-1↔j-1), 'C' (skip chunk i-1), 'G' (skip gt j-1)
    trace = [[None] * (m + 1) for _ in range(n + 1)]
    for i in range(1, n + 1):
        dp[i][0] = dp[i - 1][0] + skip_chunk_cost
        trace[i][0] = "C"
    for j in range(1, m + 1):
        dp[0][j] = dp[0][j - 1] + skip_gt_cost
        trace[0][j] = "G"
    for i in range(1, n + 1):
        for j in range(1, m + 1):
            s = sim[i - 1][j - 1]
            match_score = s if s >= match_threshold else NEG_INF
            best_val = NEG_INF
            best_tag = "G"
            if match_score != NEG_INF:
                v = dp[i - 1][j - 1] + match_score
                if v > best_val:
                    best_val, best_tag = v, "M"
            v = dp[i - 1][j] + skip_chunk_cost
            if v > best_val:
                best_val, best_tag = v, "C"
            v = dp[i][j - 1] + skip_gt_cost
            if v > best_val:
                best_val, best_tag = v, "G"
            dp[i][j] = best_val
            trace[i][j] = best_tag
    # Backtrack to recover the path
    pairs = []
    i, j = n, m
    while i > 0 or j > 0:
        t = trace[i][j]
        if t == "M":
            pairs.append((j - 1, i - 1, sim[i - 1][j - 1]))
            i -= 1
            j -= 1
        elif t == "C":
            i -= 1
        else:
            j -= 1
    pairs.reverse()
    return pairs


class AlignThread(QThread):
    """Forced-alignment thread: VAD → transcribe → align expected GT labels
    to audio chunks. Emits one segment per matched (gt, chunk) pair, in
    timeline order. Unmatched GT labels are reported but not emitted as
    segments (so the user can spot them and add them manually).
    """
    progress = pyqtSignal(int, int, str)        # done, total, stage description
    segment = pyqtSignal(dict)                  # matched (gt, start, end, prediction)
    done = pyqtSignal(int, int, int, float)     # matched, total_gt, unmatched_chunks, elapsed
    unmatched_gt = pyqtSignal(list)             # list of GT labels that didn't align
    error = pyqtSignal(str)

    def __init__(self, y, sr, y16, lang_iso, lang_name, gt_labels, vad_params):
        super().__init__()
        self.y = y
        self.sr = sr
        self.y16 = y16
        self.lang_iso = lang_iso
        self.lang_name = lang_name
        self.gt_labels = list(gt_labels)
        self.vad_params = vad_params
        self._stop = False

    def request_stop(self):
        self._stop = True

    def run(self):
        try:
            t0 = time.time()
            if ENGINE.backend is None or self.y16 is None:
                self.error.emit("Model not loaded, or 16 kHz audio not available.")
                return
            # ----- Stage 1: VAD -----
            self.progress.emit(0, 1, "Running VAD…")
            chunks = vad_segment(self.y, self.sr, **self.vad_params)
            if not chunks:
                self.done.emit(0, len(self.gt_labels), 0, time.time() - t0)
                return
            # ----- Stage 2: Transcribe each chunk -----
            n = len(chunks)
            chunk_preds = []
            for i, ch in enumerate(chunks):
                if self._stop:
                    self.done.emit(0, len(self.gt_labels), 0, time.time() - t0)
                    return
                s16 = max(0, int(ch["start"] * SR16))
                e16 = min(len(self.y16), int(ch["end"] * SR16))
                pred = ""
                if e16 > s16:
                    try:
                        pred = ENGINE.transcribe(
                            self.y16[s16:e16], self.lang_iso, self.lang_name,
                        )
                    except Exception as e:
                        log(f"  ASR failed on chunk {i + 1}: {e}")
                chunk_preds.append(pred)
                self.progress.emit(i + 1, n, f"Transcribing  ({i + 1}/{n})")
            # ----- Stage 3: DP alignment -----
            self.progress.emit(n, n, "Aligning to GT…")
            pairs = align_labels_to_chunks(chunk_preds, self.gt_labels)
            matched_gt = set()
            matched_chunks = set()
            for gt_idx, ch_idx, sim in pairs:
                matched_gt.add(gt_idx)
                matched_chunks.add(ch_idx)
                self.segment.emit({
                    "label": self.gt_labels[gt_idx],
                    "start": float(chunks[ch_idx]["start"]),
                    "end": float(chunks[ch_idx]["end"]),
                    "prediction": chunk_preds[ch_idx],
                    "_align_similarity": float(sim),
                })
            unmatched = [self.gt_labels[j] for j in range(len(self.gt_labels))
                         if j not in matched_gt]
            if unmatched:
                self.unmatched_gt.emit(unmatched)
            unmatched_chunks = n - len(matched_chunks)
            self.done.emit(
                len(pairs), len(self.gt_labels),
                unmatched_chunks, time.time() - t0,
            )
        except Exception as e:
            traceback.print_exc()
            self.error.emit(str(e))


class PredictTimingsThread(QThread):
    """For each GT segment, run ASR on its audio range and compute:
       - prediction text the model heard
       - the model's own start/end timestamps for that prediction
       - the *delta* between predicted times and GT times (signed errors)

    The prediction's start = first word's start, end = last word's end.
    For faster-whisper these come from `word_timestamps=True`. For
    IndicConformer (no word timestamps), they're estimated by evenly
    splitting the segment duration across the predicted words.

    Emits one updated dict per segment as it's processed. The dict has:
       label, start, end, prediction,
       pred_start, pred_end,
       start_err = pred_start - gt_start    (+ve = model later than GT)
       end_err   = pred_end   - gt_end      (+ve = model later than GT)
       pred_words = [(word, w_start, w_end), ...]    (word-level)
       pred_estimated = True | False         (False = real Whisper times)
    """
    progress = pyqtSignal(int, int)              # done, total
    segment_updated = pyqtSignal(int, dict)      # index, updated_segment_dict
    done = pyqtSignal(int, int, float)           # total, skipped, elapsed_s
    error = pyqtSignal(str)

    def __init__(self, segments, y16, lang_iso, lang_name, ref_paragraph=""):
        super().__init__()
        # Store an immutable snapshot of the segments to predict on
        self.segments = [dict(s) for s in segments]
        self.y16 = y16
        self.lang_iso = lang_iso
        self.lang_name = lang_name
        self.ref_paragraph = ref_paragraph or ""
        self._stop = False

    def request_stop(self):
        self._stop = True

    def run(self):
        try:
            t0 = time.time()
            if ENGINE.backend is None or self.y16 is None:
                self.error.emit("Model not loaded, or 16 kHz audio not available.")
                return
            total = len(self.segments)
            skipped = 0
            for i, seg in enumerate(self.segments):
                if self._stop:
                    break
                gt_start = float(seg.get("start", 0.0))
                gt_end = float(seg.get("end", 0.0))
                if gt_end <= gt_start:
                    skipped += 1
                    self.progress.emit(i + 1, total)
                    continue
                s16 = max(0, int(gt_start * SR16))
                e16 = min(len(self.y16), int(gt_end * SR16))
                if e16 <= s16:
                    skipped += 1
                    self.progress.emit(i + 1, total)
                    continue
                try:
                    text, words = ENGINE.transcribe_with_words(
                        self.y16[s16:e16], self.lang_iso, self.lang_name,
                        clip_offset=gt_start,
                    )
                except Exception as e:
                    log(f"  predict word-timings failed on seg {i+1}: {e}")
                    text, words = "", []
                # Decide pred_start / pred_end
                if words:
                    pred_start = float(words[0][1])
                    pred_end = float(words[-1][2])
                    # Native word timestamps only with faster-whisper
                    estimated = (ENGINE.engine_key != "fw")
                else:
                    pred_start = None
                    pred_end = None
                    estimated = True
                upd = dict(seg)
                upd["prediction"] = text
                upd["pred_start"] = pred_start
                upd["pred_end"] = pred_end
                upd["pred_words"] = words
                upd["pred_estimated"] = estimated
                if pred_start is not None:
                    upd["start_err"] = pred_start - gt_start
                    upd["end_err"] = pred_end - gt_end
                else:
                    upd["start_err"] = None
                    upd["end_err"] = None
                # Partial-paragraph progress: if this segment is a paragraph
                # (or p1/p2/...) and we have a reference paragraph, measure how
                # far the speaker actually read.
                seg_type = classify_segment_type(seg.get("label", ""))
                is_para = (seg_type == "paragraph"
                           or _RE_SEQ_PARA.match(str(seg.get("label", "")).strip()))
                if is_para and self.ref_paragraph and text:
                    prog = measure_paragraph_progress(text, self.ref_paragraph)
                    upd["_para_fraction"] = prog["fraction"]
                    upd["_para_spoken_words"] = prog["spoken_words"]
                    upd["_para_total_words"] = prog["total_words"]
                    upd["_para_last_word"] = prog["last_word"]
                self.segment_updated.emit(i, upd)
                self.progress.emit(i + 1, total)
            self.done.emit(total, skipped, time.time() - t0)
        except Exception as e:
            traceback.print_exc()
            self.error.emit(str(e))


# ---------- Main window ----------
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Manual Speech Annotator")
        self.resize(1750, 980)
        self.setMinimumSize(1280, 720)

        # State
        self.y = None
        self.y16 = None              # 16 kHz mono float32 — fed to ASR
        self.sr = None
        self.audio_path = None
        self.wave_t = None
        self.wave_y = None
        self.region_start = None
        self.region_end = None
        self._region_patch = None

        # View mode: 'waveform' / 'spec_linear' / 'spec_logdb' / 'spec_mel'
        self.view_mode = "waveform"
        self._spec_cache = {}   # {mode: (freqs, times, S, sr, n_samples)}

        # Model-load tick state
        self._model_load_t0 = None
        self._model_tick_timer = None

        # Auto-save (every 30s while there are segments to save)
        self._autosave_timer = QTimer(self)
        self._autosave_timer.setInterval(30_000)
        self._autosave_timer.timeout.connect(self._autosave_tick)
        self._autosave_timer.start()
        self._last_autosaved_count = 0  # avoid writing same content over and over

        # Playback speed (Audacity has this — slows down hard words)
        self.playback_speed = 1.0

        # Multi-file project state
        self.project_dir = None
        self.project_files = []   # list of Path

        # Drag state for the four interaction modes:
        #   'new'           — drag in empty space creates a new region
        #   'resize_start'  — grabbed the left edge of an existing region
        #   'resize_end'    — grabbed the right edge
        #   'move'          — pressed inside an existing region, drag it whole
        self._drag_mode = None
        self._drag_anchor = None        # data-x where the press started
        self._drag_moved = False
        self._press_xdata = None
        self._press_button = None
        self._initial_region = None     # snapshot of (start, end) at press time
        self._move_offset = None        # data-x captured for 'move' mode

        # Playback cursor (moving vertical line during sd.play)
        self._playback_x = None
        self._playback_start_time = None
        self._playback_start_x = None
        self._playback_end_x = None
        self._playback_timer = None

        # Scrollbar feedback-loop guard
        self._suppress_scrollbar = False

        self.segments = []  # [{'label': str, 'start': float, 'end': float}]
        self._selected_seg_idx = None
        self._suppress_spin = False
        self._threads = {}

        # Undo / redo: ring buffers of deep-copied segment snapshots.
        # We snapshot *before* each mutating action via self._push_undo().
        self._undo_stack = []   # list[(segments_snapshot, selected_idx)]
        self._redo_stack = []
        self._undo_limit = 60

        # Per-segment timing log for the paper's "annotation time" study.
        # Maps segment-identity -> seconds the annotator spent with it
        # selected. Lets us report time-per-segment without a stopwatch.
        self._timing_log = {}          # {seg_uid: cumulative_seconds}
        self._timing_active_uid = None
        self._timing_active_since = None
        self._session_start = time.time()

        self.current_words = []  # current language's wordlist (chars + words combined for picker)
        self.current_chars = []
        self.current_sentences = {}
        self.current_paragraph = ""

        ensure_dirs()
        self._build_ui()
        self._on_lang_changed()  # populate wordlist for default lang

    # ----- UI construction -----
    def _build_ui(self):
        root = QWidget()
        self.setCentralWidget(root)
        outer = QVBoxLayout(root)
        outer.setContentsMargins(8, 8, 8, 8)
        outer.setSpacing(6)

        # ===== Row 1: language + file actions =====
        r1 = QHBoxLayout()
        r1.addWidget(QLabel("Language:"))
        self.cb_lang = QComboBox()
        for disp, iso, name in LANGUAGES:
            self.cb_lang.addItem(disp, (iso, name))
        self.cb_lang.currentIndexChanged.connect(self._on_lang_changed)
        r1.addWidget(self.cb_lang)

        self.btn_load_audio = QPushButton("🎵 Load Audio")
        self.btn_load_audio.setToolTip("Load .wav / .mp3 / .flac / .ogg / .m4a")
        self.btn_load_audio.clicked.connect(self.on_load_audio)
        r1.addWidget(self.btn_load_audio)

        self.btn_import_docx = QPushButton("📄 Import Wordlist")
        self.btn_import_docx.setToolTip(
            "Import the per-language wordlist from a .docx or .pdf text material.\n"
            "Reads characters, words, sentences and the paragraph in one go.\n"
            "Result is cached to wordlists/<iso>.txt so this is one-time per language."
        )
        self.btn_import_docx.clicked.connect(self.on_import_docx)
        r1.addWidget(self.btn_import_docx)

        # PROMINENT GT loader — this is the most-asked-about button
        self.btn_load_csv = QPushButton("📥 Load GT")
        self.btn_load_csv.setStyleSheet(
            "QPushButton{font-weight:bold;color:white;background:#0891b2;"
            "border:1px solid #0e7490;padding:5px 14px;border-radius:4px}"
            "QPushButton:hover{background:#0e7490}"
        )
        self.btn_load_csv.setToolTip(
            "📥 Load ground-truth annotations from a file.\n\n"
            "Accepts: .xlsx  .csv  .tsv\n"
            "Expected columns:  label,  start_time,  end_time,  [prediction]\n"
            "Times can be HH:MM:SS.fff, MM:SS, or plain seconds.\n"
            "The first row is auto-detected as a header if it contains "
            "'label / start / end'."
        )
        self.btn_load_csv.clicked.connect(self.on_load_csv)
        r1.addWidget(self.btn_load_csv)

        self.btn_save_csv = QPushButton("💾 Save")
        self.btn_save_csv.setToolTip(
            "Save annotations.\n"
            "Pick format in the dialog:\n"
            "  • .xlsx — Excel with bold header & frozen pane\n"
            "  • .csv  — headerless, HH:MM:SS times (loadable by your app.py)\n"
            "  • .tsv  — tab-separated"
        )
        self.btn_save_csv.setStyleSheet(
            "QPushButton{font-weight:bold;color:white;background:#16a34a;"
            "border:1px solid #15803d;padding:5px 14px;border-radius:4px}"
            "QPushButton:hover{background:#15803d}"
        )
        self.btn_save_csv.clicked.connect(self.on_save_csv)
        r1.addWidget(self.btn_save_csv)

        self.btn_export_transcripts = QPushButton("📤 Export Transcripts")
        self.btn_export_transcripts.setToolTip(
            "Write transcripts/<iso>.txt with the [s1]..[s5] / [paragraph] reference text. "
            "Used by app.py for accuracy scoring on sentence/paragraph labels."
        )
        self.btn_export_transcripts.clicked.connect(self.on_export_transcripts)
        r1.addWidget(self.btn_export_transcripts)

        r1.addStretch()
        outer.addLayout(r1)

        # ===== Row 2: ASR / auto-detect controls =====
        r2 = QHBoxLayout()
        r2.addWidget(QLabel("ASR:"))

        self.btn_load_model = QPushButton("⬇ Load")
        self.btn_load_model.setToolTip(
            "Load the speech-recognition model for the current language.\n"
            "  • English   → faster-whisper large-v3\n"
            "  • Hindi / Marathi / Gujarati → IndicConformer-600M (IIT Madras)"
        )
        self.btn_load_model.setStyleSheet(
            "QPushButton{font-weight:bold;color:white;background:#2563eb;"
            "border:1px solid #1d4ed8;padding:5px 12px;border-radius:4px}"
            "QPushButton:hover{background:#1d4ed8}"
            "QPushButton:disabled{color:#9ca3af;background:#f3f4f6}"
        )
        self.btn_load_model.clicked.connect(self.on_load_model)
        r2.addWidget(self.btn_load_model)

        # Model status pill — fixed width so the longer button names below
        # don't get squashed when the model name is long. Truncation handled
        # by Qt's elidedText when name exceeds width.
        self.lbl_model_status = QLabel("")
        self.lbl_model_status.setStyleSheet(
            "color:#374151;font-size:11px;padding:2px 8px;"
            "background:#f3f4f6;border-radius:3px"
        )
        self.lbl_model_status.setFixedWidth(220)
        self.lbl_model_status.setTextInteractionFlags(Qt.TextSelectableByMouse)
        r2.addWidget(self.lbl_model_status)

        self.btn_auto_detect = QPushButton("🤖 Auto-Detect")
        self.btn_auto_detect.setStyleSheet(
            "QPushButton{font-weight:bold;color:white;background:#7c3aed;"
            "border:1px solid #6d28d9;padding:5px 12px;border-radius:4px}"
            "QPushButton:hover{background:#6d28d9}"
            "QPushButton:disabled{color:#9ca3af;background:#f3f4f6}"
        )
        self.btn_auto_detect.setToolTip(
            "🤖 Auto-Detect Segments  (VAD + ASR)\n\n"
            "Runs VAD to find speech regions, then transcribes each one. "
            "The predicted text becomes the initial label — you can edit each "
            "before saving."
        )
        self.btn_auto_detect.clicked.connect(self.on_auto_detect)
        r2.addWidget(self.btn_auto_detect)

        # Forced alignment: GT text → audio timestamps via Whisper/IndicConformer
        self.btn_align_gt = QPushButton("📐 Align GT")
        self.btn_align_gt.setStyleSheet(
            "QPushButton{font-weight:bold;color:white;background:#0891b2;"
            "border:1px solid #0e7490;padding:5px 12px;border-radius:4px}"
            "QPushButton:hover{background:#0e7490}"
            "QPushButton:disabled{color:#9ca3af;background:#f3f4f6}"
        )
        self.btn_align_gt.setToolTip(
            "📐 Align GT to Audio\n\n"
            "Given a list of expected labels (chars/words/sentences/paragraph), "
            "find each one's timestamps in the audio.\n\n"
            "How it works:\n"
            "  1. VAD splits audio into speech chunks\n"
            "  2. ASR transcribes each chunk\n"
            "  3. DP alignment maps your GT labels to chunks\n"
            "  4. Segments populated with (label, start, end, what model heard)\n\n"
            "Source can be:\n"
            "  • The loaded wordlist  (chars + words + sentences + paragraph)\n"
            "  • A pasted custom list  (one label per line)\n"
            "  • A loaded GT file  (uses its labels, finds new timestamps)"
        )
        self.btn_align_gt.clicked.connect(self.on_align_gt)
        r2.addWidget(self.btn_align_gt)

        # NEW: Run the model on each GT segment to get (predicted label,
        # predicted start, predicted end) — exactly what the user asked for.
        # This is "score my GT" — different from auto-detect (which has no GT)
        # and Align (which finds timestamps when GT has no times).
        self.btn_predict_gt = QPushButton("🎯 Predict on GT")
        self.btn_predict_gt.setStyleSheet(
            "QPushButton{font-weight:bold;color:white;background:#7c3aed;"
            "border:1px solid #6d28d9;padding:5px 14px;border-radius:4px}"
            "QPushButton:hover{background:#6d28d9}"
            "QPushButton:disabled{color:#9ca3af;background:#f3f4f6}"
        )
        self.btn_predict_gt.setToolTip(
            "Run ASR on each loaded GT segment. For every segment you get:\n"
            "  • the predicted label (what the model hears in that range)\n"
            "  • the model's own start/end timestamps  (faster-whisper only)\n"
            "  • start/end deltas vs your GT times\n\n"
            "Use this after loading a GT file (e.g. Female.xlsx) to see "
            "where the model disagrees with your annotations."
        )
        self.btn_predict_gt.clicked.connect(self.on_predict_on_gt)
        r2.addWidget(self.btn_predict_gt)

        self.btn_vad_only = QPushButton("VAD")
        self.btn_vad_only.setToolTip(
            "VAD only — split into speech regions without running ASR. "
            "Useful when no model is loaded or you want to label everything by hand."
        )
        self.btn_vad_only.clicked.connect(self.on_vad_only)
        r2.addWidget(self.btn_vad_only)

        self.btn_stop_auto = QPushButton("⏹ Stop")
        self.btn_stop_auto.setEnabled(False)
        self.btn_stop_auto.clicked.connect(self.on_stop_auto_detect)
        r2.addWidget(self.btn_stop_auto)

        # VAD param spinboxes (compact)
        r2.addWidget(QLabel("  silence:"))
        self.sp_vad_sil = QSpinBox()
        self.sp_vad_sil.setRange(50, 2000); self.sp_vad_sil.setValue(300); self.sp_vad_sil.setSuffix(" ms")
        r2.addWidget(self.sp_vad_sil)
        r2.addWidget(QLabel("threshold:"))
        self.sp_vad_db = QDoubleSpinBox()
        self.sp_vad_db.setRange(-80, 0); self.sp_vad_db.setValue(-40); self.sp_vad_db.setSuffix(" dB")
        r2.addWidget(self.sp_vad_db)
        r2.addWidget(QLabel("min:"))
        self.sp_vad_min = QSpinBox()
        self.sp_vad_min.setRange(50, 5000); self.sp_vad_min.setValue(200); self.sp_vad_min.setSuffix(" ms")
        r2.addWidget(self.sp_vad_min)
        r2.addWidget(QLabel("max:"))
        self.sp_vad_max = QSpinBox()
        self.sp_vad_max.setRange(1000, 30000); self.sp_vad_max.setValue(10000); self.sp_vad_max.setSuffix(" ms")
        r2.addWidget(self.sp_vad_max)

        r2.addStretch()
        outer.addLayout(r2)

        # ===== Status bar (fixed height — long messages don't reflow the layout) =====
        self.lbl_status = QLabel(
            "Pick a language → 🎵 Load Audio → mark region → choose label → Add Segment"
        )
        self.lbl_status.setStyleSheet(
            "padding:4px 8px;color:#374151;font-size:12px;"
            "background:#f3f4f6;border-radius:3px"
        )
        self.lbl_status.setFixedHeight(26)
        self.lbl_status.setTextInteractionFlags(Qt.TextSelectableByMouse)
        # Truncate text that doesn't fit instead of growing the label
        outer.addWidget(self.lbl_status)

        # ===== Main split: waveform (left) / labels + segments (right) =====
        splitter = QSplitter(Qt.Horizontal)

        # ----- Left: waveform + region controls -----
        left = QWidget()
        ll = QVBoxLayout(left)
        ll.setContentsMargins(0, 0, 0, 0)
        ll.setSpacing(4)

        self.fig, self.ax = plt.subplots(figsize=(12, 5))
        self.fig.set_tight_layout(True)
        # Audacity-style: start dark even before any audio is loaded
        self.fig.patch.set_facecolor(self._AUDACITY["bg"])
        self.ax.set_facecolor(self._AUDACITY["bg"])
        for spine in self.ax.spines.values():
            spine.set_color(self._AUDACITY["spine"])
        self.ax.tick_params(colors=self._AUDACITY["axis_text"], which="both")
        self.canvas = FigureCanvas(self.fig)
        self.toolbar = NavigationToolbar(self.canvas, self)
        # Connect interactive events
        self.canvas.mpl_connect("button_press_event", self._on_canvas_press)
        self.canvas.mpl_connect("motion_notify_event", self._on_canvas_motion)
        self.canvas.mpl_connect("button_release_event", self._on_canvas_release)
        self.canvas.mpl_connect("scroll_event", self._on_canvas_scroll)
        ll.addWidget(self.toolbar)
        ll.addWidget(self.canvas)

        # Horizontal scrollbar (pans the waveform when zoomed in)
        self.scrollbar = QScrollBar(Qt.Horizontal)
        self.scrollbar.setMinimum(0)
        self.scrollbar.setMaximum(0)  # configured when audio loads
        self.scrollbar.setValue(0)
        self.scrollbar.setSingleStep(10)   # 10 ms per arrow click
        self.scrollbar.setPageStep(1000)   # 1 s per page click (placeholder)
        self.scrollbar.setToolTip("Drag to scroll the waveform left/right")
        self.scrollbar.valueChanged.connect(self._on_scrollbar_changed)
        self.scrollbar.setStyleSheet(
            "QScrollBar:horizontal{background:#1f2a36;height:14px;border:1px solid #3a4d5e}"
            "QScrollBar::handle:horizontal{background:#5294c4;min-width:24px;border-radius:3px;margin:1px}"
            "QScrollBar::handle:horizontal:hover{background:#7fb1d8}"
            "QScrollBar::add-line:horizontal,QScrollBar::sub-line:horizontal{width:0;background:transparent}"
            "QScrollBar::add-page:horizontal,QScrollBar::sub-page:horizontal{background:#283645}"
        )
        ll.addWidget(self.scrollbar)

        # Zoom / view controls row (between waveform and region controls)
        view_row = QHBoxLayout()
        view_row.setSpacing(4)
        view_row.addWidget(QLabel("View:"))

        # Visualization mode dropdown — waveform vs three spectrogram flavours
        self.cb_view_mode = QComboBox()
        self.cb_view_mode.addItem("🌊 Waveform (linear)",         "waveform")
        self.cb_view_mode.addItem("🌊 Waveform (dB, −60 floor)",  "wave_db")
        self.cb_view_mode.addItem("📊 Spectrogram (linear)",      "spec_linear")
        self.cb_view_mode.addItem("📊 Spectrogram (log dB)",      "spec_logdb")
        self.cb_view_mode.addItem("🎵 Mel Spectrogram (dB)",      "spec_mel")
        self.cb_view_mode.setToolTip(
            "Waveform (linear): classic Audacity look, amplitude vs time\n"
            "Waveform (dB):     same wave shape but Y-axis is dB — quiet "
            "parts become visible (Audacity's 'Waveform (dB)' view)\n"
            "Spectrogram (linear / log dB / mel): time-frequency views"
        )
        self.cb_view_mode.currentIndexChanged.connect(self._on_view_mode_changed)
        view_row.addWidget(self.cb_view_mode)

        # Toggle: show PRED badges on every visible segment (not just selected)
        self.chk_show_preds = QCheckBox("Show predictions on all segments")
        self.chk_show_preds.setToolTip(
            "Off: only the currently-selected segment shows GT (your label) "
            "and PRED (model's original prediction).\n"
            "On: every visible segment shows both — useful for scanning a long "
            "auto-detect pass to spot where the model and your labels disagree."
        )
        self.chk_show_preds.setStyleSheet("QCheckBox{color:#374151;font-size:11px}")
        self.chk_show_preds.toggled.connect(
            lambda _: (self.y is not None) and self._redraw(xlim=self.ax.get_xlim())
        )
        view_row.addWidget(self.chk_show_preds)

        view_row.addSpacing(12)
        view_row.addWidget(QLabel("Zoom:"))

        self.btn_zoom_in = QPushButton("🔍+ Zoom In")
        self.btn_zoom_in.setToolTip("Zoom in 2× around center  (or scroll up on waveform)")
        self.btn_zoom_in.clicked.connect(self.on_zoom_in)
        view_row.addWidget(self.btn_zoom_in)

        self.btn_zoom_out = QPushButton("🔍− Zoom Out")
        self.btn_zoom_out.setToolTip("Zoom out 2× around center  (or scroll down on waveform)")
        self.btn_zoom_out.clicked.connect(self.on_zoom_out)
        view_row.addWidget(self.btn_zoom_out)

        self.btn_fit = QPushButton("Fit")
        self.btn_fit.setToolTip(
            "Fit view to (in order of priority): the selected region,\n"
            "the selected segment, all added segments, or the full audio."
        )
        self.btn_fit.clicked.connect(self.on_fit_view)
        view_row.addWidget(self.btn_fit)

        self.btn_reset_view = QPushButton("Reset View")
        self.btn_reset_view.setToolTip("Show the entire audio file")
        self.btn_reset_view.clicked.connect(self.on_reset_view)
        view_row.addWidget(self.btn_reset_view)

        self.lbl_view = QLabel("")
        self.lbl_view.setStyleSheet("color:#6b7280;font-size:11px;padding-left:8px")
        view_row.addWidget(self.lbl_view)

        view_row.addStretch()
        ll.addLayout(view_row)

        # Region row
        region_box = QGroupBox("Selected Region")
        rg = QGridLayout(region_box)
        rg.addWidget(QLabel("Start (s):"), 0, 0)
        self.sp_start = QDoubleSpinBox()
        self.sp_start.setRange(0.0, 99999.0)
        self.sp_start.setDecimals(4)
        self.sp_start.setSingleStep(0.05)
        self.sp_start.valueChanged.connect(self._on_spin_changed)
        rg.addWidget(self.sp_start, 0, 1)

        rg.addWidget(QLabel("End (s):"), 0, 2)
        self.sp_end = QDoubleSpinBox()
        self.sp_end.setRange(0.0, 99999.0)
        self.sp_end.setDecimals(4)
        self.sp_end.setSingleStep(0.05)
        self.sp_end.valueChanged.connect(self._on_spin_changed)
        rg.addWidget(self.sp_end, 0, 3)

        self.lbl_dur = QLabel("00:00:00.000 → 00:00:00.000  (0.000 s)")
        self.lbl_dur.setStyleSheet("color:#6b7280;font-size:11px;font-family:Consolas,Courier New,monospace")
        self.lbl_dur.setMinimumWidth(260)
        rg.addWidget(self.lbl_dur, 0, 4)

        self.btn_play_region = QPushButton("▶ Play Region")
        self.btn_play_region.setToolTip("Play just the selected region  (Space)")
        self.btn_play_region.clicked.connect(self.on_play_region)
        rg.addWidget(self.btn_play_region, 0, 5)

        self.btn_play_all = QPushButton("▶▶ Play All")
        self.btn_play_all.clicked.connect(self.on_play_all)
        rg.addWidget(self.btn_play_all, 0, 6)

        self.btn_stop = QPushButton("■ Stop")
        self.btn_stop.setToolTip("Stop playback  (Esc)")
        self.btn_stop.clicked.connect(self.on_stop_playback)
        rg.addWidget(self.btn_stop, 0, 7)

        # Variable-speed playback — slower speeds make hard words easier to label.
        # We change sample rate at playback time; this *will* shift the pitch but
        # for annotation review that's typically fine.
        rg.addWidget(QLabel("Speed:"), 0, 8)
        self.cb_speed = QComboBox()
        for label, val in [("0.5×", 0.5), ("0.75×", 0.75), ("1.0×", 1.0),
                           ("1.25×", 1.25), ("1.5×", 1.5), ("2.0×", 2.0)]:
            self.cb_speed.addItem(label, val)
        self.cb_speed.setCurrentIndex(2)   # default 1.0×
        self.cb_speed.setToolTip(
            "Playback speed. Below 1.0× = slow-motion (great for tricky words).\n"
            "Pitch shifts with speed — annotators learn to ignore this within a "
            "couple of minutes."
        )
        self.cb_speed.currentIndexChanged.connect(
            lambda _i: setattr(self, "playback_speed", self.cb_speed.currentData())
        )
        rg.addWidget(self.cb_speed, 0, 9)

        self.btn_clear_region = QPushButton("Clear Region")
        self.btn_clear_region.clicked.connect(self.on_clear_region)
        rg.addWidget(self.btn_clear_region, 0, 10)

        ll.addWidget(region_box)

        # Help line
        hint = QLabel(
            "<b>Drag</b> a new selection · <b>Drag the edges</b> to resize · "
            "<b>Drag inside</b> to move the whole selection · "
            "<b>Shift+click</b> = extend nearest edge · "
            "<b>Right-click</b> = set End · "
            "<b>Scroll</b> = zoom around cursor<br>"
            "Shortcuts: <b>Space</b> play/stop region · <b>Esc</b> stop · "
            "<b>Ctrl+= / Ctrl+−</b> zoom · <b>Ctrl+F</b> fit · <b>Ctrl+0</b> reset · "
            "<b>Ctrl+Enter</b> add segment · <b>Ctrl+S</b> save CSV"
        )
        hint.setStyleSheet("color:#6b7280;font-size:11px;padding:2px 4px")
        hint.setTextFormat(Qt.RichText)
        ll.addWidget(hint)

        splitter.addWidget(left)

        # ----- Right: tabs (Labels / Segments) -----
        right_tabs = QTabWidget()

        # --- Labels tab ---
        labels_tab = QWidget()
        lt = QVBoxLayout(labels_tab)

        # Big label-display + Add Segment row
        self.txt_current_label = QLineEdit()
        self.txt_current_label.setPlaceholderText("Selected / typed label appears here")
        f = self.txt_current_label.font()
        f.setPointSize(13)
        f.setBold(True)
        self.txt_current_label.setFont(f)
        self.txt_current_label.setStyleSheet("padding:6px;background:#fefce8;border:2px solid #facc15;border-radius:4px")
        lt.addWidget(QLabel("Current label:"))
        lt.addWidget(self.txt_current_label)

        self.btn_add_seg = QPushButton("➕ Add Segment")
        self.btn_add_seg.setStyleSheet(
            "QPushButton{font-weight:bold;color:white;background:#2563eb;"
            "border:1px solid #1d4ed8;padding:8px 14px;border-radius:4px;font-size:13px}"
            "QPushButton:hover{background:#1d4ed8}"
        )
        self.btn_add_seg.clicked.connect(self.on_add_segment)
        lt.addWidget(self.btn_add_seg)

        # Quick labels (s1..s5, paragraph)
        ql_box = QGroupBox("Quick Labels (sentences / paragraph)")
        qg = QGridLayout(ql_box)
        for i, ql in enumerate(QUICK_LABELS):
            b = QPushButton(ql)
            b.setStyleSheet("QPushButton{background:#dbeafe;font-weight:bold;padding:6px}")
            b.clicked.connect(lambda _, lab=ql: self._set_current_label(lab))
            qg.addWidget(b, i // 3, i % 3)
        lt.addWidget(ql_box)

        # Chars panel (scrollable grid of buttons)
        chars_box = QGroupBox("Characters / Alphabets")
        cv = QVBoxLayout(chars_box)
        self.chars_area = QScrollArea()
        self.chars_area.setWidgetResizable(True)
        self.chars_area.setMaximumHeight(120)
        self.chars_container = QWidget()
        self.chars_grid = QGridLayout(self.chars_container)
        self.chars_grid.setSpacing(2)
        self.chars_area.setWidget(self.chars_container)
        cv.addWidget(self.chars_area)
        lt.addWidget(chars_box)

        # Words list with filter
        words_box = QGroupBox("Words (click to select)")
        wv = QVBoxLayout(words_box)
        self.txt_filter = QLineEdit()
        self.txt_filter.setPlaceholderText("Filter words…")
        self.txt_filter.textChanged.connect(self._refresh_words_list)
        wv.addWidget(self.txt_filter)
        self.words_list = QListWidget()
        self.words_list.itemClicked.connect(
            lambda item: self._set_current_label(item.text())
        )
        # Use a font that handles Indic glyphs nicely
        self.words_list.setStyleSheet("QListWidget{font-size:13px}")
        wv.addWidget(self.words_list)
        lt.addWidget(words_box, 1)

        # Add-new-word row
        add_box = QGroupBox("Add new word to the list")
        ah = QHBoxLayout(add_box)
        self.txt_new_word = QLineEdit()
        self.txt_new_word.setPlaceholderText("Type a new word here…")
        self.txt_new_word.returnPressed.connect(self.on_add_new_word)
        ah.addWidget(self.txt_new_word, 1)
        self.btn_add_word = QPushButton("Add to list")
        self.btn_add_word.clicked.connect(self.on_add_new_word)
        ah.addWidget(self.btn_add_word)
        self.btn_add_word_and_use = QPushButton("Add & Use")
        self.btn_add_word_and_use.setStyleSheet("QPushButton{background:#dcfce7}")
        self.btn_add_word_and_use.clicked.connect(self.on_add_new_word_and_use)
        ah.addWidget(self.btn_add_word_and_use)
        lt.addWidget(add_box)

        right_tabs.addTab(labels_tab, "Labels")

        # --- Segments tab ---
        segs_tab = QWidget()
        sv = QVBoxLayout(segs_tab)
        sv.addWidget(QLabel("Added segments (click to highlight on waveform):"))
        self.seg_list = QListWidget()
        self.seg_list.setStyleSheet(
            "QListWidget{font-family:Consolas,Courier New,monospace;font-size:11px}"
            "QListWidget::item{padding:3px 6px}"
            "QListWidget::item:selected{background:#dbeafe}"
        )
        self.seg_list.itemClicked.connect(self._on_seg_clicked)
        # Also fire on keyboard arrow-key navigation
        self.seg_list.currentItemChanged.connect(
            lambda cur, _prev: cur and self._on_seg_clicked(cur)
        )
        sv.addWidget(self.seg_list, 1)

        seg_btns = QHBoxLayout()
        self.btn_del_seg = QPushButton("🗑 Delete Selected")
        self.btn_del_seg.clicked.connect(self.on_delete_segment)
        seg_btns.addWidget(self.btn_del_seg)
        self.btn_relabel = QPushButton("✏ Relabel from Current")
        self.btn_relabel.clicked.connect(self.on_relabel_segment)
        seg_btns.addWidget(self.btn_relabel)
        self.btn_retime = QPushButton("⏱ Update Times from Region")
        self.btn_retime.clicked.connect(self.on_retime_segment)
        seg_btns.addWidget(self.btn_retime)
        self.btn_load_to_region = QPushButton("Load to Region")
        self.btn_load_to_region.clicked.connect(self.on_load_seg_to_region)
        seg_btns.addWidget(self.btn_load_to_region)
        sv.addLayout(seg_btns)

        # Undo / redo row
        undo_row = QHBoxLayout()
        self.btn_undo = QPushButton("↩ Undo")
        self.btn_undo.setToolTip("Undo the last add/delete/relabel/retime  (Ctrl+Z)")
        self.btn_undo.clicked.connect(self.on_undo)
        undo_row.addWidget(self.btn_undo)
        self.btn_redo = QPushButton("↪ Redo")
        self.btn_redo.setToolTip("Redo  (Ctrl+Shift+Z or Ctrl+Y)")
        self.btn_redo.clicked.connect(self.on_redo)
        undo_row.addWidget(self.btn_redo)
        undo_row.addStretch()
        sv.addLayout(undo_row)

        right_tabs.addTab(segs_tab, "Segments")

        # --- Files tab (multi-file project mode) ---
        files_tab = QWidget()
        fv = QVBoxLayout(files_tab)
        proj_btns = QHBoxLayout()
        self.btn_open_folder = QPushButton("📁 Open folder…")
        self.btn_open_folder.setToolTip(
            "Pick a folder of audio files. They'll appear in the list below.\n"
            "Click a file to load it. Status icons:\n"
            "  📂  currently loaded\n"
            "  ✓   saved annotations exist\n"
            "  ✎   autosave exists (in-progress)\n"
            "  ·   not started"
        )
        self.btn_open_folder.clicked.connect(self.on_open_project_folder)
        proj_btns.addWidget(self.btn_open_folder)
        self.btn_next_unfinished = QPushButton("⏭ Next unfinished")
        self.btn_next_unfinished.setToolTip(
            "Jump to the next file in this folder that has no saved annotations."
        )
        self.btn_next_unfinished.clicked.connect(self.on_jump_next_unfinished)
        self.btn_next_unfinished.setEnabled(False)
        proj_btns.addWidget(self.btn_next_unfinished)
        fv.addLayout(proj_btns)

        self.lbl_project = QLabel("(no folder open)")
        self.lbl_project.setStyleSheet("color:#6b7280;font-size:11px;padding:2px")
        self.lbl_project.setWordWrap(True)
        fv.addWidget(self.lbl_project)

        self.proj_list = QListWidget()
        self.proj_list.setStyleSheet(
            "QListWidget{background:#fff;border:1px solid #d0d0d0;"
            "font-family:Consolas,Courier New,monospace;font-size:12px}"
            "QListWidget::item{padding:4px 6px}"
            "QListWidget::item:selected{background:#dbeafe}"
        )
        self.proj_list.itemDoubleClicked.connect(self._on_project_file_dblclick)
        fv.addWidget(self.proj_list, 1)

        proj_legend = QLabel(
            "Double-click to load.  "
            "<b>📂</b> loaded · <b>✓</b> saved · <b>✎</b> in-progress · <b>·</b> untouched"
        )
        proj_legend.setStyleSheet("color:#6b7280;font-size:10px;padding:2px")
        proj_legend.setTextFormat(Qt.RichText)
        fv.addWidget(proj_legend)
        right_tabs.addTab(files_tab, "Files")

        # --- Diff tab (GT vs PRED for the selected segment) ---
        diff_tab = QWidget()
        dv = QVBoxLayout(diff_tab)
        dv.addWidget(QLabel(
            "<b>GT</b> (your label) vs <b>PRED</b> (model's prediction) for the "
            "currently-selected segment. Differences highlighted character by "
            "character.</br>"
        ))
        self.txt_diff = QLabel(
            "<i style='color:#888'>Select a segment to see the diff.</i>"
        )
        self.txt_diff.setTextFormat(Qt.RichText)
        self.txt_diff.setWordWrap(True)
        self.txt_diff.setTextInteractionFlags(Qt.TextSelectableByMouse)
        self.txt_diff.setStyleSheet(
            "QLabel{background:#fff;border:1px solid #d0d0d0;"
            "padding:10px;font-size:14px;font-family:Consolas,Courier New,monospace}"
        )
        self.txt_diff.setAlignment(Qt.AlignTop | Qt.AlignLeft)
        diff_scroll = QScrollArea()
        diff_scroll.setWidget(self.txt_diff)
        diff_scroll.setWidgetResizable(True)
        dv.addWidget(diff_scroll, 1)
        right_tabs.addTab(diff_tab, "Diff")

        # --- Score tab (overall GT vs PRED accuracy report) ---
        score_tab = QWidget()
        sv2 = QVBoxLayout(score_tab)
        sv2.addWidget(QLabel(
            "<b>Accuracy report</b> — character-level similarity between "
            "each segment's GT (your label) and PRED (model's transcription)."
        ))
        self.btn_compute_score = QPushButton("📊 Compute accuracy now")
        self.btn_compute_score.setToolTip(
            "Compute and refresh the accuracy score across all segments that "
            "have both a GT label and a model prediction."
        )
        self.btn_compute_score.clicked.connect(self._update_score_panel)
        sv2.addWidget(self.btn_compute_score)

        # Export row: paper-ready metrics report + training manifest
        exp_row = QHBoxLayout()
        self.btn_export_report = QPushButton("📄 Export metrics report")
        self.btn_export_report.setToolTip(
            "Save a text report of all metrics (acceptance rate, edit-WER, "
            "timing, per-language breakdown) — paste-ready for your paper."
        )
        self.btn_export_report.clicked.connect(self.on_export_metrics_report)
        exp_row.addWidget(self.btn_export_report)

        self.btn_export_manifest = QPushButton("📦 Export training manifest (JSONL)")
        self.btn_export_manifest.setToolTip(
            "Export one JSON line per segment (audio path, offsets, label, "
            "prediction) — ready for ASR fine-tuning pipelines."
        )
        self.btn_export_manifest.clicked.connect(self.on_export_manifest)
        exp_row.addWidget(self.btn_export_manifest)
        sv2.addLayout(exp_row)
        self.txt_score = QLabel(
            "<i style='color:#888'>Load GT or run Auto-Detect / Align GT, "
            "then click '📊 Compute accuracy now' to see the report.</i>"
        )
        self.txt_score.setTextFormat(Qt.RichText)
        self.txt_score.setWordWrap(True)
        self.txt_score.setTextInteractionFlags(Qt.TextSelectableByMouse)
        self.txt_score.setStyleSheet(
            "QLabel{background:#fff;border:1px solid #d0d0d0;"
            "padding:10px;font-size:12px;font-family:Consolas,Courier New,monospace}"
        )
        self.txt_score.setAlignment(Qt.AlignTop | Qt.AlignLeft)
        score_scroll = QScrollArea()
        score_scroll.setWidget(self.txt_score)
        score_scroll.setWidgetResizable(True)
        sv2.addWidget(score_scroll, 1)
        right_tabs.addTab(score_tab, "Score")

        # Right panel: fixed-width so the waveform on the left stays put
        # when right-side content (tabs, segments list, etc.) changes.
        # User can still drag the splitter to make it wider/narrower.
        right_tabs.setMinimumWidth(560)
        right_tabs.setMaximumWidth(720)

        splitter.addWidget(right_tabs)
        # Left side grows; right stays at its preferred width
        splitter.setStretchFactor(0, 1)
        splitter.setStretchFactor(1, 0)
        splitter.setSizes([1200, 580])
        # Collapse-prevent: dragging the divider can't hide either side
        splitter.setCollapsible(0, False)
        splitter.setCollapsible(1, False)
        outer.addWidget(splitter, 1)

        # Keyboard shortcuts
        self._sc_space = QShortcut(QKeySequence(Qt.Key_Space), self)
        self._sc_space.activated.connect(self.on_toggle_play_region)
        self._sc_esc = QShortcut(QKeySequence(Qt.Key_Escape), self)
        self._sc_esc.activated.connect(self.on_stop_playback)
        self._sc_zoom_in = QShortcut(QKeySequence("Ctrl++"), self)
        self._sc_zoom_in.activated.connect(self.on_zoom_in)
        self._sc_zoom_in2 = QShortcut(QKeySequence("Ctrl+="), self)
        self._sc_zoom_in2.activated.connect(self.on_zoom_in)
        self._sc_zoom_out = QShortcut(QKeySequence("Ctrl+-"), self)
        self._sc_zoom_out.activated.connect(self.on_zoom_out)
        self._sc_fit = QShortcut(QKeySequence("Ctrl+F"), self)
        self._sc_fit.activated.connect(self.on_fit_view)
        self._sc_reset_view = QShortcut(QKeySequence("Ctrl+0"), self)
        self._sc_reset_view.activated.connect(self.on_reset_view)
        self._sc_add_seg = QShortcut(QKeySequence("Ctrl+Return"), self)
        self._sc_add_seg.activated.connect(self.on_add_segment)
        self._sc_save = QShortcut(QKeySequence("Ctrl+S"), self)
        self._sc_save.activated.connect(self.on_save_csv)
        # Undo / redo
        self._sc_undo = QShortcut(QKeySequence("Ctrl+Z"), self)
        self._sc_undo.activated.connect(self.on_undo)
        self._sc_redo = QShortcut(QKeySequence("Ctrl+Shift+Z"), self)
        self._sc_redo.activated.connect(self.on_redo)
        self._sc_redo2 = QShortcut(QKeySequence("Ctrl+Y"), self)
        self._sc_redo2.activated.connect(self.on_redo)
        # Delete selected segment
        self._sc_del = QShortcut(QKeySequence("Ctrl+Backspace"), self)
        self._sc_del.activated.connect(self.on_delete_segment)

    # ----- Language / wordlist handling -----
    def _current_lang(self):
        return self.cb_lang.currentData()  # (iso, name)

    def _on_lang_changed(self):
        iso, name = self._current_lang()
        # Load persisted wordlist
        words = load_wordlist(iso)
        # Always reset sentence/paragraph caches — we re-populate below from
        # whichever source is freshest (transcripts/<iso>.txt > docx/pdf).
        self.current_chars = []
        self.current_sentences = {}
        self.current_paragraph = ""

        if not words:
            # Try to auto-import from a matching docx/pdf in current dir
            auto_path = self._guess_docx_for_lang(iso)
            if auto_path:
                try:
                    parsed = parse_text_material(auto_path)
                    words = list(dict.fromkeys(parsed["chars"] + parsed["words"]))
                    self.current_chars = parsed["chars"]
                    self.current_sentences = parsed["sentences"]
                    self.current_paragraph = parsed["paragraph"]
                    save_wordlist(iso, words)
                    self._set_status(f"Auto-imported {len(words)} words/chars from {auto_path.name}")
                except Exception as e:
                    log(f"Auto-import failed: {e}")
        else:
            # Wordlist is cached on disk — but sentences/paragraph aren't.
            # Re-parse the matching docx/pdf if available so the dialog
            # doesn't say "0 sentences" on every launch after the first.
            self.current_chars = [w for w in words if len(w) <= 2]
            auto_path = self._guess_docx_for_lang(iso)
            if auto_path:
                try:
                    parsed = parse_text_material(auto_path)
                    if parsed["sentences"]:
                        self.current_sentences = parsed["sentences"]
                    if parsed["paragraph"]:
                        self.current_paragraph = parsed["paragraph"]
                    if parsed["chars"]:
                        # Refresh chars from the doc rather than the heuristic
                        self.current_chars = parsed["chars"]
                    log(f"Re-parsed {auto_path.name}: "
                        f"{len(parsed['chars'])} chars, {len(parsed['words'])} words, "
                        f"{len(parsed['sentences'])} sentences, "
                        f"{'paragraph' if parsed['paragraph'] else 'no paragraph'}")
                except Exception as e:
                    log(f"Re-parse failed for {auto_path}: {e}")

        # transcripts/<iso>.txt is the authoritative source if present —
        # this file is shared with app.py (the segmentation tool). It always
        # wins over docx/pdf since the user may have curated it.
        try:
            txt_sentences, txt_paragraph = load_transcripts_file(iso)
            if txt_sentences:
                self.current_sentences = txt_sentences
            if txt_paragraph:
                self.current_paragraph = txt_paragraph
            if txt_sentences or txt_paragraph:
                log(f"transcripts/{iso}.txt overrides → "
                    f"{len(txt_sentences)} sentences, "
                    f"{'paragraph' if txt_paragraph else 'no paragraph'}")
        except Exception as e:
            log(f"Reading transcripts/{iso}.txt failed: {e}")

        self.current_words = words
        self._refresh_words_list()
        self._refresh_chars_grid()
        # The right ASR model depends on the language — update the indicator
        self._refresh_model_status()

    def _guess_docx_for_lang(self, iso):
        """Try to find a matching Dr__Abhishek_..._Text_Material.{docx,pdf} in APP_DIR."""
        lang_token = {"en": "English", "hi": "Hindi", "gu": "Gujarati", "mr": "Marathi"}.get(iso, "")
        if not lang_token:
            return None
        # Prefer docx (faster, cleaner extraction) over pdf
        for ext in ("docx", "pdf"):
            for p in APP_DIR.glob(f"*{lang_token}*Text*Material*.{ext}"):
                return p
        return None

    def _refresh_words_list(self):
        flt = self.txt_filter.text().strip().lower()
        self.words_list.clear()
        for w in self.current_words:
            if flt and flt not in w.lower():
                continue
            self.words_list.addItem(QListWidgetItem(w))

    def _refresh_chars_grid(self):
        # Clear grid
        while self.chars_grid.count():
            item = self.chars_grid.takeAt(0)
            w = item.widget()
            if w:
                w.deleteLater()
        cols = 8
        for i, ch in enumerate(self.current_chars):
            b = QPushButton(ch)
            b.setStyleSheet("QPushButton{background:#ede9fe;padding:4px;font-size:14px;font-weight:bold;min-width:32px}")
            b.clicked.connect(lambda _, c=ch: self._set_current_label(c))
            self.chars_grid.addWidget(b, i // cols, i % cols)
        self.chars_grid.setRowStretch((len(self.current_chars) // cols) + 1, 1)

    def on_import_docx(self):
        iso, _ = self._current_lang()
        # File filter — both docx and pdf, with a combined option as default
        if _PDF:
            filt = (
                "Text material (*.docx *.pdf);;"
                "Word Document (*.docx);;"
                "PDF (*.pdf);;"
                "All Files (*)"
            )
        else:
            filt = "Word Document (*.docx);;All Files (*)"
        path, _ = QFileDialog.getOpenFileName(
            self,
            "Pick a .docx OR .pdf text-material file" +
            ("" if _PDF else "  (pip install pypdf for PDF support)"),
            str(APP_DIR), filt,
        )
        if not path:
            return
        try:
            parsed = parse_text_material(path)
            new_words = list(dict.fromkeys(parsed["chars"] + parsed["words"]))
            # Merge with existing wordlist
            existing = load_wordlist(iso)
            merged = list(dict.fromkeys(existing + new_words))
            save_wordlist(iso, merged)
            self.current_words = merged
            self.current_chars = parsed["chars"]
            self.current_sentences = parsed["sentences"]
            self.current_paragraph = parsed["paragraph"]
            # Persist sentences+paragraph to transcripts/<iso>.txt so app.py
            # (and the next session here) can use them. Don't clobber an
            # existing transcripts file if the import returned nothing.
            if parsed["sentences"] or parsed["paragraph"]:
                try:
                    save_transcripts_file(iso, parsed["sentences"], parsed["paragraph"])
                    log(f"Wrote transcripts/{iso}.txt")
                except Exception as e:
                    log(f"Could not write transcripts/{iso}.txt: {e}")
            self._refresh_words_list()
            self._refresh_chars_grid()
            self._set_status(
                f"Imported from {Path(path).name}: "
                f"{len(parsed['chars'])} chars, {len(parsed['words'])} words, "
                f"{len(parsed['sentences'])} sentences, "
                f"{'paragraph ✓' if parsed['paragraph'] else 'no paragraph'}"
            )
        except Exception as e:
            traceback.print_exc()
            QMessageBox.critical(self, "Import error", str(e))

    def on_add_new_word(self):
        word = self.txt_new_word.text().strip()
        if not word:
            return
        iso, _ = self._current_lang()
        if word in self.current_words:
            self._set_status(f"'{word}' is already in the list")
            return
        self.current_words.append(word)
        save_wordlist(iso, self.current_words)
        self.txt_new_word.clear()
        self._refresh_words_list()
        self._set_status(f"Added '{word}' to {iso}.txt ({len(self.current_words)} total)")

    def on_add_new_word_and_use(self):
        word = self.txt_new_word.text().strip()
        if not word:
            return
        self.on_add_new_word()
        self._set_current_label(word)

    def _set_current_label(self, label):
        self.txt_current_label.setText(label)

    # ----- Audio loading -----
    def on_load_audio(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "Load audio", str(APP_DIR),
            "Audio files (*.wav *.mp3 *.flac *.ogg *.m4a)"
        )
        if not path:
            return
        self._set_status(f"Loading {Path(path).name}…")
        self.btn_load_audio.setEnabled(False)
        t = AudioLoadThread(path)
        t.done.connect(self._on_audio_loaded)
        t.error.connect(self._on_audio_err)
        self._threads["audio"] = t
        t.start()

    def _on_audio_loaded(self, y, sr, y16, path):
        self.y = y
        self.y16 = y16  # 16k version for ASR (may be None if resampling failed)
        self.sr = sr
        self.audio_path = path
        # Clear spectrogram cache since the underlying audio changed
        self._spec_cache.clear()
        # Configure spinboxes to audio duration
        dur = len(y) / sr
        for sp in (self.sp_start, self.sp_end):
            sp.setRange(0.0, dur)
        self.region_start = None
        self.region_end = None
        self._region_patch = None
        self.btn_load_audio.setEnabled(True)
        asr_note = "" if self.y16 is not None else "  (⚠ resample failed: ASR off)"
        self._set_status(f"Loaded: {Path(path).name} ({dur:.2f} s, {sr} Hz){asr_note}")
        self._redraw()
        # If a previous session crashed mid-edit, offer to restore those segments
        self._check_for_autosave_on_load()
        # Refresh the project file list (updates status icons for the just-loaded file)
        if self.project_dir is not None:
            self._refresh_project_list()

    def _on_audio_err(self, msg):
        self.btn_load_audio.setEnabled(True)
        QMessageBox.critical(self, "Audio load failed", msg)
        self._set_status(f"✗ Audio load failed: {msg}")

    # ----- View mode -----
    def _on_view_mode_changed(self, _idx=None):
        new_mode = self.cb_view_mode.currentData()
        if new_mode == self.view_mode:
            return
        self.view_mode = new_mode
        if self.y is not None:
            xlim = self.ax.get_xlim() if self.ax else None
            self._redraw(xlim=xlim)

    # ----- ASR model loading -----
    def _model_for_current_lang(self):
        """Return (engine_key, model_id) for the currently-selected language,
        or None if the required backend isn't installed."""
        iso, _ = self._current_lang()
        return INDIC_MODEL if iso in ("gu", "hi", "mr") else ENGLISH_MODEL

    def _model_display_name(self, mdl):
        if not mdl:
            return "—"
        ek, mid = mdl
        if ek == "indicconformer":
            return "IndicConformer-600M (IIT Madras / AI4Bharat)"
        if ek == "fw":
            return f"faster-whisper {mid}"
        if ek == "hf_whisper":
            return f"Whisper {mid.split('/')[-1]}"
        return mid

    def _refresh_model_status(self):
        """Update the model status pill + Load Model button based on what's loaded."""
        mdl = self._model_for_current_lang()
        wanted = self._model_display_name(mdl) if mdl else "—"
        if not _ASR_AVAILABLE:
            self.lbl_model_status.setText("⚠ ASR backends not installed")
            self.btn_load_model.setEnabled(False)
            self.btn_auto_detect.setEnabled(False)
            return
        if not mdl:
            self.lbl_model_status.setText(f"⚠ No model available for this language")
            self.btn_load_model.setEnabled(False)
            return
        if ENGINE.is_loaded and ENGINE.loaded_key == tuple(mdl):
            self.lbl_model_status.setText(f"✓ Loaded: {wanted}")
            self.btn_load_model.setText("✓ Loaded")
            self.btn_load_model.setStyleSheet(
                "QPushButton{font-weight:bold;color:white;background:#16a34a;"
                "border:1px solid #15803d;padding:5px 14px;border-radius:4px}"
            )
        else:
            self.lbl_model_status.setText(f"Need: {wanted}")
            self.btn_load_model.setText("⬇ Load Model")
            self.btn_load_model.setStyleSheet(
                "QPushButton{font-weight:bold;color:white;background:#2563eb;"
                "border:1px solid #1d4ed8;padding:5px 14px;border-radius:4px}"
                "QPushButton:hover{background:#1d4ed8}"
                "QPushButton:disabled{color:#9ca3af;background:#f3f4f6}"
            )

    def on_load_model(self):
        existing = self._threads.get("model")
        if existing and existing.isRunning():
            QMessageBox.information(self, "", "Model is still loading — please wait.")
            return
        mdl = self._model_for_current_lang()
        if not mdl:
            QMessageBox.critical(self, "", "No model available for this language.")
            return
        ek, mid = mdl
        self.btn_load_model.setText("Loading…")
        self.btn_load_model.setEnabled(False)
        self.btn_auto_detect.setEnabled(False)
        self._model_load_t0 = time.time()
        self._model_tick_timer = QTimer(self)
        self._model_tick_timer.timeout.connect(
            lambda: self.lbl_model_status.setText(
                f"Loading {self._model_display_name(mdl)}… "
                f"{int(time.time() - self._model_load_t0)}s"
            )
        )
        self._model_tick_timer.start(500)
        t = ModelLoadThread(ek, mid)
        t.done.connect(self._on_model_loaded)
        t.error.connect(self._on_model_load_err)
        self._threads["model"] = t
        t.start()

    def _stop_model_tick(self):
        if self._model_tick_timer is not None:
            self._model_tick_timer.stop()
            self._model_tick_timer = None

    def _on_model_loaded(self, ok, msg, elapsed):
        self._stop_model_tick()
        self.btn_load_model.setEnabled(True)
        self.btn_auto_detect.setEnabled(True)
        if ok:
            self._refresh_model_status()
            self._set_status(f"✓ {msg}")
        else:
            self.btn_load_model.setText("⬇ Load Model")
            self._refresh_model_status()
            QMessageBox.critical(self, "Model load failed", msg)

    def _on_model_load_err(self, msg):
        self._stop_model_tick()
        self.btn_load_model.setText("⬇ Load Model")
        self.btn_load_model.setEnabled(True)
        self.btn_auto_detect.setEnabled(True)
        self._refresh_model_status()
        QMessageBox.critical(self, "Model load failed", msg)

    # ----- Auto-detect (VAD + ASR) -----
    def _vad_params(self):
        return {
            "silence_ms": self.sp_vad_sil.value(),
            "thresh_db":  self.sp_vad_db.value(),
            "min_ms":     self.sp_vad_min.value(),
            "max_ms":     self.sp_vad_max.value(),
        }

    def on_auto_detect(self):
        self._run_auto_detect(do_asr=True)

    def on_vad_only(self):
        self._run_auto_detect(do_asr=False)

    def _run_auto_detect(self, do_asr):
        if self.y is None:
            QMessageBox.warning(self, "", "Load an audio file first.")
            return
        if do_asr:
            if not _ASR_AVAILABLE:
                QMessageBox.warning(
                    self, "ASR not installed",
                    "ASR backends not installed.\n\n"
                    "Install one:\n"
                    "    pip install faster-whisper       (English)\n"
                    "    pip install torch transformers   (Indic + English)"
                )
                return
            mdl = self._model_for_current_lang()
            if not mdl:
                QMessageBox.warning(self, "", "No model available for this language.")
                return
            if not ENGINE.is_loaded or ENGINE.loaded_key != tuple(mdl):
                ret = QMessageBox.question(
                    self, "Load model?",
                    f"The required model isn't loaded:\n  {self._model_display_name(mdl)}\n\n"
                    "Load it now? (This can take a minute on first run.)"
                )
                if ret == QMessageBox.Yes:
                    self.on_load_model()
                    self._set_status(
                        "Click 🤖 Auto-Detect again once the model finishes loading."
                    )
                return
            if self.y16 is None:
                QMessageBox.warning(
                    self, "",
                    "Cannot run ASR: 16 kHz resampled audio unavailable.\n"
                    "Falling back to VAD-only mode might still work."
                )
                return

        # Wipe out any in-progress segment list and start fresh
        self.segments = []
        self._selected_seg_idx = None
        self._refresh_seg_list()
        self._redraw(xlim=self.ax.get_xlim() if self.y is not None else None)

        iso, name = self._current_lang()
        t = AutoSegmentThread(
            self.y, self.sr, self.y16, iso, name, self._vad_params(), do_asr=do_asr,
        )
        t.segment.connect(self._on_auto_segment)
        t.progress.connect(self._on_auto_progress)
        t.done.connect(self._on_auto_done)
        t.error.connect(self._on_auto_err)
        self._threads["auto"] = t

        self._set_auto_busy(True)
        mode_str = "VAD + ASR" if do_asr else "VAD only"
        self._set_status(f"🤖 Running {mode_str}…")
        t.start()

    def on_stop_auto_detect(self):
        t = self._threads.get("auto")
        if t and t.isRunning():
            t.request_stop()
            self._set_status("Stopping after current segment…")

    def _set_auto_busy(self, busy):
        # Toggle the buttons that conflict with an in-progress auto-detect
        for b in (self.btn_auto_detect, self.btn_vad_only, self.btn_align_gt,
                  self.btn_predict_gt,
                  self.btn_load_audio, self.btn_load_csv, self.btn_load_model):
            b.setEnabled(not busy)
        self.btn_stop_auto.setEnabled(busy)

    def _on_auto_segment(self, seg):
        self.segments.append(seg)
        self._refresh_seg_list()
        # Cheap incremental update of the waveform
        if self.y is not None:
            self._redraw(xlim=self.ax.get_xlim())

    def _on_auto_progress(self, done, total):
        self._set_status(f"🤖 Processing {done}/{total} segments…")

    def _on_auto_done(self, total, elapsed):
        self._set_auto_busy(False)
        speed = total / elapsed if elapsed > 0 else 0
        self._set_status(
            f"✓ Auto-detect finished: {total} segments in {elapsed:.1f}s ({speed:.1f}/s). "
            "Review and edit labels, then Save CSV."
        )

    def _on_auto_err(self, msg):
        self._set_auto_busy(False)
        QMessageBox.critical(self, "Auto-detect failed", msg)
        self._set_status(f"✗ Auto-detect failed: {msg}")

    # ----- Forced alignment: GT labels -> audio timestamps -----
    def _build_wordlist_gt_sequence(self):
        """Construct the expected GT label sequence from the loaded wordlist:
        characters first, then words, then s1..s5, then paragraph (if any).
        Matches the order speakers read in the Dr. Abhishek text material."""
        seq = []
        # Characters/alphabets
        seq.extend(self.current_chars)
        # Words (skip ones already in chars to avoid duplicates)
        seen = set(seq)
        for w in self.current_words:
            if w not in seen:
                seq.append(w); seen.add(w)
        # Sentences in s1, s2, ... order
        s_keys = sorted(
            [k for k in self.current_sentences if re.match(r"^s\d+$", k)],
            key=lambda k: int(k[1:]),
        )
        for k in s_keys:
            txt = self.current_sentences[k]
            if txt:
                seq.append(txt)
        # Paragraph
        if self.current_paragraph:
            seq.append(self.current_paragraph)
        return seq

    def _ask_gt_source(self):
        """Pop a small dialog: pick wordlist or pasted text as the GT source.
        Returns a list of GT labels, or None if cancelled."""
        wordlist_seq = self._build_wordlist_gt_sequence()

        dlg = QDialog(self)
        dlg.setWindowTitle("Align GT to Audio — pick source")
        dlg.setMinimumWidth(620)
        lay = QVBoxLayout(dlg)

        lay.addWidget(QLabel(
            "<b>Pick the expected label sequence</b>.<br>"
            "These are the labels the speaker is expected to read, in order. "
            "VAD + ASR will line them up to the audio."
        ))

        # --- Option 1: wordlist ---
        rb_wordlist = QRadioButton(
            f"Use the loaded wordlist  "
            f"({len(self.current_chars)} chars, "
            f"{len(self.current_words)} words, "
            f"{len(self.current_sentences)} sentences"
            f"{', paragraph' if self.current_paragraph else ''})"
        )
        rb_wordlist.setChecked(bool(wordlist_seq))
        rb_wordlist.setEnabled(bool(wordlist_seq))
        lay.addWidget(rb_wordlist)
        if wordlist_seq:
            preview_label = QLabel(
                "  Preview: " + ", ".join(
                    (w if len(w) <= 18 else w[:17] + "…") for w in wordlist_seq[:8]
                ) + (f"  … (+{len(wordlist_seq) - 8} more)" if len(wordlist_seq) > 8 else "")
            )
            preview_label.setStyleSheet("color:#6b7280;font-size:11px;padding-left:24px")
            preview_label.setWordWrap(True)
            lay.addWidget(preview_label)
        else:
            warn = QLabel("  ⚠ No wordlist loaded — Import one from .docx first, or use 'Paste custom list' below.")
            warn.setStyleSheet("color:#b45309;font-size:11px;padding-left:24px")
            lay.addWidget(warn)

        lay.addSpacing(6)

        # --- Option 2: paste / type ---
        rb_paste = QRadioButton("Paste a custom list (one label per line)")
        rb_paste.setChecked(not wordlist_seq)
        lay.addWidget(rb_paste)

        txt_paste = QPlainTextEdit()
        txt_paste.setPlaceholderText(
            "one label per line, e.g.:\n"
            "    अ\n    आ\n    इ\n    apple\n    elephant\n    "
            "The cat sat on the mat.\n"
        )
        txt_paste.setMinimumHeight(140)
        lay.addWidget(txt_paste)

        # Mutual-exclusion (radios already do this but make it explicit)
        group = QButtonGroup(dlg)
        group.addButton(rb_wordlist); group.addButton(rb_paste)

        # OK/Cancel
        bb = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        bb.accepted.connect(dlg.accept)
        bb.rejected.connect(dlg.reject)
        lay.addWidget(bb)

        if dlg.exec_() != QDialog.Accepted:
            return None

        if rb_wordlist.isChecked():
            return wordlist_seq
        # Paste mode: split on newlines, drop empty
        raw = txt_paste.toPlainText()
        labels = [ln.strip() for ln in raw.splitlines() if ln.strip()]
        if not labels:
            QMessageBox.warning(self, "", "No labels entered.")
            return None
        return labels

    def on_align_gt(self):
        if self.y is None:
            QMessageBox.warning(self, "", "Load an audio file first.")
            return
        if not _ASR_AVAILABLE:
            QMessageBox.warning(
                self, "ASR not installed",
                "Forced alignment needs an ASR backend.\n\n"
                "Install one:\n"
                "    pip install faster-whisper       (English)\n"
                "    pip install torch transformers   (Indic + English)"
            )
            return
        # Same model gating as auto-detect
        mdl = self._model_for_current_lang()
        if not mdl:
            QMessageBox.warning(self, "", "No model available for this language.")
            return
        if not ENGINE.is_loaded or ENGINE.loaded_key != tuple(mdl):
            ret = QMessageBox.question(
                self, "Load model?",
                f"The required model isn't loaded:\n  {self._model_display_name(mdl)}\n\n"
                "Load it now? (This can take a minute on first run.)"
            )
            if ret == QMessageBox.Yes:
                self.on_load_model()
                self._set_status(
                    "Click 📐 Align GT to Audio again once the model finishes loading."
                )
            return
        if self.y16 is None:
            QMessageBox.warning(
                self, "",
                "Cannot align: 16 kHz resampled audio unavailable."
            )
            return

        gt_labels = self._ask_gt_source()
        if not gt_labels:
            return

        # Warn before wiping existing segments
        if self.segments:
            ret = QMessageBox.question(
                self, "Replace existing segments?",
                f"You have {len(self.segments)} segments. Alignment will "
                "replace them with auto-aligned ones.\n\n"
                "(Your work is also autosaved to "
                "annotations/.autosave/ — you can recover it if needed.)",
                QMessageBox.Yes | QMessageBox.No, QMessageBox.No,
            )
            if ret != QMessageBox.Yes:
                return

        self.segments = []
        self._selected_seg_idx = None
        self._refresh_seg_list()
        self._redraw(xlim=self.ax.get_xlim() if self.y is not None else None)

        iso, name = self._current_lang()
        t = AlignThread(
            self.y, self.sr, self.y16, iso, name,
            gt_labels, self._vad_params(),
        )
        t.segment.connect(self._on_align_segment)
        t.progress.connect(self._on_align_progress)
        t.done.connect(self._on_align_done)
        t.unmatched_gt.connect(self._on_align_unmatched)
        t.error.connect(self._on_align_err)
        self._threads["align"] = t
        self._set_auto_busy(True)
        self._set_status(
            f"📐 Aligning {len(gt_labels)} GT labels to audio "
            f"(VAD → {self._model_display_name(mdl)} → DP alignment)…"
        )
        self._pending_align_total = len(gt_labels)
        t.start()

    def _on_align_segment(self, seg):
        # Stash the similarity score for the tooltip but don't write it to disk
        self.segments.append(seg)
        self._refresh_seg_list()
        if self.y is not None:
            self._redraw(xlim=self.ax.get_xlim())

    def _on_align_progress(self, done, total, stage):
        self._set_status(f"📐 {stage}  ({done}/{total})")

    def _on_align_done(self, matched, total_gt, unmatched_chunks, elapsed):
        self._set_auto_busy(False)
        coverage = (matched / total_gt * 100) if total_gt else 0
        parts = [
            f"✓ Alignment done in {elapsed:.1f}s",
            f"matched {matched}/{total_gt} GT labels ({coverage:.0f}%)",
        ]
        if unmatched_chunks:
            parts.append(f"{unmatched_chunks} audio chunks unmatched")
        parts.append("review and Save")
        self._set_status("  │  ".join(parts))

    def _on_align_unmatched(self, unmatched):
        # Show a separate message listing GT labels we couldn't find in the
        # audio — annotator can add them manually if they belong there.
        if not unmatched:
            return
        preview = ", ".join(
            (w if len(w) <= 22 else w[:21] + "…") for w in unmatched[:15]
        )
        if len(unmatched) > 15:
            preview += f"  … (+{len(unmatched) - 15} more)"
        QMessageBox.information(
            self, "Some GT labels weren't matched",
            f"{len(unmatched)} expected labels couldn't be aligned to any "
            f"audio chunk:\n\n  {preview}\n\n"
            "These are usually labels the speaker skipped, mispronounced "
            "beyond recognition, or that VAD merged into adjacent chunks. "
            "Add them manually by marking the region and typing the label."
        )

    def _on_align_err(self, msg):
        self._set_auto_busy(False)
        QMessageBox.critical(self, "Alignment failed", msg)
        self._set_status(f"✗ Alignment failed: {msg}")

    # ----- Predict on GT: for each loaded GT segment, run ASR and capture
    #       (predicted label, model's own start/end, deltas vs GT). This is
    #       the "score my ground truth" workflow.
    def on_predict_on_gt(self):
        if not self.segments:
            QMessageBox.warning(
                self, "",
                "No segments to predict on.\n\n"
                "First click 📥 Load GT and pick your annotations file "
                "(e.g. Female.xlsx). Predict-on-GT will then run ASR on "
                "each of those segments."
            )
            return
        if self.y is None or self.y16 is None:
            QMessageBox.warning(self, "", "Load the matching audio file first.")
            return
        if not _ASR_AVAILABLE:
            QMessageBox.warning(
                self, "ASR not installed",
                "Predicting on GT needs an ASR backend.\n\n"
                "Install one:\n"
                "    pip install faster-whisper       (English)\n"
                "    pip install torch transformers   (Indic + English)"
            )
            return
        mdl = self._model_for_current_lang()
        if not mdl:
            QMessageBox.warning(self, "", "No model available for this language.")
            return
        if not ENGINE.is_loaded or ENGINE.loaded_key != tuple(mdl):
            ret = QMessageBox.question(
                self, "Load model?",
                f"The required model isn't loaded:\n  {self._model_display_name(mdl)}\n\n"
                "Load it now? (This can take a minute on first run.)"
            )
            if ret == QMessageBox.Yes:
                self.on_load_model()
                self._set_status(
                    "Click 🎯 Predict on GT again once the model finishes loading."
                )
            return

        iso, name = self._current_lang()
        t = PredictTimingsThread(
            self.segments, self.y16, iso, name,
            ref_paragraph=self.current_paragraph,
        )
        t.segment_updated.connect(self._on_predict_gt_segment)
        t.progress.connect(self._on_predict_gt_progress)
        t.done.connect(self._on_predict_gt_done)
        t.error.connect(self._on_predict_gt_err)
        self._threads["predict_gt"] = t
        self._set_auto_busy(True)
        self._predict_gt_t0 = time.time()
        self._set_status(
            f"🎯 Running {self._model_display_name(mdl)} on "
            f"{len(self.segments)} GT segments…"
        )
        t.start()

    def _on_predict_gt_segment(self, idx, updated_seg):
        """A single GT segment finished — update in place and re-render."""
        if 0 <= idx < len(self.segments):
            self.segments[idx] = updated_seg
            self._refresh_seg_list()
            if self.y is not None:
                self._redraw(xlim=self.ax.get_xlim())
            # Refresh the diff panel if this is the selected segment
            if idx == self._selected_seg_idx:
                self._update_diff_panel()

    def _on_predict_gt_progress(self, done, total):
        elapsed = time.time() - getattr(self, "_predict_gt_t0", time.time())
        speed = done / elapsed if elapsed > 0 else 0
        eta = (total - done) / speed if speed > 0 else 0
        self._set_status(
            f"🎯 Predicting {done}/{total}  ({speed:.1f}/s, ETA {eta:.0f}s)"
        )

    def _on_predict_gt_done(self, total, skipped, elapsed):
        self._set_auto_busy(False)
        # Compute simple accuracy stats so the user gets immediate feedback
        n_match = 0; n_diff = 0; n_total = 0
        big_time_errors = 0
        partial_paras = []   # (label, spoken, total, frac)
        for s in self.segments:
            pred = (s.get("prediction") or "").strip()
            label = (s.get("label") or "").strip()
            if pred:
                n_total += 1
                if _fuzzy_score(pred, label) >= 0.85:
                    n_match += 1
                else:
                    n_diff += 1
            # Flag big timing errors (>0.5s shift in either edge)
            for k in ("start_err", "end_err"):
                v = s.get(k)
                if v is not None and abs(v) > 0.5:
                    big_time_errors += 1
                    break
            # Collect partial paragraphs
            frac = s.get("_para_fraction")
            if frac is not None and frac < 0.95:
                partial_paras.append((
                    label, s.get("_para_spoken_words", 0),
                    s.get("_para_total_words", 0), frac,
                ))
        pct = (n_match / n_total * 100) if n_total else 0
        speed = total / elapsed if elapsed > 0 else 0
        msg = (
            f"✓ Predict-on-GT done in {elapsed:.1f}s ({speed:.1f}/s)  │  "
            f"{n_match}/{n_total} labels match ({pct:.0f}%)  │  "
            f"{n_diff} differ"
        )
        if big_time_errors:
            msg += f"  │  {big_time_errors} have >0.5s timing error"
        if partial_paras:
            msg += f"  │  {len(partial_paras)} partial paragraph(s)"
        if skipped:
            msg += f"  │  {skipped} skipped"
        self._set_status(msg)

        # If any paragraph was only partially read, tell the annotator
        # explicitly — this is easy to miss and important for data quality.
        if partial_paras:
            lines = []
            for label, spoken, total_w, frac in partial_paras:
                lines.append(
                    f"  • {label}: spoke {spoken}/{total_w} words "
                    f"({frac*100:.0f}%)"
                )
            QMessageBox.information(
                self, "Partial paragraph(s) detected",
                "The speaker appears to have read only part of the "
                "paragraph in these segments:\n\n"
                + "\n".join(lines)
                + "\n\nThe waveform now tags these as partial. If a paragraph "
                "was split into pieces, consider relabeling them p1, p2, p3… "
                "in reading order."
            )

    def _on_predict_gt_err(self, msg):
        self._set_auto_busy(False)
        QMessageBox.critical(self, "Predict-on-GT failed", msg)
        self._set_status(f"✗ Predict-on-GT failed: {msg}")

    # ----- Waveform drawing (Audacity-style) -----
    # Palette tuned to mimic Audacity's default look on a dark background.
    _AUDACITY = {
        "bg":          "#1f2a36",  # dark navy figure / axes background
        "envelope":    "#5294c4",  # outer peak (min/max) — medium blue
        "rms":         "#a0d8f5",  # inner RMS overlay — bright cyan-blue
        "sample_line": "#5fb4e2",  # connecting line at sample-level zoom
        "sample_dot":  "#cfe9fb",  # individual sample dots at deep zoom
        "zero":        "#3a4d5e",  # center zero line
        "grid":        "#3a4d5e",  # gridlines (same as zero line, subtle)
        "axis_text":   "#c4d4e2",  # tick labels / axis names
        "title":       "#e2ecf5",
        "spine":       "#3a4d5e",
        # Selections / segments
        "region":      "#fde047",  # current selection — Audacity yellow
        "region_edge": "#facc15",
        "seg":         "#22c55e",  # added segments — green
        "seg_sel":     "#fb923c",  # selected segment — orange
        "label_bg":    "#0f1923",  # background of label boxes
        "label_text":  "#fff7ed",
    }

    @staticmethod
    def _amp_to_db_y(values, floor_db=-60.0):
        """Map signed linear amplitude into a normalized dB display y in [-1,+1].
        +1   = 0 dB peak       (loudest)
        +0   = `floor_db` floor  (silence)
        -1   = 0 dB peak (negative-going sample)
        Used by the 'Waveform (dB)' view so the wave shape is preserved but
        compressed onto a perceptually-uniform log scale, making quiet parts
        visible alongside loud ones — same trick as Audacity's Waveform (dB).
        """
        floor_lin = 10.0 ** (floor_db / 20.0)
        abs_vals = np.maximum(np.abs(values), floor_lin)
        db = 20.0 * np.log10(abs_vals)
        nrm = (db - floor_db) / (-floor_db)        # 0 at floor → 1 at 0 dB
        return np.sign(values) * nrm

    def _compute_envelope(self, x0, x1, n_pixels=2000):
        """Compute peak (min/max) and RMS envelopes for the visible range.

        Returns:
          {'mode': 'samples', 't': [..], 'y': [..]}        when very zoomed in
          {'mode': 'envelope', 't': [..], 'mins': [..],
           'maxs': [..], 'rms': [..]}                       otherwise
        """
        if self.y is None or self.sr is None:
            return None
        sr = self.sr
        total = len(self.y)
        start = max(0, int(x0 * sr))
        end = min(total, int(x1 * sr))
        if end <= start:
            return None
        visible = self.y[start:end]
        n_visible = end - start

        # If we have fewer samples than pixels, show every individual sample.
        if n_visible <= n_pixels:
            t = (start + np.arange(n_visible)) / sr
            return {"mode": "samples", "t": t, "y": visible}

        # Otherwise reshape into bins of samples_per_bin and compute min/max/RMS.
        samples_per_bin = max(1, n_visible // n_pixels)
        n_bins = n_visible // samples_per_bin
        trimmed = visible[: n_bins * samples_per_bin]
        bins = trimmed.reshape(n_bins, samples_per_bin)
        mins = bins.min(axis=1)
        maxs = bins.max(axis=1)
        # RMS in float64 to avoid overflow on int audio (we already loaded float32)
        rms = np.sqrt(np.mean(bins.astype(np.float64) ** 2, axis=1)).astype(np.float32)
        t = (start + np.arange(n_bins) * samples_per_bin + samples_per_bin / 2.0) / sr
        return {"mode": "envelope", "t": t, "mins": mins, "maxs": maxs, "rms": rms}

    def _get_or_compute_spec(self, mode):
        """Cache the full spectrogram per (mode, audio) so repeat redraws are free."""
        if self.y is None:
            return None
        cache_key = (mode, id(self.y))
        cached = self._spec_cache.get(cache_key)
        if cached is not None:
            return cached
        try:
            # Use 16k-resampled audio when available — faster, and limits the
            # plot to 0–8 kHz which is what speech actually lives in.
            audio = self.y16 if self.y16 is not None else self.y
            spec_sr = SR16 if self.y16 is not None else self.sr
            freqs, times, S = compute_spectrogram(audio, spec_sr, mode=mode)
            self._spec_cache[cache_key] = (freqs, times, S)
            return self._spec_cache[cache_key]
        except Exception as e:
            log(f"Spectrogram compute failed: {e}")
            return None

    def _redraw(self, xlim=None):
        P = self._AUDACITY
        self.ax.clear()
        # Dark Audacity-style background on both figure and axes
        self.fig.patch.set_facecolor(P["bg"])
        self.ax.set_facecolor(P["bg"])

        if self.y is None:
            self.canvas.draw_idle()
            self._update_view_label()
            return

        max_dur = len(self.y) / self.sr
        if xlim is None:
            xlim = (0.0, max_dur)
        xlim = (max(0.0, xlim[0]), min(max_dur, xlim[1]))
        if xlim[1] <= xlim[0]:
            xlim = (0.0, max_dur)

        # ---- View mode branch ----
        is_spec = self.view_mode.startswith("spec_")
        if is_spec:
            spec_mode = {"spec_linear": "linear",
                         "spec_logdb":  "log_db",
                         "spec_mel":    "mel"}[self.view_mode]
            cached = self._get_or_compute_spec(spec_mode)
            if cached is None:
                self.ax.text(0.5, 0.5,
                             "Spectrogram unavailable\n(install scipy)",
                             ha="center", va="center",
                             color=P["axis_text"], fontsize=12,
                             transform=self.ax.transAxes)
                self.canvas.draw_idle()
                return
            freqs, times, S = cached
            # Clip the color range for log-dB views so silent areas don't wash
            # out the visible dynamic range (Audacity does the same trick).
            if spec_mode in ("log_db", "mel"):
                vmax = float(S.max())
                vmin = vmax - 80.0          # 80 dB dynamic range
                cmap = "magma"
            else:
                vmax = float(np.percentile(S, 99.5))
                vmin = 0.0
                cmap = "magma"
            # imshow with extent so the spectrogram aligns to time/frequency axes
            self.ax.imshow(
                S, aspect="auto", origin="lower",
                extent=[float(times[0]), float(times[-1]),
                        float(freqs[0]), float(freqs[-1])],
                vmin=vmin, vmax=vmax, cmap=cmap,
                interpolation="nearest", zorder=2,
            )
            # Markers extend full vertical range
            y_top = float(freqs[-1])
            y_bot = float(freqs[0])
            y_span = max(1.0, y_top - y_bot)
            self.ax.set_ylim(y_bot, y_top)
            self.ax.set_ylabel(
                "Frequency (Hz)" if spec_mode != "mel" else "Frequency (Hz, mel-spaced)",
                color=P["axis_text"], fontsize=10,
            )
        else:
            # === Waveform path — two flavours depending on view_mode ===
            env = self._compute_envelope(xlim[0], xlim[1], n_pixels=2200)

            if self.view_mode == "wave_db":
                # Logarithmic dB waveform (like Audacity's "Waveform (dB)").
                # Y axis is dB on both halves with -60 dB at the center line,
                # 0 dB at the top and bottom edges. Quiet parts become visible.
                floor_db = -60.0
                # Display always uses normalized [-1, +1] range internally so
                # the segment-marker / region / playback-cursor logic below
                # doesn't need to know about dB.
                y_top = 1.0
                y_bot = -1.0
                y_span = 2.0
                if env is not None:
                    if env["mode"] == "samples":
                        y_disp = self._amp_to_db_y(env["y"], floor_db)
                        self.ax.plot(env["t"], y_disp, color=P["sample_line"],
                                     lw=1.2, zorder=2)
                        if len(env["t"]) <= 400:
                            self.ax.scatter(env["t"], y_disp,
                                            color=P["sample_dot"],
                                            s=14, zorder=4, edgecolors="none")
                    else:
                        mins_d = self._amp_to_db_y(env["mins"], floor_db)
                        maxs_d = self._amp_to_db_y(env["maxs"], floor_db)
                        rms_d = self._amp_to_db_y(env["rms"], floor_db)
                        # Outer (peak) envelope, top half goes 0..maxs_d, bottom
                        # mirrors. We use fill_between with the actual min/max
                        # display values so the shape preserves asymmetry.
                        self.ax.fill_between(
                            env["t"], mins_d, maxs_d,
                            color=P["envelope"], linewidth=0, zorder=2,
                        )
                        # RMS overlay (positive-magnitude only, mirrored)
                        self.ax.fill_between(
                            env["t"], -rms_d, rms_d,
                            color=P["rms"], linewidth=0, zorder=3,
                        )
                # Center line = noise floor (-inf dB / -60 dB in our case)
                self.ax.axhline(0, color=P["zero"], lw=0.8, zorder=4)
                self.ax.set_ylim(y_bot, y_top)
                # Custom dB ticks on both halves
                std_db = [0, -6, -12, -24, -36, -48]
                yticks, ylabels = [], []
                for db in std_db:
                    nrm = (db - floor_db) / -floor_db   # 0 at floor → 1 at 0dB
                    yticks.append(+nrm); ylabels.append(f"{db}")
                    if db != int(floor_db):
                        yticks.append(-nrm); ylabels.append(f"{db}")
                # Center tick = the floor itself
                yticks.append(0.0); ylabels.append(f"≤{int(floor_db)}")
                pairs = sorted(zip(yticks, ylabels))
                self.ax.set_yticks([p[0] for p in pairs])
                self.ax.set_yticklabels([p[1] for p in pairs])
                self.ax.set_ylabel(
                    f"Amplitude (dB,  floor = {int(floor_db)} dB)",
                    color=P["axis_text"], fontsize=10,
                )
            else:
                # Linear waveform (original Audacity look)
                ymax = float(max(np.max(np.abs(self.y)), 1e-6))
                y_top = ymax * 1.1
                y_bot = -y_top
                y_span = y_top - y_bot
                if env is not None:
                    if env["mode"] == "samples":
                        self.ax.plot(env["t"], env["y"], color=P["sample_line"],
                                     lw=1.2, zorder=2)
                        if len(env["t"]) <= 400:
                            self.ax.scatter(env["t"], env["y"],
                                            color=P["sample_dot"],
                                            s=14, zorder=4, edgecolors="none")
                    else:
                        self.ax.fill_between(
                            env["t"], env["mins"], env["maxs"],
                            color=P["envelope"], linewidth=0, zorder=2,
                        )
                        self.ax.fill_between(
                            env["t"], -env["rms"], env["rms"],
                            color=P["rms"], linewidth=0, zorder=3,
                        )
                self.ax.axhline(0, color=P["zero"], lw=0.8, zorder=4)
                self.ax.set_ylim(y_bot, y_top)
                self.ax.set_ylabel("Amplitude", color=P["axis_text"], fontsize=10)

        # Style: axis text, spines, grid (shared by both modes)
        self.ax.set_xlim(*xlim)
        self.ax.grid(axis="x", color=P["grid"], alpha=0.35,
                     linestyle="-", linewidth=0.5, zorder=1)
        self.ax.tick_params(colors=P["axis_text"], which="both", labelsize=9)
        for spine in self.ax.spines.values():
            spine.set_color(P["spine"])
        self.ax.set_xlabel("Time (s)", color=P["axis_text"], fontsize=10)

        # Existing segments (drawn over the waveform/spectrogram, under selection)
        iso, _ = self._current_lang()
        fp = get_font(iso)
        # Place bottom-of-axes markers proportional to current y-range
        marker_y = y_bot + y_span * 0.03
        tag_y = y_bot + y_span * 0.07                  # type tag — just above bar
        gt_label_y = y_top - y_span * 0.10            # GT (human) badge — top
        pred_label_y = y_bot + y_span * 0.16          # PRED (model) badge — bottom
        text_kw = {"fontproperties": fp} if fp else {}
        show_all_preds = self.chk_show_preds.isChecked()
        GT_COLOR = "#16a34a"      # green — human-authored ground truth
        PRED_COLOR = "#3b82f6"    # blue — model prediction
        DIFF_INDICATOR = "#fbbf24"  # amber — for the "diff" tick
        PREDTIME_COLOR = "#f472b6"  # pink — model's predicted start/end
        # Per-type tag colors so char/word/sent/para are visually distinct
        TYPE_TAG_COLOR = {
            "char": "#7c3aed", "word": "#0284c7",
            "sentence": "#16a34a", "paragraph": "#ea580c",
        }
        for i, seg in enumerate(self.segments):
            s, e = seg["start"], seg["end"]
            if e < xlim[0] or s > xlim[1]:
                continue
            is_sel = (i == self._selected_seg_idx)
            color = P["seg_sel"] if is_sel else P["seg"]
            pred = seg.get("prediction", "") or ""
            label = seg.get("label", "") or ""
            differs = bool(pred) and pred.strip() != label.strip()
            seg_type = classify_segment_type(label)
            tag = short_type_tag(seg)

            # Subtle background tint
            self.ax.axvspan(s, e, alpha=0.12, color=color, zorder=5)
            # Bottom-of-axes bar marker (thicker for the selected one)
            self.ax.plot([s, e], [marker_y, marker_y],
                         lw=4.0 if is_sel else 3.0,
                         color=color, solid_capstyle="butt", zorder=6)

            # --- TYPE TAG on the bar (char / word / s1..s5 / p1..p3 / para) ---
            # Always shown (small) so the annotator can see at a glance what
            # kind of unit each segment is, without opening any panel.
            tag_color = TYPE_TAG_COLOR.get(seg_type, "#64748b")
            self.ax.text(
                (s + e) / 2, tag_y, tag,
                ha="center", va="bottom",
                fontsize=8 if not is_sel else 9,
                color="#ffffff", fontweight="bold",
                bbox=dict(facecolor=tag_color,
                          alpha=0.95 if is_sel else 0.8,
                          edgecolor="none", boxstyle="round,pad=0.18"),
                zorder=7, **text_kw,
            )

            # --- Predicted timestamps (from 🎯 Predict on GT) ---
            # Draw the model's own start/end as pink brackets, plus the timing
            # error text. This makes the GT-vs-model timing discrepancy visible
            # right on the waveform.
            pstart = seg.get("pred_start")
            pend = seg.get("pred_end")
            if pstart is not None and pend is not None and (is_sel or show_all_preds):
                if not (pend < xlim[0] or pstart > xlim[1]):
                    pred_marker_y = marker_y + y_span * 0.05
                    self.ax.plot([pstart, pend], [pred_marker_y, pred_marker_y],
                                 lw=2.0, color=PREDTIME_COLOR,
                                 solid_capstyle="butt", zorder=6,
                                 linestyle=(0, (4, 2)))
                    # Vertical ticks at predicted edges
                    for xpos in (pstart, pend):
                        self.ax.plot([xpos, xpos],
                                     [pred_marker_y - y_span * 0.015,
                                      pred_marker_y + y_span * 0.015],
                                     lw=1.5, color=PREDTIME_COLOR, zorder=6)
                    if is_sel:
                        se = seg.get("start_err"); ee = seg.get("end_err")
                        if se is not None and ee is not None:
                            err_txt = (f"Δstart {se:+.2f}s   Δend {ee:+.2f}s"
                                       + ("  (est.)" if seg.get("pred_estimated") else ""))
                            self.ax.text(
                                (pstart + pend) / 2, pred_marker_y + y_span * 0.03,
                                err_txt, ha="center", va="bottom",
                                fontsize=7.5, color=PREDTIME_COLOR, fontweight="bold",
                                zorder=8,
                            )

            # Small amber tick on the bar when human label disagrees with prediction
            if differs:
                self.ax.plot([(s + e) / 2], [marker_y - y_span * 0.025],
                             marker="^", markersize=5, color=DIFF_INDICATOR,
                             markeredgecolor="white", markeredgewidth=0.5, zorder=7)

            # Decide whether to draw the GT/PRED text badges for this segment
            draw_badges = is_sel or show_all_preds
            if not draw_badges:
                continue

            # Truncate so we don't blow up the layout on long predictions
            def _trunc(t, n=70):
                t = (t or "").strip()
                return (t[:n - 1] + "…") if len(t) > n else t

            # GT (human label) — green badge above the bar
            self.ax.text(
                (s + e) / 2, gt_label_y,
                f"GT  {_trunc(label) or '∅'}",
                ha="center", va="bottom",
                fontsize=10 if is_sel else 8,
                color="#fff7ed", fontweight="bold",
                bbox=dict(facecolor=GT_COLOR,
                          alpha=0.90 if is_sel else 0.75,
                          edgecolor="#bbf7d0" if is_sel else "none",
                          boxstyle="round,pad=0.35", linewidth=1.2),
                zorder=8, **text_kw,
            )

            # PRED (model prediction) — blue badge below the bar.
            if pred:
                pred_text = (
                    f"PRED  {_trunc(pred)}"
                    if differs else "PRED  ✓ matches GT"
                )
                # For partial paragraphs, append the spoken progress
                frac = seg.get("_para_fraction")
                if frac is not None and frac < 0.99:
                    spoken = seg.get("_para_spoken_words", 0)
                    total_w = seg.get("_para_total_words", 0)
                    pred_text += f"   ⚠ partial: {spoken}/{total_w} words ({frac*100:.0f}%)"
                self.ax.text(
                    (s + e) / 2, pred_label_y,
                    pred_text,
                    ha="center", va="bottom",
                    fontsize=9 if is_sel else 7,
                    color="#fff7ed",
                    fontweight="bold" if differs else "normal",
                    bbox=dict(facecolor=PRED_COLOR,
                              alpha=0.90 if is_sel else 0.70,
                              edgecolor="#bfdbfe" if is_sel else "none",
                              boxstyle="round,pad=0.30", linewidth=1.0),
                    zorder=8, **text_kw,
                )

        # Current selection region (Audacity-style yellow translucent box)
        if self.region_start is not None and self.region_end is not None \
                and self.region_end > self.region_start:
            s, e = self.region_start, self.region_end
            self.ax.axvspan(s, e, alpha=0.22, color=P["region"], zorder=5)
            self.ax.axvline(s, color=P["region_edge"], lw=1.2,
                            linestyle="-", zorder=7)
            self.ax.axvline(e, color=P["region_edge"], lw=1.2,
                            linestyle="-", zorder=7)
            # Time labels at top — HH:MM:SS.fff format like Audacity
            label_top = y_top - y_span * 0.02
            self.ax.text(
                s, label_top, format_hms(s), ha="left", va="top",
                fontsize=8, color=P["region_edge"], fontweight="bold",
                bbox=dict(facecolor=P["label_bg"], alpha=0.7,
                          edgecolor="none", boxstyle="round,pad=0.2"),
                zorder=8,
            )
            self.ax.text(
                e, label_top, format_hms(e), ha="right", va="top",
                fontsize=8, color=P["region_edge"], fontweight="bold",
                bbox=dict(facecolor=P["label_bg"], alpha=0.7,
                          edgecolor="none", boxstyle="round,pad=0.2"),
                zorder=8,
            )

        # Live playback cursor (moves during sd.play)
        if self._playback_x is not None and xlim[0] <= self._playback_x <= xlim[1]:
            self.ax.axvline(self._playback_x, color="#fb7185",
                            lw=1.6, linestyle="-", zorder=9)
            self.ax.text(
                self._playback_x, y_bot + y_span * 0.93, "▶",
                ha="center", va="bottom", fontsize=10, color="#fb7185",
                fontweight="bold", zorder=10,
            )

        # Title — also tells you which view you're looking at
        title = Path(self.audio_path).name if self.audio_path else "Waveform"
        view_disp = {
            "waveform":    "linear waveform",
            "wave_db":     "waveform (dB)",
            "spec_linear": "linear spectrogram",
            "spec_logdb":  "log-dB spectrogram",
            "spec_mel":    "mel spectrogram",
        }.get(self.view_mode, self.view_mode)
        title += f"  │  {view_disp}"
        if self.segments:
            title += f"  │  {len(self.segments)} segments"
        self.ax.set_title(title, fontsize=10, color=P["title"])

        self.canvas.draw_idle()
        self._update_view_label()
        self._update_scrollbar()

    # ----- Canvas mouse handlers (Audacity-style) -----
    def _pixel_tolerance(self, n_pixels=8):
        """Convert n_pixels into a data-x tolerance, using the current
        axes width. With this, the edge-grab zone stays roughly the same
        size on screen no matter how zoomed in we are."""
        try:
            bbox = self.ax.get_window_extent()
            w_px = max(1.0, float(bbox.width))
        except Exception:
            w_px = max(1.0, float(self.canvas.width()))
        x0, x1 = self.ax.get_xlim()
        return n_pixels * (x1 - x0) / w_px

    def _classify_press(self, x, button):
        """Decide the drag-mode given the press location and mouse button.
        Returns one of: 'resize_start', 'resize_end', 'move', 'new'."""
        if button == 1 and self.region_start is not None and self.region_end is not None \
                and self.region_end > self.region_start:
            tol = self._pixel_tolerance(8)
            if abs(x - self.region_start) <= tol:
                return "resize_start"
            if abs(x - self.region_end) <= tol:
                return "resize_end"
            if self.region_start < x < self.region_end:
                return "move"
        return "new"

    def _on_canvas_press(self, ev):
        if ev.inaxes != self.ax or self.y is None or ev.xdata is None:
            return
        if self.toolbar.mode:  # matplotlib pan/zoom mode active
            return
        self._press_xdata = float(ev.xdata)
        self._press_button = ev.button
        self._drag_moved = False
        self._drag_anchor = float(ev.xdata)
        self._initial_region = (self.region_start, self.region_end)
        self._drag_mode = self._classify_press(self._press_xdata, ev.button)
        if self._drag_mode == "move":
            self._move_offset = float(ev.xdata)

    def _on_canvas_motion(self, ev):
        # 1. Dragging?
        if self._press_xdata is not None:
            if ev.inaxes != self.ax or ev.xdata is None or self.toolbar.mode:
                return
            if self._press_button != 1:
                return
            x = float(ev.xdata)
            # Threshold so a tiny jitter doesn't turn a click into a drag
            if not self._drag_moved and abs(x - self._press_xdata) < self._pixel_tolerance(3):
                return
            self._drag_moved = True

            if self._drag_mode == "new":
                a, b = sorted([self._drag_anchor, x])
                self._set_region(a, b, redraw=True)
            elif self._drag_mode == "resize_start":
                end = self._initial_region[1] if self._initial_region else x
                self._set_region(x, end, redraw=True)
            elif self._drag_mode == "resize_end":
                start = self._initial_region[0] if self._initial_region else x
                self._set_region(start, x, redraw=True)
            elif self._drag_mode == "move" and self._initial_region:
                delta = x - self._move_offset
                ns = self._initial_region[0] + delta
                ne = self._initial_region[1] + delta
                # Clamp to audio bounds, preserving the region's width
                dur = len(self.y) / self.sr
                width = self._initial_region[1] - self._initial_region[0]
                if ns < 0.0:
                    ns = 0.0
                    ne = ns + width
                if ne > dur:
                    ne = dur
                    ns = ne - width
                self._set_region(ns, ne, redraw=True)
            return

        # 2. Just hovering — change cursor to hint at what a press will do
        if ev.inaxes != self.ax or ev.xdata is None or self.y is None:
            self.canvas.setCursor(Qt.ArrowCursor)
            return
        if self.toolbar.mode:
            return  # leave cursor alone in pan/zoom modes
        if self.region_start is not None and self.region_end is not None \
                and self.region_end > self.region_start:
            tol = self._pixel_tolerance(8)
            x = float(ev.xdata)
            if abs(x - self.region_start) <= tol or abs(x - self.region_end) <= tol:
                self.canvas.setCursor(Qt.SizeHorCursor)
                return
            if self.region_start < x < self.region_end:
                self.canvas.setCursor(Qt.OpenHandCursor)
                return
        self.canvas.setCursor(Qt.IBeamCursor)  # text-cursor look = "place a marker"

    def _on_canvas_release(self, ev):
        if self._press_xdata is None:
            return
        if self.toolbar.mode:
            self._reset_press_state()
            return
        if ev.inaxes != self.ax or ev.xdata is None:
            self._reset_press_state()
            return

        x = float(ev.xdata)
        if self.y is not None:
            x = max(0.0, min(x, len(self.y) / self.sr))

        if not self._drag_moved:
            # Click without drag
            shift = bool(ev.guiEvent and (int(ev.guiEvent.modifiers()) & int(Qt.ShiftModifier)))
            if ev.button == 1:
                if shift and self.region_start is not None and self.region_end is not None:
                    # Shift+click extends the nearer edge to x — classic Audacity
                    if abs(x - self.region_start) < abs(x - self.region_end):
                        self._set_region(x, self.region_end, redraw=True)
                    else:
                        self._set_region(self.region_start, x, redraw=True)
                elif self._drag_mode in ("resize_start", "resize_end", "move"):
                    # Pressed inside an existing region without dragging —
                    # treat as "place playback cursor" without disturbing
                    # the selection.
                    self._playback_x = x
                    self._redraw(xlim=self.ax.get_xlim())
                else:
                    # Plain left-click in empty space — set start
                    end = self.region_end if (self.region_end is not None
                                              and self.region_end > x) else x
                    self._set_region(x, end, redraw=True)
            elif ev.button == 3:
                # Right-click anywhere — set end
                start = self.region_start if (self.region_start is not None
                                              and self.region_start < x) else x
                self._set_region(start, x, redraw=True)

        self._reset_press_state()

    def _reset_press_state(self):
        self._press_xdata = None
        self._press_button = None
        self._drag_moved = False
        self._drag_mode = None
        self._initial_region = None
        self._move_offset = None


    def _on_canvas_scroll(self, ev):
        if self.y is None or ev.inaxes != self.ax:
            return
        x0, x1 = self.ax.get_xlim()
        mx = ev.xdata if ev.xdata is not None else (x0 + x1) / 2
        f = 0.7 if ev.button == "up" else 1.42
        new_x0 = max(0.0, mx - (mx - x0) * f)
        new_x1 = min(len(self.y) / self.sr, mx + (x1 - mx) * f)
        if new_x1 - new_x0 < 0.02:
            return
        self._redraw(xlim=(new_x0, new_x1))

    # ----- Zoom / view controls -----
    def _zoom_by(self, factor):
        """factor < 1 zooms in (shrinks visible range), factor > 1 zooms out."""
        if self.y is None:
            return
        x0, x1 = self.ax.get_xlim()
        center = (x0 + x1) / 2.0
        half_width = (x1 - x0) * factor / 2.0
        max_dur = len(self.y) / self.sr
        new_x0 = max(0.0, center - half_width)
        new_x1 = min(max_dur, center + half_width)
        # Cap minimum zoom-in at 20 ms so we don't zoom into nothingness
        if new_x1 - new_x0 < 0.02:
            return
        # Cap maximum zoom-out at full audio
        if (new_x1 - new_x0) >= max_dur:
            new_x0, new_x1 = 0.0, max_dur
        self._redraw(xlim=(new_x0, new_x1))

    def on_zoom_in(self):
        if self.y is None:
            self._set_status("Load an audio file first")
            return
        self._zoom_by(0.5)

    def on_zoom_out(self):
        if self.y is None:
            self._set_status("Load an audio file first")
            return
        self._zoom_by(2.0)

    def on_fit_view(self):
        """Fit view by priority:
          1. The currently-selected (blue) region
          2. The currently-selected segment (clicked in the Segments list)
          3. The bounding box of all added segments
          4. The full audio
        """
        if self.y is None:
            self._set_status("Load an audio file first")
            return
        max_dur = len(self.y) / self.sr

        # 1. Selected region
        if (self.region_start is not None and self.region_end is not None
                and self.region_end > self.region_start):
            pad = max(0.1, (self.region_end - self.region_start) * 0.3)
            x0 = max(0.0, self.region_start - pad)
            x1 = min(max_dur, self.region_end + pad)
            self._redraw(xlim=(x0, x1))
            self._set_status(f"Fit to selected region [{self.region_start:.3f}–{self.region_end:.3f}s]")
            return

        # 2. Selected segment in the list
        if (self._selected_seg_idx is not None
                and 0 <= self._selected_seg_idx < len(self.segments)):
            seg = self.segments[self._selected_seg_idx]
            pad = max(0.1, (seg["end"] - seg["start"]) * 0.3)
            x0 = max(0.0, seg["start"] - pad)
            x1 = min(max_dur, seg["end"] + pad)
            self._redraw(xlim=(x0, x1))
            self._set_status(f"Fit to segment {self._selected_seg_idx + 1}: {seg['label']}")
            return

        # 3. All segments
        if self.segments:
            s0 = min(s["start"] for s in self.segments)
            s1 = max(s["end"] for s in self.segments)
            pad = max(0.5, (s1 - s0) * 0.05)
            x0 = max(0.0, s0 - pad)
            x1 = min(max_dur, s1 + pad)
            self._redraw(xlim=(x0, x1))
            self._set_status(f"Fit to {len(self.segments)} segments")
            return

        # 4. Full audio
        self._redraw(xlim=(0.0, max_dur))
        self._set_status("Fit to full audio")

    def on_reset_view(self):
        if self.y is None:
            self._set_status("Load an audio file first")
            return
        self._redraw(xlim=(0.0, len(self.y) / self.sr))
        self._set_status("View reset to full audio")

    def _update_view_label(self):
        """Update the tiny label next to the zoom buttons that shows the
        current visible range (handy when you've zoomed in and want to
        confirm where you are)."""
        if self.y is None or self.ax is None:
            self.lbl_view.setText("")
            return
        x0, x1 = self.ax.get_xlim()
        total = len(self.y) / self.sr
        visible = max(0.0, x1 - x0)
        zoom_ratio = total / visible if visible > 0 else 1.0
        self.lbl_view.setText(
            f"visible: {x0:.2f}–{x1:.2f}s  ({visible:.2f}s, {zoom_ratio:.1f}× zoom)"
        )

    # ----- Region state -----
    def _set_region(self, s, e, redraw=True):
        s = max(0.0, float(s))
        e = max(0.0, float(e))
        if self.y is not None:
            dur = len(self.y) / self.sr
            s = min(s, dur)
            e = min(e, dur)
        if e < s:
            s, e = e, s
        self.region_start = s
        self.region_end = e
        # Sync spinboxes without triggering callback storm
        self._suppress_spin = True
        self.sp_start.setValue(s)
        self.sp_end.setValue(e)
        self._suppress_spin = False
        self.lbl_dur.setText(
            f"{format_hms(s)} → {format_hms(e)}  ({(e - s):.3f} s)"
        )
        if redraw:
            self._redraw(xlim=self.ax.get_xlim())

    def _on_spin_changed(self, _v=None):
        if self._suppress_spin:
            return
        s = self.sp_start.value()
        e = self.sp_end.value()
        self._set_region(s, e, redraw=True)

    def on_clear_region(self):
        self.region_start = None
        self.region_end = None
        self._suppress_spin = True
        self.sp_start.setValue(0.0)
        self.sp_end.setValue(0.0)
        self._suppress_spin = False
        self.lbl_dur.setText("00:00:00.000 → 00:00:00.000  (0.000 s)")
        self._redraw(xlim=self.ax.get_xlim() if self.y is not None else None)

    # ----- Playback (with moving cursor) -----
    def _play_at_speed(self, audio_segment, start_x, end_x):
        """Play a slice of audio at self.playback_speed. We do this by
        feeding sounddevice an inflated sample rate — quick and reliable.
        The pitch shifts (no time-stretching) but for review that's the
        accepted tradeoff."""
        speed = float(self.playback_speed) if self.playback_speed > 0 else 1.0
        sd.stop()
        sd.play(audio_segment, int(round(self.sr * speed)))
        # Cursor advances by `speed` seconds per real second so it stays
        # synced with what you actually hear.
        self._start_playback_cursor(start_x, end_x, speed=speed)

    def on_play_region(self):
        if self.y is None:
            return
        if self.region_start is None or self.region_end is None \
                or self.region_end <= self.region_start:
            self._set_status("No region selected to play")
            return
        s = int(self.region_start * self.sr)
        e = int(self.region_end * self.sr)
        self._play_at_speed(self.y[s:e], self.region_start, self.region_end)

    def on_play_all(self):
        if self.y is None:
            return
        self._play_at_speed(self.y, 0.0, len(self.y) / self.sr)

    def on_toggle_play_region(self):
        """Spacebar: play the region, or stop if already playing."""
        if self._playback_timer is not None and self._playback_timer.isActive():
            self.on_stop_playback()
        else:
            self.on_play_region()

    def on_stop_playback(self):
        sd.stop()
        self._stop_playback_cursor()

    def _start_playback_cursor(self, start_x, end_x, speed=1.0):
        self._playback_start_time = time.monotonic()
        self._playback_start_x = float(start_x)
        self._playback_end_x = float(end_x)
        self._playback_x = float(start_x)
        self._playback_speed = float(speed)
        if self._playback_timer is None:
            self._playback_timer = QTimer(self)
            self._playback_timer.timeout.connect(self._tick_playback_cursor)
        self._playback_timer.start(33)  # ~30 fps

    def _tick_playback_cursor(self):
        if self._playback_start_time is None:
            self._stop_playback_cursor()
            return
        elapsed = time.monotonic() - self._playback_start_time
        # Cursor advances by `speed` seconds of source audio per second of wall
        # clock — matches what sounddevice is actually playing back.
        cur = self._playback_start_x + elapsed * getattr(self, "_playback_speed", 1.0)
        if cur >= self._playback_end_x or self.y is None:
            self._stop_playback_cursor()
            return
        self._playback_x = cur
        # Cheap redraw: keep current xlim, only the cursor moves
        self._redraw(xlim=self.ax.get_xlim())

    def _stop_playback_cursor(self):
        if self._playback_timer is not None:
            self._playback_timer.stop()
        self._playback_x = None
        self._playback_start_time = None
        if self.y is not None:
            self._redraw(xlim=self.ax.get_xlim())

    # ----- Scrollbar handling -----
    def _update_scrollbar(self):
        """Sync the horizontal scrollbar to the current view's xlim."""
        if self.y is None or self.sr is None:
            self._suppress_scrollbar = True
            self.scrollbar.setRange(0, 0)
            self._suppress_scrollbar = False
            return
        total_ms = max(1, int(round(len(self.y) / self.sr * 1000)))
        x0, x1 = self.ax.get_xlim()
        visible_ms = max(1, int(round((x1 - x0) * 1000)))
        x0_ms = max(0, int(round(x0 * 1000)))
        self._suppress_scrollbar = True
        # Slider range = [0, total - visible] so handle position = x0
        self.scrollbar.setRange(0, max(0, total_ms - visible_ms))
        self.scrollbar.setPageStep(visible_ms)
        # Small step = 5% of visible (~50 ms at default zoom)
        self.scrollbar.setSingleStep(max(1, visible_ms // 20))
        self.scrollbar.setValue(min(self.scrollbar.maximum(), x0_ms))
        self._suppress_scrollbar = False

    def _on_scrollbar_changed(self, value_ms):
        if self._suppress_scrollbar or self.y is None:
            return
        x0_cur, x1_cur = self.ax.get_xlim()
        width = x1_cur - x0_cur
        dur = len(self.y) / self.sr
        new_x0 = max(0.0, min(dur - width, value_ms / 1000.0))
        new_x1 = new_x0 + width
        self._redraw(xlim=(new_x0, new_x1))


    # ----- Undo / redo -----
    def _push_undo(self):
        """Snapshot current segment state BEFORE a mutating action."""
        import copy
        snap = (copy.deepcopy(self.segments), self._selected_seg_idx)
        self._undo_stack.append(snap)
        if len(self._undo_stack) > self._undo_limit:
            self._undo_stack.pop(0)
        # Any new action invalidates the redo history
        self._redo_stack.clear()

    def on_undo(self):
        if not self._undo_stack:
            self._set_status("Nothing to undo")
            return
        import copy
        # Save current state to redo, restore previous
        self._redo_stack.append(
            (copy.deepcopy(self.segments), self._selected_seg_idx)
        )
        segs, sel = self._undo_stack.pop()
        self.segments = segs
        self._selected_seg_idx = sel
        self._refresh_seg_list()
        self._redraw(xlim=self.ax.get_xlim() if self.y is not None else None)
        self._set_status(f"Undo  ({len(self._undo_stack)} more available)")

    def on_redo(self):
        if not self._redo_stack:
            self._set_status("Nothing to redo")
            return
        import copy
        self._undo_stack.append(
            (copy.deepcopy(self.segments), self._selected_seg_idx)
        )
        segs, sel = self._redo_stack.pop()
        self.segments = segs
        self._selected_seg_idx = sel
        self._refresh_seg_list()
        self._redraw(xlim=self.ax.get_xlim() if self.y is not None else None)
        self._set_status(f"Redo  ({len(self._redo_stack)} more available)")

    # ----- Segment list -----
    def on_add_segment(self):
        if self.y is None:
            QMessageBox.warning(self, "", "Load an audio file first.")
            return
        if self.region_start is None or self.region_end is None or self.region_end <= self.region_start:
            QMessageBox.warning(self, "", "Mark a region first (drag, click, or type Start/End).")
            return
        label = self.txt_current_label.text().strip()
        if not label:
            QMessageBox.warning(self, "", "Pick or type a label first.")
            return
        seg = {"label": label,
               "start": float(self.region_start),
               "end": float(self.region_end),
               "prediction": ""}  # manual segments have no model prediction
        self._push_undo()
        self.segments.append(seg)
        self._refresh_seg_list()
        self._set_status(f"Added: {label} [{seg['start']:.3f}–{seg['end']:.3f}s]  (total: {len(self.segments)})")
        self._redraw(xlim=self.ax.get_xlim())

    def _refresh_seg_list(self):
        self.seg_list.clear()
        for i, seg in enumerate(self.segments):
            label = (seg.get("label") or "").strip()
            pred = (seg.get("prediction") or "").strip()
            # Marker: ✓ = label & prediction agree (verified)
            #         ✎ = human edited a non-matching prediction (worth a glance)
            #         · = no prediction available (manual segment or VAD-only)
            if not pred:
                marker = "·"
            elif label == pred:
                marker = "✓"
            else:
                marker = "✎"
            # Type tag (char / word / s1.. / p1.. / sent / para)
            tag = short_type_tag(seg)
            # Truncate the label cell so the times still line up
            label_disp = label if len(label) <= 20 else label[:19] + "…"
            # Timing error suffix if predicted
            err_suffix = ""
            se = seg.get("start_err"); ee = seg.get("end_err")
            if se is not None and ee is not None:
                err_suffix = f"  Δ{se:+.2f}/{ee:+.2f}s"
            txt = (f"{i+1:03d} {marker} [{tag:<7}] {label_disp:<20}  "
                   f"{format_hms(seg['start'])} → {format_hms(seg['end'])}"
                   f"{err_suffix}")
            item = QListWidgetItem(txt)
            # Partial paragraph → orange override
            frac = seg.get("_para_fraction")
            if frac is not None and frac < 0.95:
                item.setForeground(QColor("#ea580c"))
                item.setToolTip(
                    f"Partial paragraph: spoke "
                    f"{seg.get('_para_spoken_words', 0)}/"
                    f"{seg.get('_para_total_words', 0)} words ({frac*100:.0f}%)\n\n"
                    f"Last matched word: {seg.get('_para_last_word', '')}"
                )
            elif marker == "✎":
                item.setForeground(QColor("#b45309"))  # amber: prediction differs
                tt = (f"GT (your label):  {label}\n"
                      f"PRED (model):     {pred}\n")
                if se is not None:
                    tt += (f"\nPredicted timing: {format_hms(seg.get('pred_start',0))} "
                           f"→ {format_hms(seg.get('pred_end',0))}\n"
                           f"Error: Δstart {se:+.3f}s, Δend {ee:+.3f}s"
                           + ("  (estimated)" if seg.get("pred_estimated") else ""))
                item.setToolTip(tt)
            elif marker == "✓":
                item.setForeground(QColor("#15803d"))  # green: agree
                item.setToolTip(
                    f"GT == PRED:  {label}\n\n"
                    "Marker '✓' = label matches the model's prediction exactly."
                )
            else:
                item.setToolTip(f"Label: {label}  ({tag})\n\nNo prediction yet.")
            self.seg_list.addItem(item)
        if self._selected_seg_idx is not None and 0 <= self._selected_seg_idx < len(self.segments):
            self.seg_list.setCurrentRow(self._selected_seg_idx)

    def _on_seg_clicked(self, item):
        idx = self.seg_list.row(item)
        if 0 <= idx < len(self.segments):
            # Timing: close out the previously-selected segment's timer and
            # start one for this segment. Identity = label+start so it's stable
            # across reorders. Used by the Analytics tab for the paper study.
            self._timing_flush()
            seg = self.segments[idx]
            self._timing_active_uid = self._seg_uid(seg)
            self._timing_active_since = time.monotonic()

            self._selected_seg_idx = idx
            # Zoom waveform to show this segment with some padding
            pad = max(0.5, (seg["end"] - seg["start"]) * 0.5)
            x0 = max(0.0, seg["start"] - pad)
            x1 = min(len(self.y) / self.sr if self.y is not None else seg["end"] + pad, seg["end"] + pad)
            self._redraw(xlim=(x0, x1))
            # Update the Diff tab so the right pane reflects the new selection
            self._update_diff_panel()

    @staticmethod
    def _seg_uid(seg):
        """Stable-ish identity for timing accounting."""
        return f"{seg.get('label','')}|{seg.get('start',0):.3f}"

    def _timing_flush(self):
        """Add elapsed time on the currently-active segment to its tally."""
        if self._timing_active_uid is not None and self._timing_active_since is not None:
            dt = time.monotonic() - self._timing_active_since
            if 0 < dt < 3600:  # ignore absurd gaps (left it open overnight)
                self._timing_log[self._timing_active_uid] = (
                    self._timing_log.get(self._timing_active_uid, 0.0) + dt
                )
        self._timing_active_since = time.monotonic()

    def on_delete_segment(self):
        idx = self.seg_list.currentRow()
        if idx < 0 or idx >= len(self.segments):
            return
        self._push_undo()
        seg = self.segments.pop(idx)
        self._selected_seg_idx = None
        self._refresh_seg_list()
        self._set_status(f"Deleted: {seg['label']} [{seg['start']:.3f}–{seg['end']:.3f}s]  (Ctrl+Z to undo)")
        self._redraw(xlim=self.ax.get_xlim() if self.y is not None else None)

    def on_relabel_segment(self):
        idx = self.seg_list.currentRow()
        if idx < 0 or idx >= len(self.segments):
            return
        new_label = self.txt_current_label.text().strip()
        if not new_label:
            QMessageBox.warning(self, "", "Type or pick a label first (top right).")
            return
        self._push_undo()
        old = self.segments[idx]["label"]
        self.segments[idx]["label"] = new_label
        self._refresh_seg_list()
        self._set_status(f"Relabeled: '{old}' → '{new_label}'")
        self._redraw(xlim=self.ax.get_xlim())
        self._update_diff_panel()

    def on_retime_segment(self):
        idx = self.seg_list.currentRow()
        if idx < 0 or idx >= len(self.segments):
            return
        if self.region_start is None or self.region_end is None or self.region_end <= self.region_start:
            QMessageBox.warning(self, "", "Mark a new region first.")
            return
        self._push_undo()
        self.segments[idx]["start"] = float(self.region_start)
        self.segments[idx]["end"] = float(self.region_end)
        self._refresh_seg_list()
        self._set_status(f"Updated times for segment {idx+1}")
        self._redraw(xlim=self.ax.get_xlim())

    def on_load_seg_to_region(self):
        idx = self.seg_list.currentRow()
        if idx < 0 or idx >= len(self.segments):
            return
        seg = self.segments[idx]
        self._set_current_label(seg["label"])
        self._set_region(seg["start"], seg["end"], redraw=True)

    # ----- Save/Load CSV -----
    def on_save_csv(self):
        if not self.segments:
            QMessageBox.warning(self, "", "No segments to save yet.")
            return
        # Default to xlsx if openpyxl is available — it's the user's
        # preferred format. Falls back to csv if not.
        default_ext = "xlsx" if _XLSX else "csv"
        default = "annotations." + default_ext
        if self.audio_path:
            default = Path(self.audio_path).stem + "." + default_ext
        default_path = str(OUTPUT_DIR / default)

        # Multi-format dialog filter: dialog will use the picked filter's
        # extension if the user types a bare filename without extension.
        if _XLSX:
            filt = ("Excel (*.xlsx);;CSV (*.csv);;TSV (*.tsv);;"
                    "All annotation files (*.xlsx *.csv *.tsv)")
        else:
            filt = "CSV (*.csv);;TSV (*.tsv)"
        path, sel = QFileDialog.getSaveFileName(
            self, "Save annotations", default_path, filt,
        )
        if not path:
            return
        # Append extension if the user forgot one (matches selected filter)
        if Path(path).suffix == "":
            if "xlsx" in sel.lower():
                path += ".xlsx"
            elif "tsv" in sel.lower():
                path += ".tsv"
            else:
                path += ".csv"

        # Only write the 4th `prediction` column if at least one segment has one.
        # This keeps simple manual-only files identical to the 3-column format
        # app.py expects, while still preserving model output when it exists.
        has_any_pred = any((seg.get("prediction") or "") for seg in self.segments)

        # Build rows
        rows = []
        for seg in self.segments:
            row = [seg["label"],
                   format_hms(seg["start"]),
                   format_hms(seg["end"])]
            if has_any_pred:
                row.append(seg.get("prediction") or "")
            rows.append(row)

        # For xlsx we add a header row (Excel users expect one); for csv/tsv
        # we keep it headerless so app.py can `header=None` it as-is.
        is_xlsx = Path(path).suffix.lower() in (".xlsx", ".xlsm")
        header = None
        if is_xlsx:
            header = ["label", "start", "end"] + (["prediction"] if has_any_pred else [])

        try:
            _write_table(path, rows, header=header)
            n_diff = sum(
                1 for seg in self.segments
                if (seg.get("prediction") or "").strip()
                and (seg.get("prediction") or "").strip() != (seg.get("label") or "").strip()
            )
            extra = f"  │  {n_diff} edited from PRED" if n_diff else ""
            self._set_status(f"✓ Saved {len(self.segments)} segments → {path}{extra}")
            # Clear the autosave — user has a real saved file now
            self._clear_autosave_for_current()
            # And refresh the project list (the status icon may have changed)
            if self.project_dir is not None:
                self._refresh_project_list()
        except Exception as e:
            traceback.print_exc()
            QMessageBox.critical(self, "Save error", str(e))

    def on_load_csv(self):
        # Multi-format open dialog — covers everything app.py accepts plus xlsx
        if _XLSX:
            filt = ("Annotations (*.csv *.tsv *.xlsx *.xls);;"
                    "Excel (*.xlsx *.xls);;CSV (*.csv);;TSV (*.tsv);;"
                    "All Files (*)")
        else:
            filt = "CSV/TSV (*.csv *.tsv);;CSV (*.csv);;TSV (*.tsv);;All Files (*)"
        path, _ = QFileDialog.getOpenFileName(
            self, "Load existing annotations", str(OUTPUT_DIR), filt,
        )
        if not path:
            return
        try:
            rows = _read_table(path)
        except Exception as e:
            traceback.print_exc()
            QMessageBox.critical(self, "Load error", str(e))
            return

        loaded, skipped = [], 0
        # Detect & skip an obvious header row. Accept several common column
        # names for the first cell — your existing Female.xlsx uses "Content",
        # other formats use "label", "word", "text", "transcript", etc.
        if rows:
            first = [str(c).strip().lower() for c in rows[0]]
            label_keywords = ("label", "content", "word", "text", "transcript",
                              "utterance", "phrase")
            if (len(first) >= 3 and
                any(k in first[0] for k in label_keywords) and
                ("start" in first[1] or "begin" in first[1]) and
                "end" in first[2]):
                rows = rows[1:]
                log(f"Skipped header row: {first}")

        for row in rows:
            if len(row) < 3:
                skipped += 1
                continue
            label = str(row[0]).strip()
            # parse_hms accepts HH:MM:SS.fff, MM:SS.fff, and plain seconds
            s = parse_hms(row[1])
            e = parse_hms(row[2])
            if not label or s is None or e is None or e <= s:
                skipped += 1
                continue
            # 4th column = original model prediction (optional)
            prediction = str(row[3]).strip() if len(row) >= 4 else ""
            loaded.append({
                "label": label, "start": s, "end": e,
                "prediction": prediction,
            })
        if not loaded:
            QMessageBox.warning(
                self, "",
                f"No valid rows found in {Path(path).name}.\n\n"
                "Expected format (no header for CSV; header optional for XLSX):\n"
                "    label, start, end, [prediction]\n\n"
                "Times can be HH:MM:SS.fff, MM:SS, or plain seconds."
            )
            return
        self.segments = loaded
        self._selected_seg_idx = None
        self._refresh_seg_list()
        n_with_pred = sum(1 for s in loaded if s.get("prediction"))
        msg = f"Loaded {len(loaded)} segments from {Path(path).name}"
        if n_with_pred:
            msg += f"  │  {n_with_pred} include PRED column"
        if skipped:
            msg += f" ({skipped} skipped)"
        self._set_status(msg)
        self._redraw(xlim=self.ax.get_xlim() if self.y is not None else None)

    # ----- Multi-file project mode -----
    def on_open_project_folder(self):
        folder = QFileDialog.getExistingDirectory(
            self, "Pick a folder of audio files",
            str(self.project_dir) if self.project_dir else str(APP_DIR),
        )
        if not folder:
            return
        self.project_dir = Path(folder)
        exts = {".wav", ".mp3", ".flac", ".ogg", ".m4a"}
        self.project_files = sorted(
            [p for p in self.project_dir.iterdir()
             if p.is_file() and p.suffix.lower() in exts],
            key=lambda p: p.name.lower(),
        )
        self.lbl_project.setText(
            f"📁 {self.project_dir}<br><b>{len(self.project_files)}</b> audio files"
        )
        self.lbl_project.setTextFormat(Qt.RichText)
        self.btn_next_unfinished.setEnabled(len(self.project_files) > 0)
        self._refresh_project_list()
        if len(self.project_files) == 0:
            QMessageBox.information(
                self, "Empty folder",
                "No supported audio files in that folder.\n"
                "Supported: .wav .mp3 .flac .ogg .m4a"
            )

    def _file_status(self, p):
        """Return ('icon', 'description') for a project file.
        Priorities: currently loaded > saved CSV exists > autosave exists > untouched."""
        if self.audio_path and Path(self.audio_path).name == p.name:
            return ("📂", "loaded")
        stem = p.stem
        for ext in (".xlsx", ".csv", ".tsv"):
            if (OUTPUT_DIR / f"{stem}{ext}").exists():
                return ("✓", "saved")
        autosave = AUTOSAVE_DIR / f"{p.name}.autosave.csv"
        if autosave.exists():
            return ("✎", "in-progress")
        return ("·", "untouched")

    def _refresh_project_list(self):
        self.proj_list.clear()
        for p in self.project_files:
            icon, _ = self._file_status(p)
            item = QListWidgetItem(f"{icon}  {p.name}")
            if icon == "✓":
                item.setForeground(QColor("#15803d"))
            elif icon == "✎":
                item.setForeground(QColor("#b45309"))
            elif icon == "📂":
                item.setForeground(QColor("#1d4ed8"))
            self.proj_list.addItem(item)

    def _on_project_file_dblclick(self, item):
        idx = self.proj_list.row(item)
        if not (0 <= idx < len(self.project_files)):
            return
        target = self.project_files[idx]
        # Warn if we'd lose unsaved work
        if self.segments and self.audio_path:
            ret = QMessageBox.question(
                self, "Switch file?",
                f"Switch to:\n  {target.name}\n\n"
                f"Your current segments are autosaved every 30s but you may "
                f"have unsaved changes. Continue?",
                QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes,
            )
            if ret != QMessageBox.Yes:
                return
            # Force an autosave right now so nothing is lost
            self._autosave_tick()
        self._load_audio_from_path(str(target))

    def _load_audio_from_path(self, path):
        """Programmatic version of on_load_audio — used by project mode."""
        self._set_status(f"Loading {Path(path).name}…")
        self.btn_load_audio.setEnabled(False)
        self.segments = []
        self._selected_seg_idx = None
        self._refresh_seg_list()
        t = AudioLoadThread(path)
        t.done.connect(self._on_audio_loaded)
        t.error.connect(self._on_audio_err)
        self._threads["audio"] = t
        t.start()

    def on_jump_next_unfinished(self):
        if not self.project_files:
            return
        # Find the next file (after the currently-loaded one) with no saved CSV
        cur_name = Path(self.audio_path).name if self.audio_path else None
        start_idx = 0
        if cur_name:
            for i, p in enumerate(self.project_files):
                if p.name == cur_name:
                    start_idx = i + 1
                    break
        n = len(self.project_files)
        for offset in range(n):
            cand = self.project_files[(start_idx + offset) % n]
            icon, _ = self._file_status(cand)
            if icon in ("·", "✎"):  # untouched or in-progress
                self._load_audio_from_path(str(cand))
                return
        QMessageBox.information(self, "", "All files in this folder have been saved.")

    # ----- GT vs PRED inline diff -----
    def _update_diff_panel(self):
        """Refresh the Diff tab with a coloured char-level diff for the selected segment."""
        if self.txt_diff is None:
            return
        if (self._selected_seg_idx is None or
                not (0 <= self._selected_seg_idx < len(self.segments))):
            self.txt_diff.setText(
                "<i style='color:#888'>Select a segment to see the diff.</i>"
            )
            return
        seg = self.segments[self._selected_seg_idx]
        gt = (seg.get("label") or "").strip()
        pred = (seg.get("prediction") or "").strip()
        # No prediction case
        if not pred:
            self.txt_diff.setText(
                f"<div style='color:#15803d;font-weight:bold'>GT</div>"
                f"<div style='padding:4px 0'>{self._esc(gt) or '<i>(empty)</i>'}</div>"
                f"<div style='color:#888;margin-top:14px'>"
                f"<i>This segment has no model prediction (manual segment).</i></div>"
            )
            return
        # Build char-level diff with difflib
        import difflib
        sm = difflib.SequenceMatcher(a=pred, b=gt, autojunk=False)
        gt_parts, pred_parts = [], []
        for tag, i1, i2, j1, j2 in sm.get_opcodes():
            p_chunk = self._esc(pred[i1:i2])
            g_chunk = self._esc(gt[j1:j2])
            if tag == "equal":
                gt_parts.append(g_chunk)
                pred_parts.append(p_chunk)
            elif tag == "replace":
                pred_parts.append(
                    f"<span style='background:#fecaca;color:#991b1b;"
                    f"text-decoration:line-through'>{p_chunk}</span>"
                )
                gt_parts.append(
                    f"<span style='background:#bbf7d0;color:#166534'>{g_chunk}</span>"
                )
            elif tag == "delete":
                pred_parts.append(
                    f"<span style='background:#fecaca;color:#991b1b;"
                    f"text-decoration:line-through'>{p_chunk}</span>"
                )
            elif tag == "insert":
                gt_parts.append(
                    f"<span style='background:#bbf7d0;color:#166534'>{g_chunk}</span>"
                )
        same = (gt == pred)
        header = (
            "<div style='color:#15803d;font-weight:bold;margin-top:0'>"
            "✓ GT and PRED match exactly</div>" if same else
            "<div style='color:#b45309;font-weight:bold;margin-top:0'>"
            "✎ GT differs from PRED — see highlighted changes below</div>"
        )
        self.txt_diff.setText(
            f"{header}"
            f"<div style='color:#15803d;margin-top:14px;font-weight:bold'>GT (your label)</div>"
            f"<div style='padding:4px 0;line-height:1.5'>{''.join(gt_parts) or '<i>(empty)</i>'}</div>"
            f"<div style='color:#1d4ed8;margin-top:14px;font-weight:bold'>PRED (model)</div>"
            f"<div style='padding:4px 0;line-height:1.5'>{''.join(pred_parts) or '<i>(empty)</i>'}</div>"
        )

    @staticmethod
    def _esc(s):
        """HTML-escape so labels with < > & don't break the diff display."""
        return (str(s).replace("&", "&amp;")
                       .replace("<", "&lt;")
                       .replace(">", "&gt;"))

    # ----- GT vs PRED accuracy report (Score tab) -----
    def _update_score_panel(self):
        """Compute per-segment accuracy + aggregate stats and render an HTML
        report. Uses the same fuzzy similarity as the alignment algorithm."""
        if not self.segments:
            self.txt_score.setText(
                "<i style='color:#888'>No segments yet. Load GT, run "
                "Auto-Detect, or Align GT first.</i>"
            )
            return
        rows = []
        scored = []
        n_no_pred = 0
        for i, seg in enumerate(self.segments):
            label = (seg.get("label") or "").strip()
            pred = (seg.get("prediction") or "").strip()
            if not pred:
                n_no_pred += 1
                rows.append((i + 1, label, "—", "—", "no PRED"))
                continue
            sim = _fuzzy_score(pred, label)
            scored.append(sim)
            if sim >= 0.95:
                tag = "✓"; color = "#15803d"
            elif sim >= 0.50:
                tag = "~"; color = "#b45309"
            else:
                tag = "✗"; color = "#b91c1c"
            rows.append((i + 1, label, pred, f"{sim:.3f}", tag, color))

        if scored:
            avg = sum(scored) / len(scored)
            n_perfect = sum(1 for s in scored if s >= 0.95)
            n_partial = sum(1 for s in scored if 0.50 <= s < 0.95)
            n_bad = sum(1 for s in scored if s < 0.50)
        else:
            avg = 0.0
            n_perfect = n_partial = n_bad = 0

        # Header summary
        m = self.compute_metrics()
        html = (
            f"<div style='font-size:13px;font-weight:bold;color:#111'>"
            f"Avg accuracy: <span style='color:#0e7490'>{avg:.3f}</span>  "
            f"({len(scored)} scored, {n_no_pred} no PRED)</div>"
            f"<div style='margin:6px 0 4px'>"
            f"<span style='color:#15803d'>✓ {n_perfect} ≥ 95%</span>  ·  "
            f"<span style='color:#b45309'>~ {n_partial} 50–95%</span>  ·  "
            f"<span style='color:#b91c1c'>✗ {n_bad} &lt; 50%</span>"
            f"</div>"
            # Paper-grade metrics block
            f"<div style='margin:8px 0;padding:8px;background:#f8fafc;"
            f"border:1px solid #e2e8f0;border-radius:4px;font-size:11px'>"
            f"<b>Acceptance rate:</b> {m['acceptance_rate']*100:.1f}% "
            f"({m['accepted']} kept / {m['edited']} edited)<br>"
            f"<b>Avg WER:</b> {m['avg_wer']:.3f}<br>"
            f"<b>Timing MAE:</b> start {m['mae_start']:.3f}s · end {m['mae_end']:.3f}s<br>"
            f"<b>Time/seg:</b> "
            + (f"mean {m['mean_time_s']:.1f}s · median {m['median_time_s']:.1f}s "
               f"({m['n_timed']} timed)"
               if m['n_timed'] else "<i>not yet measured</i>")
            + f"<br><b>Types:</b> "
            + " · ".join(f"{t}:{c}" for t, c in sorted(m['per_type'].items()))
            + (f" · <span style='color:#ea580c'>partial-para:{m['partial_paragraphs']}</span>"
               if m['partial_paragraphs'] else "")
            + "</div>"
        )
        # Per-segment table
        html += (
            "<table cellspacing='0' cellpadding='4' style='border-collapse:collapse;font-size:11px'>"
            "<tr style='background:#f3f4f6;font-weight:bold'>"
            "<td>#</td><td>GT</td><td>PRED</td><td style='text-align:right'>sim</td><td></td>"
            "</tr>"
        )
        for row in rows:
            if len(row) == 5:
                idx, label, pred, sim, tag = row
                color = "#888"
            else:
                idx, label, pred, sim, tag, color = row
            html += (
                f"<tr><td style='color:#666'>{idx:03d}</td>"
                f"<td>{self._esc(label)[:60]}</td>"
                f"<td style='color:#444'>{self._esc(pred)[:60]}</td>"
                f"<td style='text-align:right;color:{color}'>{sim}</td>"
                f"<td style='color:{color};font-weight:bold'>{tag}</td></tr>"
            )
        html += "</table>"
        self.txt_score.setText(html)

    # ----- Paper-grade metrics -----
    @staticmethod
    def _word_error_rate(ref, hyp):
        """Standard WER between reference and hypothesis strings."""
        r = (ref or "").split()
        h = (hyp or "").split()
        if not r:
            return 0.0 if not h else 1.0
        # Levenshtein on word tokens
        n, m = len(r), len(h)
        dp = list(range(m + 1))
        for i in range(1, n + 1):
            prev, dp[0] = dp[0], i
            for j in range(1, m + 1):
                tmp = dp[j]
                cost = 0 if r[i - 1] == h[j - 1] else 1
                dp[j] = min(prev + cost, dp[j] + 1, dp[j - 1] + 1)
                prev = tmp
        return dp[m] / n

    def compute_metrics(self):
        """Compute the full metric bundle used by both the on-screen report
        and the exported paper report. Returns a dict."""
        iso, lang_name = self._current_lang()
        self._timing_flush()  # make sure the active segment's time is counted

        n_total = len(self.segments)
        n_with_pred = 0
        accepted = 0           # prediction kept verbatim as the label
        edited = 0             # had a prediction but label differs
        sims = []
        wers = []
        start_errs = []
        end_errs = []
        abs_start_errs = []
        abs_end_errs = []
        partial_paras = 0
        per_type = {}          # type -> count
        timing_vals = []

        for seg in self.segments:
            label = (seg.get("label") or "").strip()
            pred = (seg.get("prediction") or "").strip()
            stype = classify_segment_type(label)
            per_type[stype] = per_type.get(stype, 0) + 1
            uid = self._seg_uid(seg)
            if uid in self._timing_log:
                timing_vals.append(self._timing_log[uid])
            if pred:
                n_with_pred += 1
                sims.append(_fuzzy_score(pred, label))
                wers.append(self._word_error_rate(label, pred))
                if _norm_for_compare(pred) == _norm_for_compare(label):
                    accepted += 1
                else:
                    edited += 1
            se = seg.get("start_err"); ee = seg.get("end_err")
            if se is not None and ee is not None:
                start_errs.append(se); end_errs.append(ee)
                abs_start_errs.append(abs(se)); abs_end_errs.append(abs(ee))
            frac = seg.get("_para_fraction")
            if frac is not None and frac < 0.95:
                partial_paras += 1

        def _mean(xs):
            return sum(xs) / len(xs) if xs else 0.0

        total_dur = sum(s["end"] - s["start"] for s in self.segments)
        return {
            "language": lang_name,
            "iso": iso,
            "audio": Path(self.audio_path).name if self.audio_path else "—",
            "n_total": n_total,
            "n_with_pred": n_with_pred,
            "accepted": accepted,
            "edited": edited,
            "acceptance_rate": (accepted / n_with_pred) if n_with_pred else 0.0,
            "avg_similarity": _mean(sims),
            "avg_wer": _mean(wers),
            "mae_start": _mean(abs_start_errs),
            "mae_end": _mean(abs_end_errs),
            "mean_start_err": _mean(start_errs),
            "mean_end_err": _mean(end_errs),
            "n_timed": len(timing_vals),
            "total_time_s": sum(timing_vals),
            "median_time_s": (sorted(timing_vals)[len(timing_vals)//2]
                              if timing_vals else 0.0),
            "mean_time_s": _mean(timing_vals),
            "session_s": time.time() - self._session_start,
            "per_type": per_type,
            "partial_paragraphs": partial_paras,
            "total_labeled_dur_s": total_dur,
        }

    def _format_metrics_report(self, m):
        """Plain-text, paste-ready metrics report for the paper."""
        lines = []
        lines.append("=" * 60)
        lines.append("  SPEECH ANNOTATION — METRICS REPORT")
        lines.append("=" * 60)
        lines.append(f"Audio          : {m['audio']}")
        lines.append(f"Language       : {m['language']} ({m['iso']})")
        lines.append(f"Segments       : {m['n_total']}  "
                     f"(labeled audio: {m['total_labeled_dur_s']:.1f}s)")
        lines.append("")
        lines.append("-- Segment types --")
        for t, c in sorted(m["per_type"].items()):
            lines.append(f"   {t:<10}: {c}")
        if m["partial_paragraphs"]:
            lines.append(f"   partial paragraphs flagged: {m['partial_paragraphs']}")
        lines.append("")
        lines.append("-- ASR pre-labeling quality --")
        lines.append(f"   segments with prediction : {m['n_with_pred']}")
        lines.append(f"   acceptance rate          : {m['acceptance_rate']*100:.1f}%  "
                     f"({m['accepted']} kept verbatim, {m['edited']} edited)")
        lines.append(f"   avg char similarity      : {m['avg_similarity']:.3f}")
        lines.append(f"   avg word error rate (WER): {m['avg_wer']:.3f}")
        lines.append("")
        lines.append("-- Timestamp prediction error (pred vs GT) --")
        if m["mae_start"] or m["mae_end"]:
            lines.append(f"   MAE start : {m['mae_start']:.3f}s   "
                         f"(mean signed {m['mean_start_err']:+.3f}s)")
            lines.append(f"   MAE end   : {m['mae_end']:.3f}s   "
                         f"(mean signed {m['mean_end_err']:+.3f}s)")
        else:
            lines.append("   (run 🎯 Predict on GT to populate timing errors)")
        lines.append("")
        lines.append("-- Annotation timing (this session) --")
        lines.append(f"   segments timed   : {m['n_timed']}")
        if m["n_timed"]:
            lines.append(f"   mean time/seg    : {m['mean_time_s']:.1f}s")
            lines.append(f"   median time/seg  : {m['median_time_s']:.1f}s")
            lines.append(f"   total active time: {m['total_time_s']:.1f}s")
        lines.append(f"   session length   : {m['session_s']:.0f}s")
        lines.append("=" * 60)
        return "\n".join(lines)

    def on_export_metrics_report(self):
        if not self.segments:
            QMessageBox.warning(self, "", "No segments to report on.")
            return
        m = self.compute_metrics()
        report = self._format_metrics_report(m)
        default = "metrics_report.txt"
        if self.audio_path:
            default = Path(self.audio_path).stem + "_metrics.txt"
        path, _ = QFileDialog.getSaveFileName(
            self, "Export metrics report", str(OUTPUT_DIR / default),
            "Text (*.txt)",
        )
        if not path:
            return
        try:
            Path(path).write_text(report, encoding="utf-8")
            self._set_status(f"✓ Metrics report → {path}")
            # Also echo into the Score panel so the user sees it immediately
            self.txt_score.setText(
                "<pre style='font-size:11px'>" + self._esc(report) + "</pre>"
            )
        except Exception as e:
            traceback.print_exc()
            QMessageBox.critical(self, "Export error", str(e))

    def on_export_manifest(self):
        if not self.segments:
            QMessageBox.warning(self, "", "No segments to export.")
            return
        if not self.audio_path:
            QMessageBox.warning(self, "", "Load the audio file first.")
            return
        import json
        iso, lang_name = self._current_lang()
        default = "manifest.jsonl"
        if self.audio_path:
            default = Path(self.audio_path).stem + "_manifest.jsonl"
        path, _ = QFileDialog.getSaveFileName(
            self, "Export training manifest", str(OUTPUT_DIR / default),
            "JSON Lines (*.jsonl)",
        )
        if not path:
            return
        try:
            with open(path, "w", encoding="utf-8") as f:
                for seg in self.segments:
                    rec = {
                        "audio_filepath": str(self.audio_path),
                        "offset": round(float(seg["start"]), 3),
                        "duration": round(float(seg["end"] - seg["start"]), 3),
                        "start": round(float(seg["start"]), 3),
                        "end": round(float(seg["end"]), 3),
                        "text": seg.get("label", ""),
                        "prediction": seg.get("prediction", ""),
                        "type": classify_segment_type(seg.get("label", "")),
                        "language": iso,
                    }
                    # Include partial-paragraph info if present
                    if seg.get("_para_fraction") is not None:
                        rec["para_fraction"] = round(seg["_para_fraction"], 3)
                        rec["para_spoken_words"] = seg.get("_para_spoken_words")
                        rec["para_total_words"] = seg.get("_para_total_words")
                    f.write(json.dumps(rec, ensure_ascii=False) + "\n")
            self._set_status(
                f"✓ Manifest ({len(self.segments)} segments) → {path}"
            )
        except Exception as e:
            traceback.print_exc()
            QMessageBox.critical(self, "Export error", str(e))

    # ----- Auto-save / crash recovery -----
    def _autosave_path(self):
        """Stable path keyed off the loaded audio's full name."""
        if not self.audio_path:
            return None
        # Sanitize: replace dir separators in the audio name so a deeply-nested
        # audio path doesn't escape AUTOSAVE_DIR.
        key = Path(self.audio_path).name
        return AUTOSAVE_DIR / f"{key}.autosave.csv"

    def _autosave_tick(self):
        """Called every 30s. Silently writes the current segments to a
        sidecar file so a crash / close / power loss doesn't destroy work."""
        if not self.segments or self.audio_path is None:
            return
        # Avoid pointless re-writes when nothing changed
        sig = (len(self.segments), tuple(
            (s.get("label",""), s.get("start"), s.get("end"))
            for s in self.segments
        ))
        if sig == self._last_autosaved_count:
            return
        path = self._autosave_path()
        if path is None:
            return
        try:
            has_pred = any((s.get("prediction") or "") for s in self.segments)
            rows = []
            for seg in self.segments:
                row = [seg["label"], format_hms(seg["start"]), format_hms(seg["end"])]
                if has_pred:
                    row.append(seg.get("prediction") or "")
                rows.append(row)
            _write_table(str(path), rows, header=None)
            self._last_autosaved_count = sig
            # Brief, non-intrusive note in the status bar
            self._set_status(
                f"💾 Autosaved {len(self.segments)} segments → {path.name}  ·  "
                f"{time.strftime('%H:%M:%S')}"
            )
        except Exception as e:
            log(f"Autosave failed: {e}")

    def _check_for_autosave_on_load(self):
        """Right after loading audio, see if a newer autosave exists from
        a previous crashed session — offer to restore it."""
        path = self._autosave_path()
        if path is None or not path.exists():
            return
        try:
            mtime = path.stat().st_mtime
            age_min = (time.time() - mtime) / 60.0
        except OSError:
            return
        # Don't nag for stale autosaves — older than 30 days, just leave it alone
        if age_min > 60 * 24 * 30:
            return
        # If a manually-saved file with the same stem already exists AND is
        # newer than the autosave, skip the prompt (you presumably did Save).
        manual = OUTPUT_DIR / f"{Path(self.audio_path).stem}.xlsx"
        manual_csv = OUTPUT_DIR / f"{Path(self.audio_path).stem}.csv"
        for cand in (manual, manual_csv):
            if cand.exists() and cand.stat().st_mtime > mtime:
                return
        # Prompt
        age_str = (f"{int(age_min)} min ago" if age_min < 60
                   else f"{age_min/60:.1f} hr ago")
        ret = QMessageBox.question(
            self, "Restore autosave?",
            f"An autosave was found for this audio:\n  {path.name}\n\n"
            f"Last saved {age_str}.\n\n"
            "Restore those segments? (Choose No to start fresh — the autosave "
            "will be overwritten on the next save cycle.)",
            QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel,
            QMessageBox.Yes,
        )
        if ret == QMessageBox.Yes:
            try:
                rows = _read_table(str(path))
                loaded = []
                for row in rows:
                    if len(row) < 3:
                        continue
                    s = parse_hms(row[1]); e = parse_hms(row[2])
                    label = str(row[0]).strip()
                    if not label or s is None or e is None or e <= s:
                        continue
                    pred = str(row[3]).strip() if len(row) >= 4 else ""
                    loaded.append({"label": label, "start": s, "end": e,
                                   "prediction": pred})
                self.segments = loaded
                self._refresh_seg_list()
                self._set_status(f"✓ Restored {len(loaded)} segments from autosave")
                if self.y is not None:
                    self._redraw(xlim=self.ax.get_xlim())
            except Exception as e:
                QMessageBox.critical(self, "Restore failed", str(e))

    def _clear_autosave_for_current(self):
        """After a successful manual save, dispose of the autosave."""
        path = self._autosave_path()
        if path is not None and path.exists():
            try:
                path.unlink()
                self._last_autosaved_count = 0
            except OSError:
                pass

    def on_export_transcripts(self):
        iso, _ = self._current_lang()
        if not self.current_sentences and not self.current_paragraph:
            # No cached transcripts; offer to import from docx now
            ret = QMessageBox.question(
                self, "No transcripts loaded",
                "No sentences/paragraph in memory. Import them from a .docx now?"
            )
            if ret == QMessageBox.Yes:
                self.on_import_docx()
                if not self.current_sentences and not self.current_paragraph:
                    return
            else:
                return
        try:
            save_transcripts_file(iso, self.current_sentences, self.current_paragraph)
            p = TRANSCRIPTS_DIR / f"{iso}.txt"
            self._set_status(
                f"✓ Wrote {len(self.current_sentences)} sentences"
                f"{' + paragraph' if self.current_paragraph else ''} → {p}"
            )
        except Exception as e:
            traceback.print_exc()
            QMessageBox.critical(self, "Export error", str(e))

    # ----- Status -----
    def _set_status(self, msg):
        self.lbl_status.setText(msg)
        log(msg)


def main():
    ensure_dirs()
    print("=" * 64)
    print("  Manual Speech Annotator  +  Auto-Detect (VAD + Whisper / IndicConformer)")
    print(f"  Wordlists dir   : {WORDLISTS_DIR}")
    print(f"  Transcripts dir : {TRANSCRIPTS_DIR}")
    print(f"  Output (CSV) dir: {OUTPUT_DIR}")
    print(f"  python-docx     : {'OK' if _DOCX else 'MISSING (pip install python-docx)'}")
    print(f"  openpyxl        : {'OK (.xlsx supported)' if _XLSX else 'MISSING — .xlsx disabled (pip install openpyxl)'}")
    print(f"  pypdf           : {'OK (.pdf wordlist import supported)' if _PDF else 'MISSING — .pdf import disabled (pip install pypdf)'}")
    print(f"  soundfile       : {'OK' if _SF else 'using librosa fallback'}")
    print(f"  scipy           : {'OK' if _SCIPY else 'MISSING (spectrograms disabled)'}")
    print(f"  faster-whisper  : {'OK (English ASR)' if _FW else 'not installed'}")
    print(f"  transformers    : {'OK (Indic + English ASR)' if _HF else 'not installed'}")
    if _HF:
        try:
            print(f"  CUDA            : {torch.cuda.is_available()}")
        except Exception:
            pass
    if not _ASR_AVAILABLE:
        print("  ⚠  Install at least one ASR backend to use 🤖 Auto-Detect:")
        print("       pip install faster-whisper       (English only)")
        print("       pip install torch transformers   (Indic + English)")
    print("=" * 64)
    app = QApplication(sys.argv)
    app.setStyleSheet(
        "QPushButton{padding:5px 10px;border:1px solid #d1d5db;"
        "border-radius:4px;background:#f9fafb;font-size:12px}"
        "QPushButton:hover{background:#e5e7eb}"
        "QPushButton:disabled{color:#9ca3af;background:#f3f4f6}"
        "QComboBox,QSpinBox,QDoubleSpinBox,QLineEdit{padding:4px 6px;font-size:12px}"
    )
    w = MainWindow()
    w.show()
    sys.exit(app.exec_())


if __name__ == "__main__":
    main()
