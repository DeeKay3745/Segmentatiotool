#!/usr/bin/env python3
"""
Manual Speech Annotation Tool
================================

Build ground-truth CSV files compatible with app.py.

Workflow:
  1. Pick language -> wordlist auto-loads from wordlists/<iso>.txt
     (or from the matching Dr__Abhishek_..._Text_Material.docx the first time).
  2. Load a .wav file -> waveform appears.
  3. Mark a region. Three ways:
       a. DRAG  on the waveform (click + drag horizontally)
       b. CLICK on the waveform: LEFT-click = set start, RIGHT-click = set end
       c. TYPE  numeric seconds into the Start / End spinboxes
  4. Pick a label: select from the word list, click a quick-button
     (s1..s5, paragraph, or any char button), or type freeform.
  5. New word? Type it in the "Add new word" box and click Add — it's
     appended to wordlists/<iso>.txt and stays for next session.
  6. Click "Add Segment" -> appears in the bottom list.
  7. Click "Save CSV" -> writes a 3-column CSV (label, start, end) with
     no header. Timestamps are saved as HH:MM:SS.000.
  8. Optionally click "Export transcripts/<iso>.txt" to also produce the
     sentence/paragraph reference file your app.py uses for accuracy
     scoring on s1..s5 / paragraph labels.

Dependencies:
    pip install PyQt5 matplotlib numpy soundfile sounddevice python-docx scipy
    (librosa is an optional fallback if soundfile is missing)
"""

import os
import re
import sys
import csv
import time
import traceback
from pathlib import Path
from functools import lru_cache

import numpy as np
import sounddevice as sd

import matplotlib
matplotlib.use("QtAgg")
import matplotlib.pyplot as plt
from matplotlib import font_manager as fm
from matplotlib.patches import Rectangle
from matplotlib.backends.backend_qtagg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.backends.backend_qtagg import NavigationToolbar2QT as NavigationToolbar

from PyQt5.QtCore import Qt, QThread, pyqtSignal, QTimer
from PyQt5.QtGui import QColor, QFont
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QFileDialog, QMessageBox,
    QVBoxLayout, QHBoxLayout, QGridLayout, QLabel, QPushButton, QComboBox,
    QListWidget, QListWidgetItem, QSplitter, QGroupBox, QDoubleSpinBox,
    QLineEdit, QTabWidget, QFrame, QScrollArea, QInputDialog, QCheckBox,
    QSlider,
)

# ---------- Optional deps ----------
try:
    import soundfile as sf
    _SF = True
except ImportError:
    _SF = False

try:
    from scipy.signal import resample_poly
    from math import gcd
    _SCIPY = True
except ImportError:
    _SCIPY = False

try:
    from docx import Document
    _DOCX = True
except ImportError:
    _DOCX = False


# ---------- Constants ----------
APP_DIR = Path(__file__).parent.resolve()
WORDLISTS_DIR = APP_DIR / "wordlists"
TRANSCRIPTS_DIR = APP_DIR / "transcripts"
OUTPUT_DIR = APP_DIR / "annotations"

LANGUAGES = [
    ("Gujarati", "gu", "gujarati"),
    ("Hindi", "hi", "hindi"),
    ("Marathi", "mr", "marathi"),
    ("English", "en", "english"),
]

QUICK_LABELS = ["s1", "s2", "s3", "s4", "s5", "paragraph"]

# Multilingual section header keywords for docx parsing
SECTION_KEYWORDS = {
    "chars": [
        "alphabet", "alphabets", "vowels", "consonants",
        "स्वर", "व्यंजन", "वर्णमाला", "मूळाक्षरे", "अक्षरे",
        "મૂળાક્ષરો", "સ્વર", "વ્યંજન",
    ],
    "words": [
        "words", "word",
        "शब्द", "शब्दों",
        "શબ્દો", "શબ્દ",
    ],
    "sentences": [
        "sentences", "sentence",
        "वाक्य", "वाक्ये", "वाक्यांश",
        "વાક્યો", "વાક્ય",
    ],
    "paragraph": [
        "paragraph", "passage",
        "अनुच्छेद", "गद्यांश", "परिच्छेद",
        "ફકરો", "પેરેગ્રાફ", "ગદ્યાંશ",
    ],
}


def log(msg):
    print(f"[{time.strftime('%H:%M:%S')}] {msg}", flush=True)


def seconds_to_hhmmss_ms(sec):
    """Convert seconds -> HH:MM:SS.000 format."""
    sec = max(0.0, float(sec))
    total_ms = int(round(sec * 1000.0))
    hours = total_ms // 3_600_000
    total_ms %= 3_600_000
    minutes = total_ms // 60_000
    total_ms %= 60_000
    seconds = total_ms // 1000
    millis = total_ms % 1000
    return f"{hours:02d}:{minutes:02d}:{seconds:02d}.{millis:03d}"


def hhmmss_ms_to_seconds(value):
    """Accept either float seconds or HH:MM:SS.000 and return seconds."""
    text = str(value).strip()
    if not text:
        raise ValueError("empty timestamp")
    # Plain seconds from older CSV files remain supported.
    try:
        return float(text)
    except ValueError:
        pass
    m = re.match(r"^(?:(\d+):)?([0-5]?\d):([0-5]?\d)(?:[\.,](\d{1,3}))?$", text)
    if not m:
        raise ValueError(f"bad timestamp: {value}")
    h = int(m.group(1) or 0)
    mi = int(m.group(2))
    se = int(m.group(3))
    ms = int((m.group(4) or "0").ljust(3, "0")[:3])
    return h * 3600 + mi * 60 + se + ms / 1000.0


# ---------- Audio I/O ----------
def load_audio(path):
    """Load audio as float32 mono with native sample rate."""
    if _SF:
        try:
            data, sr = sf.read(str(path), dtype="float32", always_2d=True)
            y = data.mean(axis=1) if data.shape[1] > 1 else data[:, 0]
            return y.astype(np.float32), sr
        except Exception:
            pass
    import librosa
    y, sr = librosa.load(str(path), sr=None, mono=True)
    return y.astype(np.float32), sr


# ---------- DOCX text material parser ----------
def _extract_docx_paragraphs(path):
    """Return list of cleaned non-empty paragraph strings from a .docx.

    Word stores Shift+Enter "soft" line breaks as `\\n` inside a single
    paragraph object — but each visual line is semantically a separate
    entry for our purposes (header on one line, body on the next). We
    split on `\\n` so the section parser sees the right structure.
    Also reads text inside tables so docs that put content in a 1-cell
    table (like the English file's blank table) don't drop anything.
    """
    if not _DOCX:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")
    doc = Document(str(path))
    out = []

    def _push(raw):
        if not raw:
            return
        for sub in raw.split("\n"):
            sub = sub.strip()
            if sub:
                out.append(sub)

    for p in doc.paragraphs:
        _push(p.text)
    # Include any text trapped inside tables, in document order is hard
    # to guarantee with python-docx; for the Dr. Abhishek files the tables
    # are empty, so we just append after paragraphs.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for cp in cell.paragraphs:
                    _push(cp.text)
    return out


def _looks_like_header(line):
    """Strip ornament chars and return a short normalized version, or '' if not header-ish."""
    # Strip markdown stars, parens, colons, brackets, dashes, asterisks
    clean = re.sub(r"[\*_:()\[\]\-\.]", " ", line)
    clean = re.sub(r"\s+", " ", clean).strip().lower()
    # Strip leading "a)", "b)", "part-i", etc.
    clean = re.sub(r"^[a-z]\)\s*", "", clean)
    clean = re.sub(r"^part[\s\-]*[ivx]+\s*", "", clean)
    # Headers are short
    return clean if len(clean) <= 40 else ""


def _classify_header(line):
    """If the line is a section header, return its category (chars/words/sentences/paragraph), else None."""
    norm = _looks_like_header(line)
    if not norm:
        return None
    tokens = set(norm.split())
    for category, kws in SECTION_KEYWORDS.items():
        for kw in kws:
            kw_norm = kw.lower()
            if kw_norm in tokens or kw_norm == norm:
                return category
            # Indic keywords may be one of several tokens
            if any(kw_norm == t for t in tokens):
                return category
    return None


def parse_text_material(path):
    """
    Parse a Dr. Abhishek-style .docx and return:
      {
        'chars':     [str],
        'words':     [str],
        'sentences': {'s1': str, 's2': str, ...},
        'paragraph': str,
      }
    Best-effort across English/Hindi/Marathi/Gujarati layouts.
    """
    paragraphs = _extract_docx_paragraphs(path)
    out = {"chars": [], "words": [], "sentences": {}, "paragraph": ""}

    # Find section header lines
    sections = []  # list of (idx, category)
    seen_cats = set()
    for i, line in enumerate(paragraphs):
        cat = _classify_header(line)
        if cat and cat not in seen_cats:
            sections.append((i, cat))
            seen_cats.add(cat)

    if not sections:
        # Fallback: dump everything into words
        for line in paragraphs:
            for tok in re.split(r"[\s,;।]+", line):
                tok = tok.strip(" *_:.-,()")
                if tok and not re.fullmatch(r"[\W_]+", tok):
                    out["words"].append(tok)
        out["words"] = list(dict.fromkeys(out["words"]))
        return out

    n = len(paragraphs)
    for idx, (start_i, cat) in enumerate(sections):
        end_i = sections[idx + 1][0] if idx + 1 < len(sections) else n
        body = paragraphs[start_i + 1 : end_i]
        # Skip lines inside body that are themselves header-like
        # (e.g. "b) व्यंजन:" in Hindi after "a) स्वर:", or "Words (Take a pause)"
        # right after "Read the words twice" — both got classified as section
        # starts but only the first wins; the rest must not pollute body.)
        body = [ln for ln in body if _classify_header(ln) is None]

        if cat == "chars":
            for line in body:
                # Tokens are usually whitespace-separated single chars
                for tok in re.split(r"[\s,;।]+", line):
                    tok = tok.strip(" *_:.-,()")
                    # Allow ligatures (Indic combining chars can make a "char" 2-3 codepoints)
                    if tok and 1 <= len(tok) <= 4 and not re.fullmatch(r"[\W_]+", tok):
                        out["chars"].append(tok)

        elif cat == "words":
            for line in body:
                # Strip leading bullets and stars
                line = re.sub(r"^[\-\*•]+\s*", "", line)
                for tok in re.split(r"[\s,;।]+", line):
                    tok = tok.strip(" *_:.,()")
                    if tok and not re.fullmatch(r"[\W_]+", tok) and len(tok) >= 1:
                        # Skip tokens that are just numbers or just punctuation
                        out["words"].append(tok)

        elif cat == "sentences":
            # Each non-empty line in this section is one sentence. If the
            # line carries an explicit (sN) prefix, use it as the key;
            # otherwise auto-number as s1, s2, s3, ... in document order.
            # This handles all three observed shapes:
            #   "(s1) Each untimely…"       (English, Marathi)
            #   "- पुलिस को देखते ही…"        (legacy bulleted text)
            #   "ધીરુભાઈ અંબાણી એક…"          (Word native list items, Gujarati/Hindi)
            auto_idx = 0
            used_labels = set()
            for line in body:
                line_stripped = line.strip()
                if not line_stripped:
                    continue
                m = re.match(r"^\s*\*?\s*\(?(s\d+)\)?\s*[:\.]?\s*(.+)$",
                             line_stripped, re.IGNORECASE)
                if m and m.group(2).strip():
                    key = m.group(1).lower()
                    text = m.group(2).strip()
                else:
                    # Strip any stray leading bullet character
                    text = re.sub(r"^[\-\*•]\s+", "", line_stripped).strip()
                    if not text:
                        continue
                    auto_idx += 1
                    # Find next free auto key
                    while f"s{auto_idx}" in used_labels:
                        auto_idx += 1
                    key = f"s{auto_idx}"
                used_labels.add(key)
                # If same key appears more than once, append to existing
                if key in out["sentences"]:
                    out["sentences"][key] += " " + text
                else:
                    out["sentences"][key] = text

        elif cat == "paragraph":
            joined = " ".join(l.strip() for l in body if l.strip())
            out["paragraph"] = re.sub(r"\s+", " ", joined).strip()

    # Dedupe preserving order
    out["chars"] = list(dict.fromkeys(out["chars"]))
    out["words"] = list(dict.fromkeys(out["words"]))
    return out


# ---------- Persistence ----------
def ensure_dirs():
    for d in (WORDLISTS_DIR, TRANSCRIPTS_DIR, OUTPUT_DIR):
        d.mkdir(exist_ok=True, parents=True)


def load_wordlist(iso):
    """Load wordlists/<iso>.txt as a list of strings (preserves order, dedupes)."""
    p = WORDLISTS_DIR / f"{iso}.txt"
    if not p.exists():
        return []
    items = []
    seen = set()
    for line in p.read_text(encoding="utf-8").splitlines():
        w = line.strip()
        if w and w not in seen:
            seen.add(w)
            items.append(w)
    return items


def save_wordlist(iso, words):
    p = WORDLISTS_DIR / f"{iso}.txt"
    p.write_text("\n".join(words) + ("\n" if words else ""), encoding="utf-8")


def save_transcripts_file(iso, sentences, paragraph):
    """Write transcripts/<iso>.txt in the format app.py expects."""
    p = TRANSCRIPTS_DIR / f"{iso}.txt"
    lines = [f"# Reference text for {iso}", ""]
    for key in sorted(sentences.keys(), key=lambda k: (k != "paragraph", k)):
        lines.append(f"[{key}]")
        lines.append(sentences[key])
        lines.append("")
    if paragraph:
        lines.append("[paragraph]")
        lines.append(paragraph)
        lines.append("")
    p.write_text("\n".join(lines), encoding="utf-8")


# ---------- Font helper (for matplotlib waveform labels in Indic) ----------
@lru_cache(maxsize=8)
def get_font(iso):
    fams = {
        "gu": ["Noto Sans Gujarati", "Lohit Gujarati", "Shruti", "Arial Unicode MS"],
        "hi": ["Noto Sans Devanagari", "Lohit Devanagari", "Mangal", "Arial Unicode MS"],
        "mr": ["Noto Sans Devanagari", "Lohit Devanagari", "Mangal", "Arial Unicode MS"],
        "en": ["Arial", "Helvetica", "DejaVu Sans"],
    }.get(iso, ["DejaVu Sans"])
    installed = {f.name for f in fm.fontManager.ttflist}
    for f in fams:
        if f in installed:
            return fm.FontProperties(family=f)
    return None


# ---------- Background audio loader ----------
class AudioLoadThread(QThread):
    done = pyqtSignal(np.ndarray, int, str)
    error = pyqtSignal(str)

    def __init__(self, path):
        super().__init__()
        self.path = path

    def run(self):
        try:
            y, sr = load_audio(self.path)
            self.done.emit(y, sr, self.path)
        except Exception as e:
            traceback.print_exc()
            self.error.emit(str(e))


# ---------- Main window ----------
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Manual Speech Annotator")
        self.resize(1500, 900)

        # State
        self.y = None
        self.sr = None
        self.audio_path = None
        self.wave_t = None
        self.wave_y = None
        self.region_start = None
        self.region_end = None
        self._region_patch = None
        self._drag_anchor = None
        self._drag_moved = False
        self._press_xdata = None
        self._press_button = None
        self._drag_mode = None          # None | "new" | "start" | "end" | "move"
        self._move_region_len = 0.0
        self._move_offset = 0.0
        self._is_updating_scroll = False
        self.segments = []  # [{'label': str, 'start': float, 'end': float}]
        self._selected_seg_idx = None
        self._suppress_spin = False
        self._threads = {}
        self.current_words = []  # current language's wordlist (chars + words combined for picker)
        self.current_chars = []
        self.current_sentences = {}
        self.current_paragraph = ""

        ensure_dirs()
        self._build_ui()
        self._on_lang_changed()  # populate wordlist for default lang

    # ----- UI construction -----
    def _build_ui(self):
        root = QWidget()
        self.setCentralWidget(root)
        outer = QVBoxLayout(root)
        outer.setContentsMargins(8, 8, 8, 8)
        outer.setSpacing(6)

        # ===== Row 1: language + file actions =====
        r1 = QHBoxLayout()
        r1.addWidget(QLabel("Language:"))
        self.cb_lang = QComboBox()
        for disp, iso, name in LANGUAGES:
            self.cb_lang.addItem(disp, (iso, name))
        self.cb_lang.currentIndexChanged.connect(self._on_lang_changed)
        r1.addWidget(self.cb_lang)

        self.btn_load_audio = QPushButton("Load .wav")
        self.btn_load_audio.clicked.connect(self.on_load_audio)
        r1.addWidget(self.btn_load_audio)

        self.btn_import_docx = QPushButton("Import Wordlist from .docx")
        self.btn_import_docx.clicked.connect(self.on_import_docx)
        r1.addWidget(self.btn_import_docx)

        self.btn_load_csv = QPushButton("Load existing CSV")
        self.btn_load_csv.clicked.connect(self.on_load_csv)
        r1.addWidget(self.btn_load_csv)

        self.btn_save_csv = QPushButton("Save CSV")
        self.btn_save_csv.setStyleSheet(
            "QPushButton{font-weight:bold;color:white;background:#16a34a;"
            "border:1px solid #15803d;padding:5px 14px;border-radius:4px}"
            "QPushButton:hover{background:#15803d}"
        )
        self.btn_save_csv.clicked.connect(self.on_save_csv)
        r1.addWidget(self.btn_save_csv)

        self.btn_export_transcripts = QPushButton("Export transcripts/<iso>.txt")
        self.btn_export_transcripts.clicked.connect(self.on_export_transcripts)
        r1.addWidget(self.btn_export_transcripts)

        r1.addStretch()
        outer.addLayout(r1)

        # ===== Status bar =====
        self.lbl_status = QLabel("Pick a language → Load .wav → mark region → choose label → Add Segment")
        self.lbl_status.setStyleSheet("padding:4px 6px;color:#374151;font-size:12px;background:#f3f4f6;border-radius:3px")
        outer.addWidget(self.lbl_status)

        # ===== Main split: waveform (left) / labels + segments (right) =====
        splitter = QSplitter(Qt.Horizontal)

        # ----- Left: waveform + region controls -----
        left = QWidget()
        ll = QVBoxLayout(left)
        ll.setContentsMargins(0, 0, 0, 0)
        ll.setSpacing(4)

        self.fig, self.ax = plt.subplots(figsize=(12, 5))
        self.fig.set_tight_layout(True)
        # Audacity-style: start dark even before any audio is loaded
        self.fig.patch.set_facecolor(self._AUDACITY["bg"])
        self.ax.set_facecolor(self._AUDACITY["bg"])
        for spine in self.ax.spines.values():
            spine.set_color(self._AUDACITY["spine"])
        self.ax.tick_params(colors=self._AUDACITY["axis_text"], which="both")
        self.canvas = FigureCanvas(self.fig)
        self.toolbar = NavigationToolbar(self.canvas, self)
        # Connect interactive events
        self.canvas.mpl_connect("button_press_event", self._on_canvas_press)
        self.canvas.mpl_connect("motion_notify_event", self._on_canvas_motion)
        self.canvas.mpl_connect("button_release_event", self._on_canvas_release)
        self.canvas.mpl_connect("scroll_event", self._on_canvas_scroll)
        ll.addWidget(self.toolbar)
        ll.addWidget(self.canvas)

        # Zoom / view controls row (between waveform and region controls)
        view_row = QHBoxLayout()
        view_row.setSpacing(4)
        view_row.addWidget(QLabel("View:"))

        self.btn_zoom_in = QPushButton("🔍+ Zoom In")
        self.btn_zoom_in.setToolTip("Zoom in 2× around center  (or scroll up on waveform)")
        self.btn_zoom_in.clicked.connect(self.on_zoom_in)
        view_row.addWidget(self.btn_zoom_in)

        self.btn_zoom_out = QPushButton("🔍− Zoom Out")
        self.btn_zoom_out.setToolTip("Zoom out 2× around center  (or scroll down on waveform)")
        self.btn_zoom_out.clicked.connect(self.on_zoom_out)
        view_row.addWidget(self.btn_zoom_out)

        self.btn_fit = QPushButton("Fit")
        self.btn_fit.setToolTip(
            "Fit view to (in order of priority): the selected region,\n"
            "the selected segment, all added segments, or the full audio."
        )
        self.btn_fit.clicked.connect(self.on_fit_view)
        view_row.addWidget(self.btn_fit)

        self.btn_reset_view = QPushButton("Reset View")
        self.btn_reset_view.setToolTip("Show the entire audio file")
        self.btn_reset_view.clicked.connect(self.on_reset_view)
        view_row.addWidget(self.btn_reset_view)

        self.lbl_view = QLabel("")
        self.lbl_view.setStyleSheet("color:#6b7280;font-size:11px;padding-left:8px")
        view_row.addWidget(self.lbl_view)

        view_row.addStretch()
        ll.addLayout(view_row)

        # Audacity-style horizontal scroller / panner row
        scroll_row = QHBoxLayout()
        scroll_row.setSpacing(5)
        scroll_row.addWidget(QLabel("Scroll:"))

        self.btn_back_big = QPushButton("⏪")
        self.btn_back_big.setToolTip("Move view backward by 50% of the visible window")
        self.btn_back_big.clicked.connect(lambda: self.pan_view(-0.50))
        scroll_row.addWidget(self.btn_back_big)

        self.btn_back_small = QPushButton("◀")
        self.btn_back_small.setToolTip("Move view backward by 10% of the visible window")
        self.btn_back_small.clicked.connect(lambda: self.pan_view(-0.10))
        scroll_row.addWidget(self.btn_back_small)

        self.wave_scroll = QSlider(Qt.Horizontal)
        self.wave_scroll.setRange(0, 10000)
        self.wave_scroll.setEnabled(False)
        self.wave_scroll.setToolTip("Drag this to move forward/backward through the waveform after zooming")
        self.wave_scroll.valueChanged.connect(self._on_wave_scroll_changed)
        scroll_row.addWidget(self.wave_scroll, 1)

        self.btn_forward_small = QPushButton("▶")
        self.btn_forward_small.setToolTip("Move view forward by 10% of the visible window")
        self.btn_forward_small.clicked.connect(lambda: self.pan_view(0.10))
        scroll_row.addWidget(self.btn_forward_small)

        self.btn_forward_big = QPushButton("⏩")
        self.btn_forward_big.setToolTip("Move view forward by 50% of the visible window")
        self.btn_forward_big.clicked.connect(lambda: self.pan_view(0.50))
        scroll_row.addWidget(self.btn_forward_big)

        ll.addLayout(scroll_row)

        # Region row
        region_box = QGroupBox("Selected Region")
        rg = QGridLayout(region_box)
        rg.addWidget(QLabel("Start (s):"), 0, 0)
        self.sp_start = QDoubleSpinBox()
        self.sp_start.setRange(0.0, 99999.0)
        self.sp_start.setDecimals(4)
        self.sp_start.setSingleStep(0.05)
        self.sp_start.valueChanged.connect(self._on_spin_changed)
        rg.addWidget(self.sp_start, 0, 1)

        rg.addWidget(QLabel("End (s):"), 0, 2)
        self.sp_end = QDoubleSpinBox()
        self.sp_end.setRange(0.0, 99999.0)
        self.sp_end.setDecimals(4)
        self.sp_end.setSingleStep(0.05)
        self.sp_end.valueChanged.connect(self._on_spin_changed)
        rg.addWidget(self.sp_end, 0, 3)

        self.lbl_dur = QLabel("dur: 0.000 s")
        self.lbl_dur.setStyleSheet("color:#6b7280;font-size:11px")
        rg.addWidget(self.lbl_dur, 0, 4)

        self.btn_play_region = QPushButton("▶ Play Region")
        self.btn_play_region.clicked.connect(self.on_play_region)
        rg.addWidget(self.btn_play_region, 0, 5)

        self.btn_play_all = QPushButton("▶▶ Play All")
        self.btn_play_all.clicked.connect(self.on_play_all)
        rg.addWidget(self.btn_play_all, 0, 6)

        self.btn_stop = QPushButton("■ Stop")
        self.btn_stop.clicked.connect(sd.stop)
        rg.addWidget(self.btn_stop, 0, 7)

        self.btn_clear_region = QPushButton("Clear Region")
        self.btn_clear_region.clicked.connect(self.on_clear_region)
        rg.addWidget(self.btn_clear_region, 0, 8)

        ll.addWidget(region_box)

        # Help line
        hint = QLabel(
            "Tip: <b>Drag empty area</b> to select · <b>drag yellow edges</b> to adjust · "
            "<b>drag inside yellow region</b> to move it · <b>Right-click</b> = set End. "
            "Mouse wheel zooms; the scroller moves through long audio."
        )
        hint.setStyleSheet("color:#6b7280;font-size:11px;padding:2px 4px")
        hint.setTextFormat(Qt.RichText)
        ll.addWidget(hint)

        splitter.addWidget(left)

        # ----- Right: tabs (Labels / Segments) -----
        right_tabs = QTabWidget()

        # --- Labels tab ---
        labels_tab = QWidget()
        lt = QVBoxLayout(labels_tab)

        # Big label-display + Add Segment row
        self.txt_current_label = QLineEdit()
        self.txt_current_label.setPlaceholderText("Selected / typed label appears here")
        f = self.txt_current_label.font()
        f.setPointSize(13)
        f.setBold(True)
        self.txt_current_label.setFont(f)
        self.txt_current_label.setStyleSheet("padding:6px;background:#fefce8;border:2px solid #facc15;border-radius:4px")
        lt.addWidget(QLabel("Current label:"))
        lt.addWidget(self.txt_current_label)

        self.btn_add_seg = QPushButton("➕ Add Segment")
        self.btn_add_seg.setStyleSheet(
            "QPushButton{font-weight:bold;color:white;background:#2563eb;"
            "border:1px solid #1d4ed8;padding:8px 14px;border-radius:4px;font-size:13px}"
            "QPushButton:hover{background:#1d4ed8}"
        )
        self.btn_add_seg.clicked.connect(self.on_add_segment)
        lt.addWidget(self.btn_add_seg)

        # Quick labels (s1..s5, paragraph)
        ql_box = QGroupBox("Quick Labels (sentences / paragraph)")
        qg = QGridLayout(ql_box)
        for i, ql in enumerate(QUICK_LABELS):
            b = QPushButton(ql)
            b.setStyleSheet("QPushButton{background:#dbeafe;font-weight:bold;padding:6px}")
            b.clicked.connect(lambda _, lab=ql: self._set_current_label(lab))
            qg.addWidget(b, i // 3, i % 3)
        lt.addWidget(ql_box)

        # Chars panel (scrollable grid of buttons)
        chars_box = QGroupBox("Characters / Alphabets")
        cv = QVBoxLayout(chars_box)
        self.chars_area = QScrollArea()
        self.chars_area.setWidgetResizable(True)
        self.chars_area.setMaximumHeight(120)
        self.chars_container = QWidget()
        self.chars_grid = QGridLayout(self.chars_container)
        self.chars_grid.setSpacing(2)
        self.chars_area.setWidget(self.chars_container)
        cv.addWidget(self.chars_area)
        lt.addWidget(chars_box)

        # Words list with filter
        words_box = QGroupBox("Words (click to select)")
        wv = QVBoxLayout(words_box)
        self.txt_filter = QLineEdit()
        self.txt_filter.setPlaceholderText("Filter words…")
        self.txt_filter.textChanged.connect(self._refresh_words_list)
        wv.addWidget(self.txt_filter)
        self.words_list = QListWidget()
        self.words_list.itemClicked.connect(
            lambda item: self._set_current_label(item.text())
        )
        # Use a font that handles Indic glyphs nicely
        self.words_list.setStyleSheet("QListWidget{font-size:13px}")
        wv.addWidget(self.words_list)
        lt.addWidget(words_box, 1)

        # Add-new-word row
        add_box = QGroupBox("Add new word to the list")
        ah = QHBoxLayout(add_box)
        self.txt_new_word = QLineEdit()
        self.txt_new_word.setPlaceholderText("Type a new word here…")
        self.txt_new_word.returnPressed.connect(self.on_add_new_word)
        ah.addWidget(self.txt_new_word, 1)
        self.btn_add_word = QPushButton("Add to list")
        self.btn_add_word.clicked.connect(self.on_add_new_word)
        ah.addWidget(self.btn_add_word)
        self.btn_add_word_and_use = QPushButton("Add & Use")
        self.btn_add_word_and_use.setStyleSheet("QPushButton{background:#dcfce7}")
        self.btn_add_word_and_use.clicked.connect(self.on_add_new_word_and_use)
        ah.addWidget(self.btn_add_word_and_use)
        lt.addWidget(add_box)

        right_tabs.addTab(labels_tab, "Labels")

        # --- Segments tab ---
        segs_tab = QWidget()
        sv = QVBoxLayout(segs_tab)
        sv.addWidget(QLabel("Added segments (click to highlight on waveform):"))
        self.seg_list = QListWidget()
        self.seg_list.setStyleSheet(
            "QListWidget{font-family:Consolas,Courier New,monospace;font-size:11px}"
            "QListWidget::item{padding:3px 6px}"
            "QListWidget::item:selected{background:#dbeafe}"
        )
        self.seg_list.itemClicked.connect(self._on_seg_clicked)
        sv.addWidget(self.seg_list, 1)

        seg_btns = QHBoxLayout()
        self.btn_del_seg = QPushButton("🗑 Delete Selected")
        self.btn_del_seg.clicked.connect(self.on_delete_segment)
        seg_btns.addWidget(self.btn_del_seg)
        self.btn_relabel = QPushButton("✏ Relabel from Current")
        self.btn_relabel.clicked.connect(self.on_relabel_segment)
        seg_btns.addWidget(self.btn_relabel)
        self.btn_retime = QPushButton("⏱ Update Times from Region")
        self.btn_retime.clicked.connect(self.on_retime_segment)
        seg_btns.addWidget(self.btn_retime)
        self.btn_load_to_region = QPushButton("Load to Region")
        self.btn_load_to_region.clicked.connect(self.on_load_seg_to_region)
        seg_btns.addWidget(self.btn_load_to_region)
        sv.addLayout(seg_btns)

        right_tabs.addTab(segs_tab, "Segments")

        splitter.addWidget(right_tabs)
        splitter.setSizes([1000, 500])
        outer.addWidget(splitter, 1)

    # ----- Language / wordlist handling -----
    def _current_lang(self):
        return self.cb_lang.currentData()  # (iso, name)

    def _on_lang_changed(self):
        iso, name = self._current_lang()
        # Load persisted wordlist
        words = load_wordlist(iso)
        if not words:
            # Try to auto-import from a matching docx in current dir
            auto_path = self._guess_docx_for_lang(iso)
            if auto_path:
                try:
                    parsed = parse_text_material(auto_path)
                    words = list(dict.fromkeys(parsed["chars"] + parsed["words"]))
                    self.current_chars = parsed["chars"]
                    self.current_sentences = parsed["sentences"]
                    self.current_paragraph = parsed["paragraph"]
                    save_wordlist(iso, words)
                    self._set_status(f"Auto-imported {len(words)} words/chars from {auto_path.name}")
                except Exception as e:
                    log(f"Auto-import failed: {e}")
                    self.current_chars = []
                    self.current_sentences = {}
                    self.current_paragraph = ""
            else:
                self.current_chars = []
                self.current_sentences = {}
                self.current_paragraph = ""
        else:
            # Read chars/sentences/paragraph if we have a matching docx cached, otherwise just empty
            # The chars detection: pick out single-codepoint words from the loaded list as a heuristic
            self.current_chars = [w for w in words if len(w) <= 2]
            self.current_sentences = {}
            self.current_paragraph = ""

        self.current_words = words
        self._refresh_words_list()
        self._refresh_chars_grid()

    def _guess_docx_for_lang(self, iso):
        """Try to find a matching Dr__Abhishek_..._Text_Material.docx in APP_DIR."""
        lang_token = {"en": "English", "hi": "Hindi", "gu": "Gujarati", "mr": "Marathi"}.get(iso, "")
        if not lang_token:
            return None
        for p in APP_DIR.glob(f"*{lang_token}*Text*Material*.docx"):
            return p
        return None

    def _refresh_words_list(self):
        flt = self.txt_filter.text().strip().lower()
        self.words_list.clear()
        for w in self.current_words:
            if flt and flt not in w.lower():
                continue
            self.words_list.addItem(QListWidgetItem(w))

    def _refresh_chars_grid(self):
        # Clear grid
        while self.chars_grid.count():
            item = self.chars_grid.takeAt(0)
            w = item.widget()
            if w:
                w.deleteLater()
        cols = 12
        for i, ch in enumerate(self.current_chars):
            b = QPushButton(ch)
            b.setStyleSheet("QPushButton{background:#ede9fe;padding:4px;font-size:14px;font-weight:bold;min-width:32px}")
            b.clicked.connect(lambda _, c=ch: self._set_current_label(c))
            self.chars_grid.addWidget(b, i // cols, i % cols)
        self.chars_grid.setRowStretch((len(self.current_chars) // cols) + 1, 1)

    def on_import_docx(self):
        iso, _ = self._current_lang()
        path, _ = QFileDialog.getOpenFileName(
            self, "Pick a .docx text-material file for this language",
            str(APP_DIR), "Word Documents (*.docx)"
        )
        if not path:
            return
        try:
            parsed = parse_text_material(path)
            new_words = list(dict.fromkeys(parsed["chars"] + parsed["words"]))
            # Merge with existing wordlist
            existing = load_wordlist(iso)
            merged = list(dict.fromkeys(existing + new_words))
            save_wordlist(iso, merged)
            self.current_words = merged
            self.current_chars = parsed["chars"]
            self.current_sentences = parsed["sentences"]
            self.current_paragraph = parsed["paragraph"]
            self._refresh_words_list()
            self._refresh_chars_grid()
            self._set_status(
                f"Imported {len(parsed['chars'])} chars, {len(parsed['words'])} words, "
                f"{len(parsed['sentences'])} sentences, "
                f"{'1' if parsed['paragraph'] else '0'} paragraph from {Path(path).name}"
            )
        except Exception as e:
            traceback.print_exc()
            QMessageBox.critical(self, "Import error", str(e))

    def on_add_new_word(self):
        word = self.txt_new_word.text().strip()
        if not word:
            return
        iso, _ = self._current_lang()
        if word in self.current_words:
            self._set_status(f"'{word}' is already in the list")
            return
        self.current_words.append(word)
        save_wordlist(iso, self.current_words)
        self.txt_new_word.clear()
        self._refresh_words_list()
        self._set_status(f"Added '{word}' to {iso}.txt ({len(self.current_words)} total)")

    def on_add_new_word_and_use(self):
        word = self.txt_new_word.text().strip()
        if not word:
            return
        self.on_add_new_word()
        self._set_current_label(word)

    def _set_current_label(self, label):
        self.txt_current_label.setText(label)

    # ----- Audio loading -----
    def on_load_audio(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "Load audio", str(APP_DIR),
            "Audio files (*.wav *.mp3 *.flac *.ogg *.m4a)"
        )
        if not path:
            return
        self._set_status(f"Loading {Path(path).name}…")
        self.btn_load_audio.setEnabled(False)
        t = AudioLoadThread(path)
        t.done.connect(self._on_audio_loaded)
        t.error.connect(self._on_audio_err)
        self._threads["audio"] = t
        t.start()

    def _on_audio_loaded(self, y, sr, path):
        self.y = y
        self.sr = sr
        self.audio_path = path
        # No pre-decimation needed: _compute_envelope() computes peak/RMS
        # for the currently-visible range on every redraw.
        # Configure spinboxes to audio duration
        dur = len(y) / sr
        for sp in (self.sp_start, self.sp_end):
            sp.setRange(0.0, dur)
        self.region_start = None
        self.region_end = None
        self._region_patch = None
        self.btn_load_audio.setEnabled(True)
        self._set_status(f"Loaded: {Path(path).name} ({dur:.2f} s, {sr} Hz)")
        self._redraw()

    def _on_audio_err(self, msg):
        self.btn_load_audio.setEnabled(True)
        QMessageBox.critical(self, "Audio load failed", msg)
        self._set_status(f"✗ Audio load failed: {msg}")

    # ----- Waveform drawing (Audacity-style) -----
    # Palette tuned to mimic Audacity's default look on a dark background.
    _AUDACITY = {
        "bg":          "#1f2a36",  # dark navy figure / axes background
        "envelope":    "#5294c4",  # outer peak (min/max) — medium blue
        "rms":         "#a0d8f5",  # inner RMS overlay — bright cyan-blue
        "sample_line": "#5fb4e2",  # connecting line at sample-level zoom
        "sample_dot":  "#cfe9fb",  # individual sample dots at deep zoom
        "zero":        "#3a4d5e",  # center zero line
        "grid":        "#3a4d5e",  # gridlines (same as zero line, subtle)
        "axis_text":   "#c4d4e2",  # tick labels / axis names
        "title":       "#e2ecf5",
        "spine":       "#3a4d5e",
        # Selections / segments
        "region":      "#fde047",  # current selection — Audacity yellow
        "region_edge": "#facc15",
        "seg":         "#22c55e",  # added segments — green
        "seg_sel":     "#fb923c",  # selected segment — orange
        "label_bg":    "#0f1923",  # background of label boxes
        "label_text":  "#fff7ed",
    }

    def _compute_envelope(self, x0, x1, n_pixels=2000):
        """Compute peak (min/max) and RMS envelopes for the visible range.

        Returns:
          {'mode': 'samples', 't': [..], 'y': [..]}        when very zoomed in
          {'mode': 'envelope', 't': [..], 'mins': [..],
           'maxs': [..], 'rms': [..]}                       otherwise
        """
        if self.y is None or self.sr is None:
            return None
        sr = self.sr
        total = len(self.y)
        start = max(0, int(x0 * sr))
        end = min(total, int(x1 * sr))
        if end <= start:
            return None
        visible = self.y[start:end]
        n_visible = end - start

        # If we have fewer samples than pixels, show every individual sample.
        if n_visible <= n_pixels:
            t = (start + np.arange(n_visible)) / sr
            return {"mode": "samples", "t": t, "y": visible}

        # Otherwise reshape into bins of samples_per_bin and compute min/max/RMS.
        samples_per_bin = max(1, n_visible // n_pixels)
        n_bins = n_visible // samples_per_bin
        trimmed = visible[: n_bins * samples_per_bin]
        bins = trimmed.reshape(n_bins, samples_per_bin)
        mins = bins.min(axis=1)
        maxs = bins.max(axis=1)
        # RMS in float64 to avoid overflow on int audio (we already loaded float32)
        rms = np.sqrt(np.mean(bins.astype(np.float64) ** 2, axis=1)).astype(np.float32)
        t = (start + np.arange(n_bins) * samples_per_bin + samples_per_bin / 2.0) / sr
        return {"mode": "envelope", "t": t, "mins": mins, "maxs": maxs, "rms": rms}

    def _redraw(self, xlim=None):
        P = self._AUDACITY
        self.ax.clear()
        # Dark Audacity-style background on both figure and axes
        self.fig.patch.set_facecolor(P["bg"])
        self.ax.set_facecolor(P["bg"])

        if self.y is None:
            self.canvas.draw_idle()
            self._update_view_label()
            return

        max_dur = len(self.y) / self.sr
        if xlim is None:
            xlim = (0.0, max_dur)
        xlim = (max(0.0, xlim[0]), min(max_dur, xlim[1]))
        if xlim[1] <= xlim[0]:
            xlim = (0.0, max_dur)

        # Y range: a touch above the global peak so the wave doesn't kiss the top
        ymax = float(max(np.max(np.abs(self.y)), 1e-6))
        y_top = ymax * 1.1

        env = self._compute_envelope(xlim[0], xlim[1], n_pixels=2200)
        if env is not None:
            if env["mode"] == "samples":
                # Sample-level zoom: stem-and-dot like Audacity's "sample view"
                self.ax.plot(env["t"], env["y"], color=P["sample_line"],
                             lw=1.2, zorder=2)
                # Only show dots if not too many — keeps things crisp
                if len(env["t"]) <= 400:
                    self.ax.scatter(env["t"], env["y"], color=P["sample_dot"],
                                    s=14, zorder=4, edgecolors="none")
            else:
                # Standard view: outer peak envelope + inner RMS overlay
                self.ax.fill_between(
                    env["t"], env["mins"], env["maxs"],
                    color=P["envelope"], linewidth=0, zorder=2,
                )
                self.ax.fill_between(
                    env["t"], -env["rms"], env["rms"],
                    color=P["rms"], linewidth=0, zorder=3,
                )

        # Center zero line
        self.ax.axhline(0, color=P["zero"], lw=0.8, zorder=4)

        # Style: axis text, spines, grid
        self.ax.set_xlim(*xlim)
        self.ax.set_ylim(-y_top, y_top)
        self.ax.grid(axis="x", color=P["grid"], alpha=0.35,
                     linestyle="-", linewidth=0.5, zorder=1)
        self.ax.tick_params(colors=P["axis_text"], which="both", labelsize=9)
        for spine in self.ax.spines.values():
            spine.set_color(P["spine"])
        self.ax.set_xlabel("Time (s)", color=P["axis_text"], fontsize=10)
        self.ax.set_ylabel("Amplitude", color=P["axis_text"], fontsize=10)

        # Existing segments (drawn over the waveform, under selection)
        iso, _ = self._current_lang()
        fp = get_font(iso)
        for i, seg in enumerate(self.segments):
            s, e = seg["start"], seg["end"]
            if e < xlim[0] or s > xlim[1]:
                continue
            is_sel = (i == self._selected_seg_idx)
            color = P["seg_sel"] if is_sel else P["seg"]
            # Subtle background tint
            self.ax.axvspan(s, e, alpha=0.12, color=color, zorder=5)
            # Bottom-of-axes bar marker
            self.ax.plot([s, e], [-y_top * 0.94, -y_top * 0.94],
                         lw=3.5, color=color, solid_capstyle="butt", zorder=6)
            if is_sel:
                kw = {"fontproperties": fp} if fp else {}
                self.ax.text(
                    (s + e) / 2, y_top * 0.78, seg["label"],
                    ha="center", va="bottom", fontsize=10,
                    color=P["label_text"], fontweight="bold",
                    bbox=dict(facecolor=color, alpha=0.85,
                              edgecolor="none", boxstyle="round,pad=0.35"),
                    zorder=8, **kw,
                )

        # Current selection region (Audacity-style yellow translucent box)
        if self.region_start is not None and self.region_end is not None \
                and self.region_end > self.region_start:
            s, e = self.region_start, self.region_end
            self.ax.axvspan(s, e, alpha=0.22, color=P["region"], zorder=5)
            self.ax.axvline(s, color=P["region_edge"], lw=1.2,
                            linestyle="-", zorder=7)
            self.ax.axvline(e, color=P["region_edge"], lw=1.2,
                            linestyle="-", zorder=7)
            # Time labels at top, light text on dark
            self.ax.text(
                s, y_top * 0.98, f"{s:.3f}s", ha="left", va="top",
                fontsize=8, color=P["region_edge"], fontweight="bold",
                bbox=dict(facecolor=P["label_bg"], alpha=0.7,
                          edgecolor="none", boxstyle="round,pad=0.2"),
                zorder=8,
            )
            self.ax.text(
                e, y_top * 0.98, f"{e:.3f}s", ha="right", va="top",
                fontsize=8, color=P["region_edge"], fontweight="bold",
                bbox=dict(facecolor=P["label_bg"], alpha=0.7,
                          edgecolor="none", boxstyle="round,pad=0.2"),
                zorder=8,
            )

        # Title
        title = Path(self.audio_path).name if self.audio_path else "Waveform"
        if self.segments:
            title += f"  │  {len(self.segments)} segments added"
        self.ax.set_title(title, fontsize=10, color=P["title"])

        self.canvas.draw_idle()
        if hasattr(self, "wave_scroll"):
            self._sync_wave_scroll()
        self._update_view_label()

    # ----- Canvas mouse handlers -----
    def _selection_hit_test(self, x):
        """Return how a click should interact with the current region.

        Audacity-like behavior:
        - near left edge  -> drag start
        - near right edge -> drag end
        - inside region   -> move whole selection
        - outside region  -> create a new selection
        """
        if (self.region_start is None or self.region_end is None
                or self.region_end <= self.region_start):
            return "new"
        x0, x1 = self.ax.get_xlim()
        visible = max(0.001, x1 - x0)
        # handle hit zone: 1.5% of visible window, clamped 20–200 ms
        tol = min(0.200, max(0.020, visible * 0.015))
        if abs(x - self.region_start) <= tol:
            return "start"
        if abs(x - self.region_end) <= tol:
            return "end"
        if self.region_start < x < self.region_end:
            return "move"
        return "new"

    def _on_canvas_press(self, ev):
        if ev.inaxes != self.ax or self.y is None or ev.xdata is None:
            return
        # Ignore if a matplotlib pan/zoom mode is active
        if self.toolbar.mode:
            return
        x = max(0.0, min(float(ev.xdata), len(self.y) / self.sr))
        self._press_xdata = x
        self._press_button = ev.button
        self._drag_moved = False
        self._drag_anchor = x
        self._drag_mode = self._selection_hit_test(x) if ev.button == 1 else None
        if self._drag_mode == "move" and self.region_start is not None:
            self._move_region_len = self.region_end - self.region_start
            self._move_offset = x - self.region_start

    def _on_canvas_motion(self, ev):
        if self._press_xdata is None or ev.inaxes != self.ax or ev.xdata is None:
            return
        if self.toolbar.mode or self._press_button != 1:
            return
        dur = len(self.y) / self.sr
        x = max(0.0, min(float(ev.xdata), dur))
        if abs(x - self._press_xdata) <= 0.002:  # >2 ms is a drag
            return
        self._drag_moved = True

        if self._drag_mode == "start":
            end = self.region_end if self.region_end is not None else x
            self._set_region(min(x, end), max(x, end), redraw=True)
        elif self._drag_mode == "end":
            start = self.region_start if self.region_start is not None else x
            self._set_region(min(start, x), max(start, x), redraw=True)
        elif self._drag_mode == "move":
            length = max(0.0, self._move_region_len)
            new_start = x - self._move_offset
            new_start = min(max(0.0, new_start), max(0.0, dur - length))
            self._set_region(new_start, new_start + length, redraw=True)
        else:
            a, b = sorted([self._drag_anchor, x])
            self._set_region(a, b, redraw=True)

    def _on_canvas_release(self, ev):
        if self._press_xdata is None:
            return
        if self.toolbar.mode:
            self._press_xdata = None
            return
        if ev.inaxes != self.ax or ev.xdata is None:
            self._press_xdata = None
            return
        x = float(ev.xdata)
        x = max(0.0, min(x, len(self.y) / self.sr if self.y is not None else x))
        if not self._drag_moved:
            # Treat as click
            if ev.button == 1:
                hit = self._selection_hit_test(x)
                if hit == "start" and self.region_end is not None:
                    self._set_region(x, self.region_end, redraw=True)
                elif hit == "end" and self.region_start is not None:
                    self._set_region(self.region_start, x, redraw=True)
                elif hit == "move":
                    # Clicking inside the region keeps the selection unchanged.
                    self._redraw(xlim=self.ax.get_xlim())
                else:
                    # Left click -> set start, preserving a valid end if possible
                    if self.region_end is not None and x < self.region_end:
                        self._set_region(x, self.region_end, redraw=True)
                    else:
                        self._set_region(x, x, redraw=True)
            elif ev.button == 3:
                # Right click -> set end
                start = self.region_start if (self.region_start is not None and self.region_start < x) else x
                self._set_region(start, x, redraw=True)
        self._press_xdata = None
        self._press_button = None
        self._drag_mode = None
        self._drag_moved = False

    def _on_canvas_scroll(self, ev):
        if self.y is None or ev.inaxes != self.ax:
            return
        # Shift + wheel pans horizontally; normal wheel zooms around cursor.
        if ev.key == "shift":
            self.pan_view(-0.10 if ev.button == "up" else 0.10)
            return
        x0, x1 = self.ax.get_xlim()
        mx = ev.xdata if ev.xdata is not None else (x0 + x1) / 2
        f = 0.7 if ev.button == "up" else 1.42
        new_x0 = max(0.0, mx - (mx - x0) * f)
        new_x1 = min(len(self.y) / self.sr, mx + (x1 - mx) * f)
        if new_x1 - new_x0 < 0.02:
            return
        self._redraw(xlim=(new_x0, new_x1))

    # ----- Zoom / view controls -----
    def _zoom_by(self, factor):
        """factor < 1 zooms in (shrinks visible range), factor > 1 zooms out."""
        if self.y is None:
            return
        x0, x1 = self.ax.get_xlim()
        center = (x0 + x1) / 2.0
        half_width = (x1 - x0) * factor / 2.0
        max_dur = len(self.y) / self.sr
        new_x0 = max(0.0, center - half_width)
        new_x1 = min(max_dur, center + half_width)
        # Cap minimum zoom-in at 20 ms so we don't zoom into nothingness
        if new_x1 - new_x0 < 0.02:
            return
        # Cap maximum zoom-out at full audio
        if (new_x1 - new_x0) >= max_dur:
            new_x0, new_x1 = 0.0, max_dur
        self._redraw(xlim=(new_x0, new_x1))

    def on_zoom_in(self):
        if self.y is None:
            self._set_status("Load an audio file first")
            return
        self._zoom_by(0.5)

    def on_zoom_out(self):
        if self.y is None:
            self._set_status("Load an audio file first")
            return
        self._zoom_by(2.0)

    def on_fit_view(self):
        """Fit view by priority:
          1. The currently-selected (blue) region
          2. The currently-selected segment (clicked in the Segments list)
          3. The bounding box of all added segments
          4. The full audio
        """
        if self.y is None:
            self._set_status("Load an audio file first")
            return
        max_dur = len(self.y) / self.sr

        # 1. Selected region
        if (self.region_start is not None and self.region_end is not None
                and self.region_end > self.region_start):
            pad = max(0.1, (self.region_end - self.region_start) * 0.3)
            x0 = max(0.0, self.region_start - pad)
            x1 = min(max_dur, self.region_end + pad)
            self._redraw(xlim=(x0, x1))
            self._set_status(f"Fit to selected region [{self.region_start:.3f}–{self.region_end:.3f}s]")
            return

        # 2. Selected segment in the list
        if (self._selected_seg_idx is not None
                and 0 <= self._selected_seg_idx < len(self.segments)):
            seg = self.segments[self._selected_seg_idx]
            pad = max(0.1, (seg["end"] - seg["start"]) * 0.3)
            x0 = max(0.0, seg["start"] - pad)
            x1 = min(max_dur, seg["end"] + pad)
            self._redraw(xlim=(x0, x1))
            self._set_status(f"Fit to segment {self._selected_seg_idx + 1}: {seg['label']}")
            return

        # 3. All segments
        if self.segments:
            s0 = min(s["start"] for s in self.segments)
            s1 = max(s["end"] for s in self.segments)
            pad = max(0.5, (s1 - s0) * 0.05)
            x0 = max(0.0, s0 - pad)
            x1 = min(max_dur, s1 + pad)
            self._redraw(xlim=(x0, x1))
            self._set_status(f"Fit to {len(self.segments)} segments")
            return

        # 4. Full audio
        self._redraw(xlim=(0.0, max_dur))
        self._set_status("Fit to full audio")

    def on_reset_view(self):
        if self.y is None:
            self._set_status("Load an audio file first")
            return
        self._redraw(xlim=(0.0, len(self.y) / self.sr))
        self._set_status("View reset to full audio")

    def pan_view(self, fraction):
        """Move the visible waveform window left/right by a fraction of its width."""
        if self.y is None:
            self._set_status("Load an audio file first")
            return
        x0, x1 = self.ax.get_xlim()
        width = x1 - x0
        max_dur = len(self.y) / self.sr
        if width >= max_dur:
            return
        shift = width * float(fraction)
        new_x0 = min(max(0.0, x0 + shift), max(0.0, max_dur - width))
        new_x1 = new_x0 + width
        self._redraw(xlim=(new_x0, new_x1))

    def _sync_wave_scroll(self):
        """Keep the bottom scroller aligned with the current visible x-range."""
        if self.y is None:
            self.wave_scroll.setEnabled(False)
            return
        max_dur = len(self.y) / self.sr
        x0, x1 = self.ax.get_xlim()
        width = max(0.0, x1 - x0)
        can_scroll = max_dur > 0 and width < (max_dur - 1e-6)
        self.wave_scroll.setEnabled(can_scroll)
        self._is_updating_scroll = True
        if can_scroll:
            denom = max(max_dur - width, 1e-9)
            val = int(round((max(0.0, min(x0, denom)) / denom) * 10000))
            self.wave_scroll.setValue(max(0, min(10000, val)))
        else:
            self.wave_scroll.setValue(0)
        self._is_updating_scroll = False

    def _on_wave_scroll_changed(self, value):
        if self._is_updating_scroll or self.y is None:
            return
        max_dur = len(self.y) / self.sr
        x0, x1 = self.ax.get_xlim()
        width = x1 - x0
        if width >= max_dur:
            return
        new_x0 = (float(value) / 10000.0) * (max_dur - width)
        self._redraw(xlim=(new_x0, new_x0 + width))

    def _update_view_label(self):
        """Update the tiny label next to the zoom buttons that shows the
        current visible range (handy when you've zoomed in and want to
        confirm where you are)."""
        if self.y is None or self.ax is None:
            self.lbl_view.setText("")
            return
        x0, x1 = self.ax.get_xlim()
        total = len(self.y) / self.sr
        visible = max(0.0, x1 - x0)
        zoom_ratio = total / visible if visible > 0 else 1.0
        self.lbl_view.setText(
            f"visible: {x0:.2f}–{x1:.2f}s  ({visible:.2f}s, {zoom_ratio:.1f}× zoom)"
        )

    # ----- Region state -----
    def _set_region(self, s, e, redraw=True):
        s = max(0.0, float(s))
        e = max(0.0, float(e))
        if self.y is not None:
            dur = len(self.y) / self.sr
            s = min(s, dur)
            e = min(e, dur)
        if e < s:
            s, e = e, s
        self.region_start = s
        self.region_end = e
        # Sync spinboxes without triggering callback storm
        self._suppress_spin = True
        self.sp_start.setValue(s)
        self.sp_end.setValue(e)
        self._suppress_spin = False
        self.lbl_dur.setText(f"dur: {(e - s):.3f} s")
        if redraw:
            self._redraw(xlim=self.ax.get_xlim())

    def _on_spin_changed(self, _v=None):
        if self._suppress_spin:
            return
        s = self.sp_start.value()
        e = self.sp_end.value()
        self._set_region(s, e, redraw=True)

    def on_clear_region(self):
        self.region_start = None
        self.region_end = None
        self._suppress_spin = True
        self.sp_start.setValue(0.0)
        self.sp_end.setValue(0.0)
        self._suppress_spin = False
        self.lbl_dur.setText("dur: 0.000 s")
        self._redraw(xlim=self.ax.get_xlim() if self.y is not None else None)

    # ----- Playback -----
    def on_play_region(self):
        if self.y is None:
            return
        if self.region_start is None or self.region_end is None or self.region_end <= self.region_start:
            self._set_status("No region selected to play")
            return
        s = int(self.region_start * self.sr)
        e = int(self.region_end * self.sr)
        sd.stop()
        sd.play(self.y[s:e], self.sr)

    def on_play_all(self):
        if self.y is None:
            return
        sd.stop()
        sd.play(self.y, self.sr)

    # ----- Segment list -----
    def on_add_segment(self):
        if self.y is None:
            QMessageBox.warning(self, "", "Load an audio file first.")
            return
        if self.region_start is None or self.region_end is None or self.region_end <= self.region_start:
            QMessageBox.warning(self, "", "Mark a region first (drag, click, or type Start/End).")
            return
        label = self.txt_current_label.text().strip()
        if not label:
            QMessageBox.warning(self, "", "Pick or type a label first.")
            return
        seg = {"label": label, "start": float(self.region_start), "end": float(self.region_end)}
        self.segments.append(seg)
        self._refresh_seg_list()
        self._set_status(f"Added: {label} [{seconds_to_hhmmss_ms(seg['start'])}–{seconds_to_hhmmss_ms(seg['end'])}]  (total: {len(self.segments)})")
        self._redraw(xlim=self.ax.get_xlim())

    def _refresh_seg_list(self):
        self.seg_list.clear()
        for i, seg in enumerate(self.segments):
            start_ts = seconds_to_hhmmss_ms(seg["start"])
            end_ts = seconds_to_hhmmss_ms(seg["end"])
            txt = f"{i+1:03d}  {seg['label']:<20}  {start_ts} → {end_ts}  (dur {seg['end']-seg['start']:.3f}s)"
            self.seg_list.addItem(QListWidgetItem(txt))
        if self._selected_seg_idx is not None and 0 <= self._selected_seg_idx < len(self.segments):
            self.seg_list.setCurrentRow(self._selected_seg_idx)

    def _on_seg_clicked(self, item):
        idx = self.seg_list.row(item)
        if 0 <= idx < len(self.segments):
            self._selected_seg_idx = idx
            seg = self.segments[idx]
            # Zoom waveform to show this segment with some padding
            pad = max(0.5, (seg["end"] - seg["start"]) * 0.5)
            x0 = max(0.0, seg["start"] - pad)
            x1 = min(len(self.y) / self.sr if self.y is not None else seg["end"] + pad, seg["end"] + pad)
            self._redraw(xlim=(x0, x1))

    def on_delete_segment(self):
        idx = self.seg_list.currentRow()
        if idx < 0 or idx >= len(self.segments):
            return
        seg = self.segments.pop(idx)
        self._selected_seg_idx = None
        self._refresh_seg_list()
        self._set_status(f"Deleted: {seg['label']} [{seconds_to_hhmmss_ms(seg['start'])}–{seconds_to_hhmmss_ms(seg['end'])}]")
        self._redraw(xlim=self.ax.get_xlim() if self.y is not None else None)

    def on_relabel_segment(self):
        idx = self.seg_list.currentRow()
        if idx < 0 or idx >= len(self.segments):
            return
        new_label = self.txt_current_label.text().strip()
        if not new_label:
            QMessageBox.warning(self, "", "Type or pick a label first (top right).")
            return
        old = self.segments[idx]["label"]
        self.segments[idx]["label"] = new_label
        self._refresh_seg_list()
        self._set_status(f"Relabeled: '{old}' → '{new_label}'")
        self._redraw(xlim=self.ax.get_xlim())

    def on_retime_segment(self):
        idx = self.seg_list.currentRow()
        if idx < 0 or idx >= len(self.segments):
            return
        if self.region_start is None or self.region_end is None or self.region_end <= self.region_start:
            QMessageBox.warning(self, "", "Mark a new region first.")
            return
        self.segments[idx]["start"] = float(self.region_start)
        self.segments[idx]["end"] = float(self.region_end)
        self._refresh_seg_list()
        self._set_status(f"Updated times for segment {idx+1}")
        self._redraw(xlim=self.ax.get_xlim())

    def on_load_seg_to_region(self):
        idx = self.seg_list.currentRow()
        if idx < 0 or idx >= len(self.segments):
            return
        seg = self.segments[idx]
        self._set_current_label(seg["label"])
        self._set_region(seg["start"], seg["end"], redraw=True)

    # ----- Save/Load CSV -----
    def on_save_csv(self):
        if not self.segments:
            QMessageBox.warning(self, "", "No segments to save yet.")
            return
        # Default filename: <audio-name>.csv in annotations/
        default = "annotations.csv"
        if self.audio_path:
            default = Path(self.audio_path).stem + ".csv"
        default_path = str(OUTPUT_DIR / default)
        path, _ = QFileDialog.getSaveFileName(
            self, "Save segments to CSV", default_path, "CSV (*.csv)"
        )
        if not path:
            return
        try:
            with open(path, "w", newline="", encoding="utf-8-sig") as f:
                w = csv.writer(f)
                # No header — app.py reads with header=None
                for seg in self.segments:
                    w.writerow([seg["label"], seconds_to_hhmmss_ms(seg["start"]), seconds_to_hhmmss_ms(seg["end"])])
            self._set_status(f"✓ Saved {len(self.segments)} segments → {path}")
        except Exception as e:
            traceback.print_exc()
            QMessageBox.critical(self, "Save error", str(e))

    def on_load_csv(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "Load existing GT CSV", str(OUTPUT_DIR), "CSV (*.csv)"
        )
        if not path:
            return
        try:
            loaded = []
            with open(path, "r", encoding="utf-8-sig", newline="") as f:
                reader = csv.reader(f)
                for row in reader:
                    if len(row) < 3:
                        continue
                    label = str(row[0]).strip()
                    try:
                        s = hhmmss_ms_to_seconds(row[1])
                        e = hhmmss_ms_to_seconds(row[2])
                    except ValueError:
                        continue
                    if not label or e <= s:
                        continue
                    loaded.append({"label": label, "start": s, "end": e})
            if not loaded:
                QMessageBox.warning(self, "", "No valid rows found in that CSV.")
                return
            self.segments = loaded
            self._selected_seg_idx = None
            self._refresh_seg_list()
            self._set_status(f"Loaded {len(loaded)} segments from {Path(path).name}")
            self._redraw(xlim=self.ax.get_xlim() if self.y is not None else None)
        except Exception as e:
            traceback.print_exc()
            QMessageBox.critical(self, "Load error", str(e))

    def on_export_transcripts(self):
        iso, _ = self._current_lang()
        if not self.current_sentences and not self.current_paragraph:
            # No cached transcripts; offer to import from docx now
            ret = QMessageBox.question(
                self, "No transcripts loaded",
                "No sentences/paragraph in memory. Import them from a .docx now?"
            )
            if ret == QMessageBox.Yes:
                self.on_import_docx()
                if not self.current_sentences and not self.current_paragraph:
                    return
            else:
                return
        try:
            save_transcripts_file(iso, self.current_sentences, self.current_paragraph)
            p = TRANSCRIPTS_DIR / f"{iso}.txt"
            self._set_status(
                f"✓ Wrote {len(self.current_sentences)} sentences"
                f"{' + paragraph' if self.current_paragraph else ''} → {p}"
            )
        except Exception as e:
            traceback.print_exc()
            QMessageBox.critical(self, "Export error", str(e))

    # ----- Status -----
    def _set_status(self, msg):
        self.lbl_status.setText(msg)
        log(msg)


def main():
    ensure_dirs()
    print("=" * 60)
    print("  Manual Speech Annotator")
    print(f"  Wordlists dir   : {WORDLISTS_DIR}")
    print(f"  Transcripts dir : {TRANSCRIPTS_DIR}")
    print(f"  Output (CSV) dir: {OUTPUT_DIR}")
    print(f"  python-docx     : {'OK' if _DOCX else 'MISSING (pip install python-docx)'}")
    print(f"  soundfile       : {'OK' if _SF else 'using librosa fallback'}")
    print("=" * 60)
    app = QApplication(sys.argv)
    app.setStyleSheet(
        "QPushButton{padding:5px 10px;border:1px solid #d1d5db;"
        "border-radius:4px;background:#f9fafb;font-size:12px}"
        "QPushButton:hover{background:#e5e7eb}"
        "QPushButton:disabled{color:#9ca3af;background:#f3f4f6}"
        "QComboBox,QSpinBox,QDoubleSpinBox,QLineEdit{padding:4px 6px;font-size:12px}"
    )
    w = MainWindow()
    w.show()
    sys.exit(app.exec_())


if __name__ == "__main__":
    main()